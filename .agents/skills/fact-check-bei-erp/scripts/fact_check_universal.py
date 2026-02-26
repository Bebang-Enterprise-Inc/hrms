#!/usr/bin/env python3
from __future__ import annotations

"""
Universal GLM-5 Fact Checker.

Verifies claims from DECISIONS.md files, business documents, legal contracts,
and tabular files against mixed-format source evidence.
"""

import argparse
import csv
import io
import json
import os
import re
import subprocess
import sys
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Optional


if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")


ZAI_ENDPOINT = "https://api.z.ai/api/coding/paas/v4/chat/completions"
DEFAULT_PARALLEL_WORKERS = 3
MAX_PARALLEL_WORKERS = 3
MAX_API_RETRIES = 6
SUPPORTED_EXTENSIONS = {".md", ".txt", ".csv", ".json", ".xlsx", ".xls", ".docx", ".pdf"}


@dataclass
class Claim:
    claim_id: str
    text: str
    locator: str
    source_ref: str = ""
    meta: dict | None = None


@dataclass
class SourceChunk:
    source: str
    locator: str
    text: str


def _safe_text(value: object) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    return re.sub(r"\s+", " ", text)


def get_api_key(
    env_var: str = "ZAI_API_KEY",
    doppler_project: Optional[str] = None,
    doppler_config: Optional[str] = None,
) -> str:
    key = os.getenv(env_var, "").strip()
    if key:
        return key

    if doppler_project and doppler_config:
        try:
            result = subprocess.run(
                [
                    "doppler",
                    "secrets",
                    "get",
                    env_var,
                    "--plain",
                    "--project",
                    doppler_project,
                    "--config",
                    doppler_config,
                ],
                capture_output=True,
                text=True,
                check=False,
            )
            key = result.stdout.strip()
            if key:
                return key
        except Exception:
            pass

    print(
        f"ERROR: Missing API key. Set {env_var} or provide Doppler project/config.",
        file=sys.stderr,
    )
    sys.exit(1)


def call_glm(
    api_key: str,
    messages: list[dict],
    model: str = "glm-5",
    max_tokens: int = 4096,
    temperature: float = 0.0,
) -> Optional[str]:
    import requests

    for attempt in range(MAX_API_RETRIES):
        try:
            resp = requests.post(
                ZAI_ENDPOINT,
                headers={
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {api_key}",
                },
                json={
                    "model": model,
                    "messages": messages,
                    "max_tokens": max_tokens,
                    "temperature": temperature,
                },
                timeout=180,
            )
            if resp.status_code == 200:
                data = resp.json()
                return data["choices"][0]["message"]["content"]
            if resp.status_code in (429, 500, 502, 503, 504):
                retry_after = resp.headers.get("Retry-After", "").strip()
                if retry_after.isdigit():
                    wait = int(retry_after)
                else:
                    wait = min(90, (2 ** attempt) + 2)
                if attempt < MAX_API_RETRIES - 1:
                    print(
                        f"  API {resp.status_code}, retrying in {wait}s (attempt {attempt + 1}/{MAX_API_RETRIES})...",
                        file=sys.stderr,
                    )
                    time.sleep(wait)
                    continue
                print(
                    f"  API {resp.status_code} after {MAX_API_RETRIES} attempts: {resp.text[:300]}",
                    file=sys.stderr,
                )
                return None
            print(f"  API error {resp.status_code}: {resp.text[:300]}", file=sys.stderr)
            return None
        except Exception as exc:
            print(f"  Request error: {exc}", file=sys.stderr)
            if attempt < MAX_API_RETRIES - 1:
                wait = min(30, attempt + 2)
                print(
                    f"  Retrying after request error in {wait}s (attempt {attempt + 1}/{MAX_API_RETRIES})...",
                    file=sys.stderr,
                )
                time.sleep(wait)
                continue
            return None
    return None


def _extract_json_object(text: str) -> Optional[dict]:
    text = text.strip()
    decoder = json.JSONDecoder()
    try:
        obj = json.loads(text)
        if isinstance(obj, dict):
            return obj
    except Exception:
        pass

    for i, ch in enumerate(text):
        if ch != "{":
            continue
        try:
            obj, _end = decoder.raw_decode(text[i:])
            if isinstance(obj, dict) and "verdict" in obj:
                return obj
        except Exception:
            continue
    return None


def _validate_result(obj: dict) -> tuple[bool, dict]:
    allowed_verdicts = {
        "SUPPORTED",
        "PARTIAL",
        "NOT_FOUND",
        "CONTRADICTED",
        "INSUFFICIENT_CONTEXT",
        "ERROR",
    }
    allowed_mismatch = {
        "NONE",
        "NUMBER",
        "DATE",
        "ENTITY",
        "ATTRIBUTION",
        "LEGAL_INTERPRETATION",
        "SCOPE",
        "OTHER",
    }

    verdict = str(obj.get("verdict", "")).upper()
    if verdict not in allowed_verdicts:
        return False, {}

    try:
        conf = float(obj.get("confidence", 0.0))
    except Exception:
        conf = 0.0
    conf = max(0.0, min(1.0, conf))

    mismatch = str(obj.get("mismatch_type", "OTHER")).upper()
    if mismatch not in allowed_mismatch:
        mismatch = "OTHER"

    citations = obj.get("citations", [])
    clean_citations = []
    if isinstance(citations, list):
        for c in citations[:5]:
            if not isinstance(c, dict):
                continue
            clean_citations.append(
                {
                    "source": _safe_text(c.get("source", "")),
                    "locator": _safe_text(c.get("locator", "")),
                    "evidence": _safe_text(c.get("evidence", ""))[:260],
                }
            )

    checks = obj.get("checks", [])
    clean_checks = []
    if isinstance(checks, list):
        clean_checks = [_safe_text(v)[:120] for v in checks[:8] if _safe_text(v)]

    return True, {
        "verdict": verdict,
        "confidence": conf,
        "reason": _safe_text(obj.get("reason", "")),
        "mismatch_type": mismatch,
        "citations": clean_citations,
        "checks": clean_checks,
    }


def _tokenize(text: str) -> set[str]:
    return set(re.findall(r"\b[a-zA-Z0-9_%-]{3,}\b", text.lower()))


def _extract_numbers(text: str) -> set[str]:
    return set(re.findall(r"-?\d+(?:[.,]\d+)?", text))


def _extract_dates(text: str) -> set[str]:
    patterns = [
        r"\b\d{4}-\d{2}-\d{2}\b",
        r"\b\d{1,2}/\d{1,2}/\d{2,4}\b",
        r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{1,2}(?:,\s*\d{4})?\b",
    ]
    out: set[str] = set()
    lower = text.lower()
    for p in patterns:
        out.update(re.findall(p, lower, flags=re.IGNORECASE))
    return out


def _score_chunk(claim_text: str, chunk_text: str) -> float:
    claim_tokens = _tokenize(claim_text)
    if not claim_tokens:
        return 0.0
    chunk_tokens = _tokenize(chunk_text)
    overlap = len(claim_tokens & chunk_tokens) / max(len(claim_tokens), 1)

    claim_nums = _extract_numbers(claim_text)
    chunk_nums = _extract_numbers(chunk_text)
    num_score = 0.0 if not claim_nums else len(claim_nums & chunk_nums) / len(claim_nums)

    claim_dates = _extract_dates(claim_text)
    chunk_dates = _extract_dates(chunk_text)
    date_score = 0.0 if not claim_dates else len(claim_dates & chunk_dates) / len(claim_dates)

    return (0.7 * overlap) + (0.2 * num_score) + (0.1 * date_score)


def _iter_text_paragraphs(text: str) -> Iterable[tuple[str, str]]:
    lines = text.splitlines()
    buf: list[str] = []
    start = 1
    for idx, line in enumerate(lines, start=1):
        if line.strip():
            if not buf:
                start = idx
            buf.append(line.strip())
        elif buf:
            yield f"lines {start}-{idx - 1}", " ".join(buf)
            buf = []
    if buf:
        yield f"lines {start}-{len(lines)}", " ".join(buf)


def _chunk_long_text(locator_prefix: str, text: str, chunk_chars: int = 1500) -> list[tuple[str, str]]:
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        return []
    if len(text) <= chunk_chars:
        return [(locator_prefix, text)]
    chunks = []
    start = 0
    idx = 1
    while start < len(text):
        end = min(len(text), start + chunk_chars)
        chunks.append((f"{locator_prefix} chunk {idx}", text[start:end]))
        idx += 1
        if end == len(text):
            break
        start = max(start + chunk_chars - 250, end)
    return chunks


def _load_md_txt_json(path: Path) -> list[SourceChunk]:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return []
    chunks = []
    for locator, paragraph in _iter_text_paragraphs(text):
        for loc, txt in _chunk_long_text(locator, paragraph):
            chunks.append(SourceChunk(source=path.name, locator=loc, text=txt))
    return chunks


def _load_csv(path: Path, max_rows: int) -> list[SourceChunk]:
    chunks: list[SourceChunk] = []
    try:
        with path.open("r", encoding="utf-8", errors="replace", newline="") as f:
            reader = csv.DictReader(f)
            for row_idx, row in enumerate(reader, start=2):
                if row_idx - 1 > max_rows:
                    break
                pairs = []
                for col, val in row.items():
                    v = _safe_text(val)
                    if v:
                        pairs.append(f"{col}={v}")
                if pairs:
                    text = "; ".join(pairs)
                    chunks.append(SourceChunk(source=path.name, locator=f"row {row_idx}", text=text[:2000]))
    except Exception:
        return []
    return chunks


def _load_xlsx(path: Path, max_rows: int) -> list[SourceChunk]:
    chunks: list[SourceChunk] = []
    try:
        import pandas as pd
    except Exception:
        return chunks

    try:
        sheets = pd.read_excel(str(path), sheet_name=None)
    except Exception:
        return chunks

    for sheet_name, df in sheets.items():
        df = df.fillna("")
        for idx, row in df.iterrows():
            row_num = int(idx) + 2
            if row_num - 1 > max_rows:
                break
            pairs = []
            for col in df.columns:
                v = _safe_text(row[col])
                if v:
                    pairs.append(f"{col}={v}")
            if pairs:
                text = "; ".join(pairs)
                chunks.append(
                    SourceChunk(
                        source=path.name,
                        locator=f"{sheet_name}!row {row_num}",
                        text=text[:2200],
                    )
                )
    return chunks


def _load_docx(path: Path) -> list[SourceChunk]:
    chunks: list[SourceChunk] = []
    try:
        from docx import Document
    except Exception:
        return chunks
    try:
        doc = Document(str(path))
    except Exception:
        return chunks

    pidx = 0
    for p in doc.paragraphs:
        text = _safe_text(p.text)
        if not text:
            continue
        pidx += 1
        for loc, txt in _chunk_long_text(f"paragraph {pidx}", text):
            chunks.append(SourceChunk(source=path.name, locator=loc, text=txt))

    tidx = 0
    for table in doc.tables:
        tidx += 1
        for ridx, row in enumerate(table.rows, start=1):
            vals = [_safe_text(cell.text) for cell in row.cells]
            vals = [v for v in vals if v]
            if vals:
                text = " | ".join(vals)
                chunks.append(SourceChunk(source=path.name, locator=f"table {tidx} row {ridx}", text=text[:2200]))

    return chunks


def _load_pdf(path: Path) -> list[SourceChunk]:
    chunks: list[SourceChunk] = []
    try:
        from pypdf import PdfReader
    except Exception:
        return chunks
    try:
        reader = PdfReader(str(path))
    except Exception:
        return chunks

    for pnum, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        text = re.sub(r"\s+", " ", text).strip()
        if not text:
            continue
        for loc, txt in _chunk_long_text(f"page {pnum}", text, chunk_chars=1700):
            chunks.append(SourceChunk(source=path.name, locator=loc, text=txt))
    return chunks


def load_source_chunks(
    source_dir: str,
    max_files: int = 300,
    max_chunks: int = 4000,
    max_rows_per_table: int = 800,
) -> list[SourceChunk]:
    root = Path(source_dir)
    if not root.exists():
        print(f"ERROR: Source directory not found: {source_dir}", file=sys.stderr)
        sys.exit(1)

    chunks: list[SourceChunk] = []
    file_count = 0
    skip_tokens = ("GLM_AUDIT", "NLI_AUDIT", "manual_verification", "AUDIT_REPORT")

    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        if any(token in path.name for token in skip_tokens):
            continue

        file_count += 1
        if file_count > max_files:
            print(f"  Source file cap reached ({max_files}); skipping remaining files.")
            break

        ext = path.suffix.lower()
        rel_source = str(path.relative_to(root)).replace("\\", "/")
        loaded: list[SourceChunk] = []
        if ext in (".md", ".txt", ".json"):
            loaded = _load_md_txt_json(path)
        elif ext == ".csv":
            loaded = _load_csv(path, max_rows=max_rows_per_table)
        elif ext in (".xlsx", ".xls"):
            loaded = _load_xlsx(path, max_rows=max_rows_per_table)
        elif ext == ".docx":
            loaded = _load_docx(path)
        elif ext == ".pdf":
            loaded = _load_pdf(path)

        if not loaded:
            continue

        for c in loaded:
            c.source = rel_source
            chunks.append(c)
            if len(chunks) >= max_chunks:
                print(f"  Source chunk cap reached ({max_chunks}); truncating.")
                return chunks

    return chunks


def _looks_like_verifiable_sentence(sentence: str, mode: str) -> bool:
    s = sentence.strip()
    if len(s) < 16:
        return False

    number_like = bool(re.search(r"\d|%|[$€£¥₱]", s))
    date_like = bool(re.search(r"\b\d{4}-\d{2}-\d{2}\b|\b\d{1,2}/\d{1,2}/\d{2,4}\b", s))
    citation_like = bool(re.search(r"\b(section|clause|article|schedule|appendix)\b", s, flags=re.IGNORECASE))
    named_entity_like = len(re.findall(r"\b[A-Z][a-z]{2,}\b", s)) >= 2
    legal_like = bool(
        re.search(
            r"\b(shall|must|may not|governing law|effective date|termination|indemnif|liability|confidential)\b",
            s,
            flags=re.IGNORECASE,
        )
    )

    if mode == "contract":
        return legal_like or date_like or number_like or citation_like or named_entity_like
    if mode == "dataset":
        return number_like or date_like
    return number_like or date_like or citation_like or named_entity_like


def _sentences_from_text(text: str) -> list[str]:
    text = re.sub(r"\n{2,}", "\n", text)
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    sentences: list[str] = []
    for line in lines:
        parts = re.split(r"(?<=[.!?])\s+(?=[A-Z0-9$€£¥₱(])", line)
        for p in parts:
            p = p.strip()
            if p:
                sentences.append(p)
    return sentences


def extract_decision_claims(decisions_file: str, max_claims: int) -> list[Claim]:
    text = Path(decisions_file).read_text(encoding="utf-8", errors="replace")
    lines = text.splitlines()
    claims: list[Claim] = []
    section = ""

    for idx, line in enumerate(lines, start=1):
        if line.startswith("## "):
            section = line.strip("# ").strip()
            continue
        m = re.match(r"^\|\s*([A-Z]+-\d+)\s*\|(.+)$", line)
        if not m:
            continue
        claim_id = m.group(1)
        cols = [c.strip() for c in m.group(2).split("|")]
        if len(cols) < 5:
            continue
        decision = cols[0]
        value = cols[1] if len(cols) > 1 else ""
        confirmed_by = cols[2] if len(cols) > 2 else ""
        date = cols[3] if len(cols) > 3 else ""
        source_ref = cols[4] if len(cols) > 4 else ""
        claims.append(
            Claim(
                claim_id=claim_id,
                text=f"{decision} {value}".strip(),
                locator=f"line {idx}",
                source_ref=source_ref,
                meta={"section": section, "confirmed_by": confirmed_by, "date": date, "raw": line.strip()},
            )
        )
        if len(claims) >= max_claims:
            break
    return claims


def _extract_text_from_target(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in (".md", ".txt", ".json", ".csv"):
        return path.read_text(encoding="utf-8", errors="replace")
    if ext in (".xlsx", ".xls"):
        chunks = _load_xlsx(path, max_rows=800)
        return "\n".join(f"{c.locator}: {c.text}" for c in chunks)
    if ext == ".docx":
        chunks = _load_docx(path)
        return "\n".join(f"{c.locator}: {c.text}" for c in chunks)
    if ext == ".pdf":
        chunks = _load_pdf(path)
        return "\n".join(f"{c.locator}: {c.text}" for c in chunks)
    try:
        from markitdown import MarkItDown

        md = MarkItDown()
        return md.convert(str(path)).text_content
    except Exception:
        return path.read_text(encoding="utf-8", errors="replace")


def extract_generic_claims(target_file: str, mode: str, max_claims: int) -> list[Claim]:
    text = _extract_text_from_target(Path(target_file))
    sentences = _sentences_from_text(text)
    claims: list[Claim] = []
    for idx, sentence in enumerate(sentences, start=1):
        if not _looks_like_verifiable_sentence(sentence, mode=mode):
            continue
        claims.append(Claim(claim_id=f"C-{idx:04d}", text=sentence, locator=f"sentence {idx}"))
        if len(claims) >= max_claims:
            break
    return claims


def select_top_chunks(claim: Claim, chunks: list[SourceChunk], top_k: int) -> list[SourceChunk]:
    scored: list[tuple[float, SourceChunk]] = []
    for chunk in chunks:
        score = _score_chunk(claim.text, chunk.text)
        if score > 0.03:
            scored.append((score, chunk))
    scored.sort(key=lambda t: t[0], reverse=True)
    return [c for _, c in scored[:top_k]]


def build_prompt(mode: str, claim: Claim, contexts: list[SourceChunk], numeric_tolerance: float) -> str:
    evidence_blocks = []
    for ctx in contexts:
        evidence_blocks.append(f"SOURCE={ctx.source}\nLOCATOR={ctx.locator}\nTEXT={ctx.text[:2200]}")
    evidence = "\n\n---\n\n".join(evidence_blocks)

    mode_note = ""
    if mode == "contract":
        mode_note = (
            "Contract review mode: enforce defined-term consistency, party attribution, "
            "obligation polarity (shall/must/may), dates, and clause scope."
        )
    elif mode == "dataset":
        mode_note = (
            f"Dataset review mode: prioritize numeric/date consistency; allow +/- {numeric_tolerance} "
            "absolute tolerance only when numbers are otherwise aligned by unit/context."
        )

    src_hint = f"Cited source field (if present): {claim.source_ref}" if claim.source_ref else ""
    meta_hint = f"Metadata: {json.dumps(claim.meta, ensure_ascii=True)}" if claim.meta else ""

    return (
        "You are a strict fact checker.\n"
        f"{mode_note}\n"
        f"CLAIM_ID: {claim.claim_id}\n"
        f"CLAIM_LOCATOR: {claim.locator}\n"
        f"CLAIM_TEXT: {claim.text}\n"
        f"{src_hint}\n"
        f"{meta_hint}\n\n"
        "EVIDENCE_SNIPPETS:\n"
        f"{evidence}\n\n"
        "Return JSON only (no markdown):\n"
        "{"
        "\"verdict\":\"SUPPORTED|PARTIAL|NOT_FOUND|CONTRADICTED|INSUFFICIENT_CONTEXT\","
        "\"confidence\":0.0,"
        "\"reason\":\"short reason\","
        "\"mismatch_type\":\"NONE|NUMBER|DATE|ENTITY|ATTRIBUTION|LEGAL_INTERPRETATION|SCOPE|OTHER\","
        "\"citations\":[{\"source\":\"...\",\"locator\":\"...\",\"evidence\":\"...\"}],"
        "\"checks\":[\"check 1\", \"check 2\"]"
        "}\n"
        "Rules: prefer NOT_FOUND when evidence is absent; use CONTRADICTED when evidence conflicts."
    )


def verify_claim(
    api_key: str,
    model: str,
    mode: str,
    claim: Claim,
    source_chunks: list[SourceChunk],
    top_k: int,
    numeric_tolerance: float,
) -> dict:
    contexts = select_top_chunks(claim, source_chunks, top_k=top_k)
    if not contexts:
        return {
            "verdict": "NOT_FOUND",
            "confidence": 0.0,
            "reason": "No relevant source chunks found.",
            "mismatch_type": "OTHER",
            "citations": [],
            "checks": [],
        }

    response = call_glm(
        api_key=api_key,
        messages=[{"role": "user", "content": build_prompt(mode, claim, contexts, numeric_tolerance)}],
        model=model,
        max_tokens=8192 if "glm-5" in model else 2048,
        temperature=0.0,
    )
    if not response:
        return {
            "verdict": "ERROR",
            "confidence": 0.0,
            "reason": "GLM API call failed",
            "mismatch_type": "OTHER",
            "citations": [],
            "checks": [],
        }

    parsed = _extract_json_object(response)
    if not parsed:
        return {
            "verdict": "ERROR",
            "confidence": 0.0,
            "reason": f"Could not parse model JSON: {response[:180]}",
            "mismatch_type": "OTHER",
            "citations": [],
            "checks": [],
        }

    ok, normalized = _validate_result(parsed)
    if not ok:
        return {
            "verdict": "ERROR",
            "confidence": 0.0,
            "reason": f"Invalid verdict JSON: {json.dumps(parsed)[:180]}",
            "mismatch_type": "OTHER",
            "citations": [],
            "checks": [],
        }
    return normalized


def detect_mode(target_file: str, mode: str) -> str:
    if mode != "auto":
        return mode
    path = Path(target_file)
    if path.name.upper() == "DECISIONS.MD":
        return "decisions"
    ext = path.suffix.lower()
    if ext in (".xlsx", ".xls", ".csv"):
        return "dataset"
    if ext in (".docx", ".pdf"):
        return "contract"
    return "document"


def render_report(
    results: list[dict],
    target_file: str,
    source_dir: str,
    mode: str,
    model: str,
) -> str:
    total = len(results)
    counts: dict[str, int] = {}
    for r in results:
        verdict = r["result"]["verdict"]
        counts[verdict] = counts.get(verdict, 0) + 1

    def pct(n: int) -> str:
        return "0%" if total == 0 else f"{(n * 100) // total}%"

    lines = [
        "# Universal Fact-Check Audit",
        "",
        f"**Target:** `{Path(target_file).name}`",
        f"**Mode:** `{mode}`",
        f"**Sources:** `{source_dir}`",
        f"**Engine:** `{model}`",
        f"**Date:** {time.strftime('%Y-%m-%d %H:%M')}",
        "",
        "## Summary",
        "",
        "| Verdict | Count | % |",
        "|---------|-------|---|",
    ]
    for verdict in ["SUPPORTED", "PARTIAL", "NOT_FOUND", "CONTRADICTED", "INSUFFICIENT_CONTEXT", "ERROR"]:
        n = counts.get(verdict, 0)
        lines.append(f"| {verdict} | {n} | {pct(n)} |")
    lines.append(f"| **Total** | **{total}** | **100%** |")
    lines.append("")

    flagged = [r for r in results if r["result"]["verdict"] in {"CONTRADICTED", "NOT_FOUND", "PARTIAL", "INSUFFICIENT_CONTEXT", "ERROR"}]
    if flagged:
        lines.append("## Claims Needing Review")
        lines.append("")
        for item in flagged:
            claim: Claim = item["claim"]
            res = item["result"]
            lines.append(f"### {claim.claim_id} ({claim.locator}) - {res['verdict']}")
            lines.append(f"**Claim:** {claim.text}")
            lines.append(f"**Confidence:** {res.get('confidence', 0):.2f}")
            lines.append(f"**Mismatch Type:** {res.get('mismatch_type', 'OTHER')}")
            lines.append(f"**Reason:** {res.get('reason', '')}")
            if res.get("citations"):
                lines.append("**Citations:**")
                for c in res["citations"]:
                    lines.append(f"- `{c.get('source', '')}` @ `{c.get('locator', '')}`: {c.get('evidence', '')}")
            if res.get("checks"):
                lines.append(f"**Checks:** {', '.join(res['checks'])}")
            lines.append("")

    return "\n".join(lines)


def main() -> None:
    parser = argparse.ArgumentParser(description="Universal GLM-5 fact checker.")
    parser.add_argument("target", help="Claims source file (DECISIONS.md, document, contract, spreadsheet)")
    parser.add_argument("--sources", required=True, help="Directory containing source evidence files")
    parser.add_argument("--output", default=None, help="Output audit markdown path")
    parser.add_argument("--mode", choices=["auto", "decisions", "document", "contract", "dataset"], default="auto")
    parser.add_argument("--model", default="glm-5")
    parser.add_argument(
        "--parallel",
        type=int,
        default=DEFAULT_PARALLEL_WORKERS,
        help=f"Worker count (hard cap: {MAX_PARALLEL_WORKERS})",
    )
    parser.add_argument("--top-k", type=int, default=8, help="Top evidence chunks per claim")
    parser.add_argument("--max-claims", type=int, default=250, help="Max claims extracted from target")
    parser.add_argument("--max-source-files", type=int, default=300)
    parser.add_argument("--max-source-chunks", type=int, default=4000)
    parser.add_argument("--max-rows-per-table", type=int, default=800)
    parser.add_argument("--numeric-tolerance", type=float, default=0.0)
    parser.add_argument("--api-key-env", default="ZAI_API_KEY")
    parser.add_argument("--doppler-project", default=None)
    parser.add_argument("--doppler-config", default=None)
    args = parser.parse_args()

    target_path = Path(args.target)
    if not target_path.exists():
        print(f"ERROR: Target file not found: {args.target}", file=sys.stderr)
        sys.exit(1)
    if not Path(args.sources).exists():
        print(f"ERROR: Sources path not found: {args.sources}", file=sys.stderr)
        sys.exit(1)

    mode = detect_mode(args.target, args.mode)
    if args.output is None:
        args.output = str(Path(args.sources) / f"{target_path.stem}_UNIVERSAL_GLM_AUDIT.md")

    workers = max(1, min(args.parallel, MAX_PARALLEL_WORKERS))
    if args.parallel > MAX_PARALLEL_WORKERS:
        print(
            f"  Requested parallel={args.parallel} capped to {MAX_PARALLEL_WORKERS} to reduce GLM rate-limit failures"
        )
    api_key = get_api_key(
        env_var=args.api_key_env,
        doppler_project=args.doppler_project,
        doppler_config=args.doppler_config,
    )

    print("Universal Fact Checker")
    print(f"  Target: {args.target}")
    print(f"  Mode: {mode}")
    print(f"  Sources: {args.sources}")
    print(f"  Output: {args.output}")
    print(f"  Engine: {args.model}")
    print(f"  Parallel: {workers}")
    print("")

    print("Loading source evidence...")
    source_chunks = load_source_chunks(
        source_dir=args.sources,
        max_files=args.max_source_files,
        max_chunks=args.max_source_chunks,
        max_rows_per_table=args.max_rows_per_table,
    )
    print(f"  Loaded {len(source_chunks)} evidence chunks")
    if not source_chunks:
        print("ERROR: No readable source evidence found.", file=sys.stderr)
        sys.exit(1)

    print("\nExtracting claims...")
    claims = (
        extract_decision_claims(args.target, max_claims=args.max_claims)
        if mode == "decisions"
        else extract_generic_claims(args.target, mode=mode, max_claims=args.max_claims)
    )
    print(f"  Extracted {len(claims)} claims")
    if not claims:
        print("ERROR: No verifiable claims extracted from target.", file=sys.stderr)
        sys.exit(1)

    strict_mode = True

    results: list[dict] = [None] * len(claims)  # type: ignore[assignment]
    errors: list[str] = []

    if workers == 1:
        for idx, claim in enumerate(claims, start=1):
            print(f"  [{idx}/{len(claims)}] {claim.claim_id}: {claim.text[:80]}...", end="", flush=True)
            res = verify_claim(
                api_key=api_key,
                model=args.model,
                mode=mode,
                claim=claim,
                source_chunks=source_chunks,
                top_k=args.top_k,
                numeric_tolerance=args.numeric_tolerance,
            )
            print(f" -> {res['verdict']} ({res.get('confidence', 0):.2f})")
            results[idx - 1] = {"claim": claim, "result": res}
            if res["verdict"] == "ERROR":
                errors.append(f"{claim.claim_id}: {res.get('reason', '')[:120]}")
                if strict_mode:
                    break
            time.sleep(0.6)
    else:
        import threading

        lock = threading.Lock()
        completed = [0]
        abort_event = threading.Event()

        def _run_one(i: int) -> tuple[int, dict]:
            if abort_event.is_set():
                return i, {
                    "claim": claims[i],
                    "result": {
                        "verdict": "INSUFFICIENT_CONTEXT",
                        "confidence": 0.0,
                        "reason": "Skipped after prior error in strict mode.",
                        "mismatch_type": "OTHER",
                        "citations": [],
                        "checks": [],
                    },
                }
            claim = claims[i]
            time.sleep(i * 0.12)
            res = verify_claim(
                api_key=api_key,
                model=args.model,
                mode=mode,
                claim=claim,
                source_chunks=source_chunks,
                top_k=args.top_k,
                numeric_tolerance=args.numeric_tolerance,
            )
            with lock:
                completed[0] += 1
                n = completed[0]
            print(f"  [{n}/{len(claims)}] {claim.claim_id}: {claim.text[:55]}... -> {res['verdict']} ({res.get('confidence', 0):.2f})")
            if strict_mode and res["verdict"] == "ERROR":
                abort_event.set()
            return i, {"claim": claim, "result": res}

        with ThreadPoolExecutor(max_workers=workers) as exe:
            futures = {exe.submit(_run_one, i): i for i in range(len(claims))}
            for fut in as_completed(futures):
                i, item = fut.result()
                results[i] = item
                if item["result"]["verdict"] == "ERROR":
                    errors.append(f"{item['claim'].claim_id}: {item['result'].get('reason', '')[:120]}")

    if errors and strict_mode:
        print("\nHARD STOP: GLM-5 error(s) detected.", file=sys.stderr)
        for e in errors[:20]:
            print(f"  - {e}", file=sys.stderr)
        print("Run failed due to strict no-fallback policy.", file=sys.stderr)
        sys.exit(10)

    rendered = render_report(
        results=[r for r in results if r is not None],
        target_file=args.target,
        source_dir=args.sources,
        mode=mode,
        model=args.model,
    )
    Path(args.output).parent.mkdir(parents=True, exist_ok=True)
    Path(args.output).write_text(rendered, encoding="utf-8")
    print(f"\nReport written to: {args.output}")

    verdicts = [r["result"]["verdict"] for r in results if r is not None]
    contradicted = verdicts.count("CONTRADICTED")
    needs_review = sum(v in {"PARTIAL", "NOT_FOUND", "INSUFFICIENT_CONTEXT"} for v in verdicts)
    print(
        f"Final: {verdicts.count('SUPPORTED')} SUPPORTED, {verdicts.count('PARTIAL')} PARTIAL, "
        f"{verdicts.count('NOT_FOUND')} NOT_FOUND, {contradicted} CONTRADICTED, "
        f"{verdicts.count('INSUFFICIENT_CONTEXT')} INSUFFICIENT_CONTEXT"
    )

    if contradicted > 0:
        sys.exit(1)
    if needs_review > 0:
        sys.exit(2)
    sys.exit(0)


if __name__ == "__main__":
    main()
