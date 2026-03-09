"""
BEI Weekly Meta Ads Report — Reusable DOCX Template (v2 — CEO-Grade)
Brand: Bebang Halo-Halo | Palette: Jelly Green / Yema Gold / Ube Purple
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

PALETTE = {
    'primary': '04400A', 'secondary': 'C8900A', 'accent': '801297',
    'light': 'E6ECE7', 'light_gold': 'F8F0D9', 'light_purple': 'F2E5F5',
    'text': '1A1A1A', 'text_light': '666666', 'bg_alt': 'E6ECE7',
    'bg_page': 'F9F5EB', 'white': 'FFFFFF', 'border': 'D4D0C8',
    'green_mid': '2D7A35', 'error': 'CC0000', 'green_good': 'D9F2D9',
    'red_bad': 'F9D6D6', 'yellow_warn': 'FFF3CD',
}
_DARK_BG = {PALETTE['primary'], PALETTE['accent'], PALETTE['green_mid']}

# -- Helpers ----------------------------------------------------------------

def hex_to_rgb(h): return RGBColor(int(h[:2], 16), int(h[2:4], 16), int(h[4:6], 16))

def _run(p, text, size=10, bold=False, color=PALETTE['text'], italic=False):
    r = p.add_run(text)
    r.font.size, r.font.name, r.bold, r.italic = Pt(size), 'Calibri', bold, italic
    r.font.color.rgb = hex_to_rgb(color)
    return r

def set_cell_bg(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex); shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        el = OxmlElement(f'w:{edge}')
        for k, d in [('w:val', 'single'), ('w:sz', '4'), ('w:color', PALETTE['border'])]:
            el.set(qn(k), val.get(k.split(':')[1], d))
        el.set(qn('w:space'), '0')
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_styled_para(doc, text, size=10, bold=False, color=None,
                    alignment=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    _run(p, text, size, bold, color or PALETTE['text'])
    if alignment: p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def add_section_heading(doc, text, level=1):
    if level == 1:
        p = add_styled_para(doc, text, 16, True, PALETTE['primary'], space_before=12, space_after=6)
        pBdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
        for a, v in [('w:val','single'),('w:sz','12'),('w:color',PALETTE['primary']),('w:space','4')]:
            bottom.set(qn(a), v)
        pBdr.append(bottom); p._p.get_or_add_pPr().append(pBdr)
        return p
    if level == 2:
        return add_styled_para(doc, text, 13, True, PALETTE['secondary'], space_before=10, space_after=4)
    return add_styled_para(doc, text, 11, True, PALETTE['primary'], space_before=6, space_after=3)

def create_branded_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; set_cell_bg(c, PALETTE['primary'])
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(c.paragraphs[0], h, 9, True, PALETTE['white'])
    for ri, rd in enumerate(rows):
        bg = PALETTE['bg_alt'] if ri % 2 == 0 else PALETTE['white']
        for ci, v in enumerate(rd):
            c = table.rows[ri+1].cells[ci]; set_cell_bg(c, bg)
            _run(c.paragraphs[0], str(v), 9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for r in table.rows: r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

def format_php(amount):
    if amount is None or amount == 0: return '\u2014'
    return f'(PHP {abs(amount):,.2f})' if amount < 0 else f'PHP {amount:,.2f}'

def _callout(doc, title, body, border_color=None):
    bc = border_color or PALETTE['secondary']
    bg = PALETTE['light_gold'] if bc == PALETTE['secondary'] else PALETTE['light']
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]; set_cell_bg(cell, bg)
    for edge in ('top','bottom','start','end'):
        set_cell_border(cell, **{edge: {'val':'single','sz':'12','color':bc}})
    _run(cell.paragraphs[0], title, 10, True, bc)
    _run(cell.add_paragraph(), body, 9, color=PALETTE['text'])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def _kpi_bar(doc, kpis):
    tbl = doc.add_table(rows=2, cols=len(kpis)); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value, color) in enumerate(kpis):
        vc = tbl.rows[0].cells[i]; set_cell_bg(vc, color)
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = PALETTE['white'] if color in _DARK_BG else PALETTE['text']
        _run(vc.paragraphs[0], str(value), 22, True, tc)
        lc = tbl.rows[1].cells[i]; set_cell_bg(lc, PALETTE['white'])
        lc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(lc.paragraphs[0], label, 8, True, PALETTE['text_light'])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def _delta_arrow(pct):
    if pct is None or pct == 0: return '\u2192 0%'
    arrow = '\u2191' if pct > 0 else '\u2193'
    return f'{arrow} {abs(pct):.1f}%'

def _roas_bg(roas):
    if roas is None: return PALETTE['white']
    if roas >= 3.0: return PALETTE['green_good']
    if roas < 1.0: return PALETTE['red_bad']
    return PALETTE['yellow_warn']

def _cpa_bg(cpa, idx):
    if cpa is not None and cpa < 150: return 'D9F2D9'
    if cpa is not None and cpa > 200: return 'F9D6D6'
    return PALETTE['light_gold'] if cpa is not None else (PALETTE['bg_alt'] if idx%2==0 else PALETTE['white'])

# -- Main generator ---------------------------------------------------------

def generate_weekly_report(data: dict, output_path: str) -> str:
    """Generate branded BEI weekly Meta Ads report DOCX (v2 — CEO-grade). Returns output_path."""
    doc = Document()
    for s in doc.sections:
        s.page_width, s.page_height = Cm(21), Cm(29.7)
        s.left_margin, s.right_margin = Cm(3.18), Cm(2.54)
        s.top_margin = s.bottom_margin = Cm(1.91)
    ns = doc.styles['Normal']; ns.font.name = 'Calibri'
    ns.font.size = Pt(10); ns.font.color.rgb = hex_to_rgb(PALETTE['text'])
    we = data.get('week_ending', datetime.now().strftime('%Y-%m-%d'))

    # -- 1. Cover Page --
    for _ in range(4): doc.add_paragraph().paragraph_format.space_after = Pt(0)
    cv = doc.add_table(rows=1, cols=1); cv.alignment = WD_TABLE_ALIGNMENT.CENTER
    cc = cv.rows[0].cells[0]; set_cell_bg(cc, PALETTE['primary'])
    p = cc.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = _run(p, 'BEBANG HALO-HALO', 14, color=PALETTE['white']); r.font.letter_spacing = Pt(3)
    p2 = cc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2, 'Weekly Meta Ads Intelligence', 28, True, PALETTE['white'])
    p3 = cc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(20)
    _run(p3, f'Week ending {we}', 14, color=PALETTE['light'])
    doc.add_page_break()

    # -- 2. Executive Summary (Decision-First) --
    add_section_heading(doc, '1. Executive Summary')
    _callout(doc, 'CEO BRIEFING', data.get('ai_analysis', ''), PALETTE['secondary'])

    # Revenue-first KPI bar
    roas = data.get('portfolio_roas', 0)
    roas_str = f'{roas:.1f}x' if roas else '\u2014'
    _kpi_bar(doc, [
        ('Revenue', format_php(data.get('total_revenue')), PALETTE['primary']),
        ('ROAS', roas_str, PALETTE['green_mid']),
        ('Spend', format_php(data.get('total_spend')), PALETTE['secondary']),
        ('Purchases', str(data.get('total_purchases', 0)), PALETTE['accent']),
        ('Avg CPA', format_php(data.get('avg_cpa')), PALETTE['green_mid']),
    ])

    # WoW deltas
    wow = data.get('wow_deltas', {})
    if wow:
        add_section_heading(doc, 'Week-over-Week Changes', level=3)
        delta_items = [
            ('Revenue', wow.get('revenue_pct')),
            ('ROAS', wow.get('roas_pct')),
            ('Spend', wow.get('spend_pct')),
            ('Purchases', wow.get('purchases_pct')),
            ('CPA', wow.get('cpa_pct')),
        ]
        tbl = doc.add_table(rows=1, cols=len(delta_items))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (label, pct) in enumerate(delta_items):
            c = tbl.rows[0].cells[i]
            if pct is not None and pct != 0:
                bg = PALETTE['green_good'] if (pct > 0 and label != 'CPA') or (pct < 0 and label == 'CPA') else PALETTE['red_bad']
            else:
                bg = PALETTE['white']
            set_cell_bg(c, bg)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(c.paragraphs[0], _delta_arrow(pct), 11, True)
            p_label = c.add_paragraph(); p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p_label, label, 8, True, PALETTE['text_light'])
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # -- 3. Funnel Allocation --
    funnel = data.get('funnel_allocation', {})
    if funnel and any(funnel.get(k) for k in ('tof_pct', 'mof_pct', 'bof_pct')):
        add_section_heading(doc, '2. Funnel Allocation')
        funnel_rows = [
            ['TOF (Awareness)', f"{funnel.get('tof_pct', 0):.0f}%", '20-30%',
             f"{funnel.get('tof_roas', 0):.1f}x"],
            ['MOF (Consideration)', f"{funnel.get('mof_pct', 0):.0f}%", '30-40%',
             f"{funnel.get('mof_roas', 0):.1f}x"],
            ['BOF (Conversion)', f"{funnel.get('bof_pct', 0):.0f}%", '30-40%',
             f"{funnel.get('bof_roas', 0):.1f}x"],
        ]
        t = create_branded_table(doc, ['Stage', 'Actual %', 'Ideal %', 'ROAS'], funnel_rows,
                                 col_widths=[2.0, 1.0, 1.0, 1.0])
        # Color-code ROAS cells
        for ri in range(3):
            roas_val = [funnel.get('tof_roas',0), funnel.get('mof_roas',0), funnel.get('bof_roas',0)][ri]
            set_cell_bg(t.rows[ri+1].cells[3], _roas_bg(roas_val))
    else:
        add_section_heading(doc, '2. Funnel Allocation')
        add_styled_para(doc, 'Funnel data not available.', color=PALETTE['text_light'])

    # -- 4. Campaign Performance (with Revenue & ROAS) --
    add_section_heading(doc, '3. Campaign Performance')
    campaigns = data.get('campaigns', [])
    if campaigns:
        hdrs = ['Campaign', 'Funnel', 'Spend', 'Revenue', 'ROAS', 'Purchases', 'CPA', 'CTR', 'Freq']
        tbl = doc.add_table(rows=1+len(campaigns), cols=len(hdrs))
        tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(hdrs):
            c = tbl.rows[0].cells[i]; set_cell_bg(c, PALETTE['primary'])
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(c.paragraphs[0], h, 8, True, PALETTE['white'])
        for idx, cp in enumerate(campaigns):
            cpa = cp.get('cpa_7d'); bg = _cpa_bg(cpa, idx)
            camp_roas = cp.get('roas_7d', 0)
            vals = [
                cp.get('campaign_name', '')[:30],
                cp.get('funnel_stage', ''),
                format_php(cp.get('spend_7d')),
                format_php(cp.get('revenue_7d')),
                f"{camp_roas:.1f}x" if camp_roas else '\u2014',
                str(cp.get('purchases_7d', 0)),
                format_php(cpa),
                f"{cp.get('ctr_7d', 0):.2f}%",
                f"{cp.get('frequency_7d', 0):.1f}",
            ]
            for ci, v in enumerate(vals):
                c = tbl.rows[idx+1].cells[ci]; set_cell_bg(c, bg)
                _run(c.paragraphs[0], v, 8)
            # Color ROAS cell
            set_cell_bg(tbl.rows[idx+1].cells[4], _roas_bg(camp_roas))
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
    else:
        add_styled_para(doc, 'No campaign data for this period.', color=PALETTE['text_light'])

    # -- 5. Creative Performance --
    add_section_heading(doc, '4. Creative Performance by Type')
    creatives = data.get('creative_performance', [])
    if creatives:
        rows = []
        for cr in creatives:
            cr_roas = cr.get('roas', 0)
            rows.append([
                cr.get('creative_type', 'Unknown'),
                str(cr.get('ad_count', 0)),
                format_php(cr.get('spend')),
                format_php(cr.get('revenue')),
                f"{cr_roas:.1f}x" if cr_roas else '\u2014',
                f"{cr.get('avg_ctr', 0):.2f}%",
            ])
        t = create_branded_table(doc, ['Type', 'Ads', 'Spend', 'Revenue', 'ROAS', 'CTR'], rows,
                                 col_widths=[1.5, 0.5, 1.2, 1.2, 0.8, 0.8])
        # Color ROAS cells
        for ri, cr in enumerate(creatives):
            set_cell_bg(t.rows[ri+1].cells[4], _roas_bg(cr.get('roas')))
    else:
        add_styled_para(doc, 'No creative performance data.', color=PALETTE['text_light'])

    # -- 6. Flagged Ads --
    add_section_heading(doc, '5. Flagged Ads')
    flagged = data.get('flagged_ads', [])
    if flagged:
        rows = [[f.get('ad_name','')[:30], f.get('campaign_name','')[:20], f.get('flag_reason',''),
                 format_php(f.get('spend_7d')), format_php(f.get('cpa_7d')),
                 f"{f.get('ctr_7d',0):.2f}%", f"{f.get('frequency',0):.1f}"] for f in flagged]
        t = create_branded_table(doc, ['Ad Name','Campaign','Flag','Spend','CPA','CTR','Freq'], rows)
        for ri, f in enumerate(flagged):
            if any(k in (f.get('flag_reason') or '').lower() for k in ('critical','high_cpa','zero_purchases')):
                for ci in range(7): set_cell_bg(t.rows[ri+1].cells[ci], PALETTE['red_bad'])
    else:
        add_styled_para(doc, 'No flagged ads this week.', color=PALETTE['text_light'])
    doc.add_page_break()

    # -- 7. Organic Post Winners --
    add_section_heading(doc, '6. Boost Candidates')
    boosts = data.get('boost_candidates', [])[:5]
    if boosts:
        rows = [[(b.get('post_text') or '')[:60],
                 str(b.get('engagement_score',0)),
                 str(b.get('days_old', '')),
                 b.get('suggested_boost_budget', ''),
                 str(b.get('likes',0)), str(b.get('comments',0))]
                for b in boosts]
        t = create_branded_table(doc, ['Post', 'Score', 'Age', 'Suggested Budget', 'Likes', 'Comments'],
                                 rows, col_widths=[2.0, 0.6, 0.5, 1.3, 0.5, 0.6])
        set_cell_bg(t.rows[1].cells[0], PALETTE['light_gold'])
    else:
        add_styled_para(doc, 'No boost candidates identified.', color=PALETTE['text_light'])

    # -- 8. Weekly Trend (4-week, with Revenue & ROAS) --
    add_section_heading(doc, '7. 4-Week Trend')
    trend = data.get('weekly_trend', [])
    if trend:
        rows = []
        for t_row in trend:
            t_roas = t_row.get('roas', 0)
            rows.append([
                t_row.get('week', ''),
                t_row.get('campaign_name', '')[:25],
                t_row.get('funnel_stage', ''),
                format_php(t_row.get('spend')),
                str(t_row.get('purchases', 0)),
                format_php(t_row.get('revenue')),
                f"{t_roas:.1f}x" if t_roas else '\u2014',
                format_php(t_row.get('cpa')),
            ])
        create_branded_table(doc, ['Week', 'Campaign', 'Funnel', 'Spend', 'Purch', 'Revenue', 'ROAS', 'CPA'],
                             rows, col_widths=[0.7, 1.3, 0.5, 1.0, 0.5, 1.0, 0.6, 0.9])
    else:
        add_styled_para(doc, 'No trend data available.', color=PALETTE['text_light'])
    doc.add_page_break()

    # -- 9. Recommendations (with PHP amounts & expected impact) --
    add_section_heading(doc, '8. Recommendations')
    recs = data.get('recommendations', [])
    for i, rec in enumerate(recs, 1):
        pri = (rec.get('priority') or 'MEDIUM').upper()
        pc = PALETTE['accent'] if pri == 'HIGH' else (PALETTE['secondary'] if pri == 'MEDIUM' else PALETTE['primary'])
        p = doc.add_paragraph()
        _run(p, f'{i}. [{pri}] ', 10, True, pc)
        _run(p, rec.get('action', ''), 10, True)
        # PHP amount callout
        if rec.get('php_amount'):
            _run(p, f'  ({rec["php_amount"]})', 10, True, PALETTE['secondary'])
        if rec.get('reasoning'):
            p2 = doc.add_paragraph(); p2.paragraph_format.left_indent = Pt(18)
            _run(p2, rec['reasoning'], 9, color=PALETTE['text_light'])
        if rec.get('expected_impact'):
            p3 = doc.add_paragraph(); p3.paragraph_format.left_indent = Pt(18)
            _run(p3, f'Expected: {rec["expected_impact"]}', 9, True, PALETTE['green_mid'])
    if not recs:
        add_styled_para(doc, 'No recommendations generated.', color=PALETTE['text_light'])

    # -- Footer --
    doc.add_paragraph()
    fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(fp, '-- End of Report --', 9, italic=True, color=PALETTE['text_light'])
    fp2 = doc.add_paragraph(); fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(fp2, f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")} PHT | AI-Powered by Claude', 8, color=PALETTE['text_light'])

    doc.save(output_path)
    return output_path


# -- Sample data & test -----------------------------------------------------

def _sample_data():
    C = lambda n,o,fs,s,p,r,ro,cpa,ctr,f: dict(campaign_name=n, objective=o, funnel_stage=fs,
        spend_7d=s, purchases_7d=p, revenue_7d=r, roas_7d=ro, cpa_7d=cpa, ctr_7d=ctr, frequency_7d=f, status='ACTIVE')
    return {
        'week_ending': '2026-03-09',
        'total_spend': 208683.91, 'total_revenue': 502841.50, 'total_purchases': 1772, 'avg_cpa': 117.77,
        'portfolio_roas': 2.41,
        'ai_analysis': ('Revenue hit PHP 502K this week with a portfolio ROAS of 2.41x — strong but down '
            'from 2.6x last week due to rising TOF spend. BOF campaigns deliver 3.8x ROAS and should '
            'receive PHP 30K more budget shifted from underperforming TOF awareness campaigns. '
            'Retargeting frequency at 3.2 needs immediate creative refresh.'),
        'wow_deltas': {'spend_pct': 5.2, 'purchases_pct': 3.1, 'cpa_pct': 2.0, 'revenue_pct': -1.8, 'roas_pct': -7.3},
        'funnel_allocation': {'tof_pct': 35, 'mof_pct': 30, 'bof_pct': 35,
                              'tof_roas': 0.8, 'mof_roas': 1.9, 'bof_roas': 3.8},
        'campaigns': [
            C('Lookalike 1%','CONVERSIONS','BOF',82500,698,313500,3.8,118.19,2.34,1.8),
            C('Retargeting 7d','CONVERSIONS','BOF',48200,472,120500,2.5,102.12,3.12,3.2),
            C('Broad Interest','OUTCOME_TRAFFIC','MOF',42085,380,45200,1.07,110.75,0.89,2.1),
            C('Brand Awareness','OUTCOME_AWARENESS','TOF',35898,222,23641,0.66,161.70,1.45,1.5),
        ],
        'creative_performance': [
            {'creative_type': 'VIDEO', 'ad_count': 12, 'spend': 95000, 'revenue': 280000, 'roas': 2.95, 'avg_ctr': 2.1},
            {'creative_type': 'IMAGE', 'ad_count': 18, 'spend': 78000, 'revenue': 165000, 'roas': 2.12, 'avg_ctr': 1.8},
            {'creative_type': 'CAROUSEL', 'ad_count': 5, 'spend': 35683, 'revenue': 57841, 'roas': 1.62, 'avg_ctr': 1.2},
        ],
        'flagged_ads': [
            {'ad_name':'Brand_Video_03','campaign_name':'Brand Awareness',
             'flag_reason':'ZERO_PURCHASES','spend_7d':8200,'cpa_7d':0,'ctr_7d':0.45,'frequency':2.8},
            {'ad_name':'Broad_Carousel_01','campaign_name':'Broad Interest',
             'flag_reason':'HIGH_CPA','spend_7d':5200,'cpa_7d':260,'ctr_7d':0.62,'frequency':2.3},
        ],
        'boost_candidates': [
            {'post_text':'FurMom/FurDads flex nyo mga bebe nyo!','engagement_score':5910,
             'likes':237,'comments':1876,'shares':45,'days_old':3,'suggested_boost_budget':'PHP 5,000-15,000 (3-5 days)'},
            {'post_text':'TAG MO NA BEB!','engagement_score':2982,
             'likes':106,'comments':852,'shares':32,'days_old':5,'suggested_boost_budget':'PHP 2,000-5,000 (3 days)'},
        ],
        'weekly_trend': [
            {'week':'Feb 17','campaign_name':'Lookalike 1%','funnel_stage':'BOF','spend':75000,'purchases':650,'revenue':290000,'cpa':115.38,'roas':3.87},
            {'week':'Feb 24','campaign_name':'Lookalike 1%','funnel_stage':'BOF','spend':79000,'purchases':670,'revenue':305000,'cpa':117.91,'roas':3.86},
            {'week':'Mar 03','campaign_name':'Lookalike 1%','funnel_stage':'BOF','spend':82500,'purchases':698,'revenue':313500,'cpa':118.19,'roas':3.80},
        ],
        'recommendations': [
            {'priority':'HIGH','action':'Shift budget from TOF to BOF','php_amount':'PHP 30,000',
             'reasoning':'TOF delivering 0.8x ROAS vs BOF 3.8x. Currently 35% TOF (ideal 20-25%).','expected_impact':'Portfolio ROAS from 2.4x to 2.8x'},
            {'priority':'HIGH','action':'Refresh retargeting creatives','php_amount':'PHP 5,000 creative cost',
             'reasoning':'Frequency at 3.2 and rising. Audience seeing same ads too often.','expected_impact':'Reduce frequency to <2.5, improve CPA by 10-15%'},
            {'priority':'MEDIUM','action':'Boost FurMom/FurDads post','php_amount':'PHP 10,000 (5 days)',
             'reasoning':'1,876 organic comments — viral potential. Convert to paid reach.','expected_impact':'Est. 50-80 additional purchases at <PHP 125 CPA'},
            {'priority':'LOW','action':'Pause Brand_Video_03','php_amount':'PHP 8,200/week saved',
             'reasoning':'Zero purchases on PHP 8.2K spend. Creative fatigue.','expected_impact':'Reallocate to Lookalike which delivers 3.8x ROAS'},
        ],
    }

if __name__ == '__main__':
    import os
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'sample_weekly_report.docx')
    generate_weekly_report(_sample_data(), out)
    print(f'Sample report saved to: {out}')
