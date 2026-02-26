#!/usr/bin/env python3
"""
Create a practical, beginner-friendly 90-day EA training manual for Kath.
"""

from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PALETTE: Dict[str, str] = {
    "primary": "04400A",
    "secondary": "C8900A",
    "accent": "801297",
    "light": "E6ECE7",
    "light_gold": "F8F0D9",
    "light_purple": "F2E5F5",
    "text": "1A1A1A",
    "text_light": "666666",
    "bg_alt": "E6ECE7",
    "white": "FFFFFF",
    "border": "D4D0C8",
    "success_bg": "D1FAE5",
    "warning_bg": "FEF3C7",
    "risk_bg": "FEE2E2",
}

FLOW_STEP_COLORS: Dict[str, Dict[str, str]] = {
    "start": {"bg": "D1FAE5", "border": "2D7A35"},
    "process": {"bg": "DBEAFE", "border": "2563EB"},
    "decision": {"bg": "FEF3C7", "border": "C8900A"},
    "support": {"bg": "F2E5F5", "border": "801297"},
    "end": {"bg": "E5E7EB", "border": "6B7280"},
}


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def ensure_style(doc: Document, name: str, style_type: WD_STYLE_TYPE):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def setup_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.15)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(PALETTE["text"])

    h1 = ensure_style(doc, "DesignH1", WD_STYLE_TYPE.PARAGRAPH)
    h1.font.name = "Calibri Light"
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = rgb(PALETTE["primary"])
    h1.paragraph_format.space_before = Pt(22)
    h1.paragraph_format.space_after = Pt(8)

    h2 = ensure_style(doc, "DesignH2", WD_STYLE_TYPE.PARAGRAPH)
    h2.font.name = "Calibri"
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = rgb(PALETTE["secondary"])
    h2.paragraph_format.space_before = Pt(17)
    h2.paragraph_format.space_after = Pt(6)

    h3 = ensure_style(doc, "DesignH3", WD_STYLE_TYPE.PARAGRAPH)
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = rgb(PALETTE["primary"])
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)

    body = ensure_style(doc, "DesignBody", WD_STYLE_TYPE.PARAGRAPH)
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    body.font.color.rgb = rgb(PALETTE["text"])
    body.paragraph_format.line_spacing = 1.15
    body.paragraph_format.space_after = Pt(8)

    caption = ensure_style(doc, "DesignCaption", WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(PALETTE["text_light"])


def set_cell_bg(cell, color_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            edge_data = kwargs[edge]
            edge_tag = OxmlElement(f"w:{edge}")
            for key, value in edge_data.items():
                edge_tag.set(qn(f"w:{key}"), str(value))
            tc_borders.append(edge_tag)
    tc_pr.append(tc_borders)


def set_cell_margins(cell, top: float = 0.05, bottom: float = 0.05, left: float = 0.1, right: float = 0.1) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        m = OxmlElement(f"w:{edge}")
        m.set(qn("w:w"), str(int(val * 1440)))
        m.set(qn("w:type"), "dxa")
        tc_mar.append(m)
    tc_pr.append(tc_mar)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    color_hex: str = "1A1A1A",
    size: int = 10,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = rgb(color_hex)


def remove_table_borders(table) -> None:
    tbl = table._element
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        borders.append(border)
    tbl_pr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)


def add_page_number_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p.add_run("Page ")
    run_page = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_page._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run_page._r.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_page._r.append(fld_end)

    p.add_run(" of ")
    run_total = p.add_run()
    fld2_begin = OxmlElement("w:fldChar")
    fld2_begin.set(qn("w:fldCharType"), "begin")
    run_total._r.append(fld2_begin)
    instr2 = OxmlElement("w:instrText")
    instr2.set(qn("xml:space"), "preserve")
    instr2.text = "NUMPAGES"
    run_total._r.append(instr2)
    fld2_end = OxmlElement("w:fldChar")
    fld2_end.set(qn("w:fldCharType"), "end")
    run_total._r.append(fld2_end)

    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = rgb(PALETTE["text_light"])


def add_cover_page(doc: Document, title: str, subtitle: str, date_text: str) -> None:
    band = doc.add_table(rows=2, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(band)

    top_cell = band.rows[0].cells[0]
    set_cell_bg(top_cell, PALETTE["primary"])
    top_cell.paragraphs[0].paragraph_format.space_before = Pt(14)
    top_cell.paragraphs[0].paragraph_format.space_after = Pt(14)

    bottom_cell = band.rows[1].cells[0]
    set_cell_bg(bottom_cell, PALETTE["secondary"])
    bottom_cell.paragraphs[0].paragraph_format.space_before = Pt(2)
    bottom_cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    spacer = doc.add_paragraph("", style="DesignBody")
    spacer.paragraph_format.space_after = Pt(28)

    p1 = doc.add_paragraph(style="DesignH1")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(title)
    r1.font.size = Pt(30)
    r1.font.color.rgb = rgb(PALETTE["primary"])

    p2 = doc.add_paragraph(style="DesignH2")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(subtitle)

    p3 = doc.add_paragraph(style="DesignBody")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(date_text)

    p4 = doc.add_paragraph(style="DesignCaption")
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run("Practical training manual for a new executive assistant.")

    doc.add_page_break()


def add_section_divider(doc: Document) -> None:
    line = doc.add_table(rows=1, cols=1)
    remove_table_borders(line)
    set_cell_bg(line.rows[0].cells[0], PALETTE["light"])
    line.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(2)
    line.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Pt(2)
    doc.add_paragraph("", style="DesignBody")


def add_callout(doc: Document, title: str, content: str, *, kind: str = "info") -> None:
    kind_map = {
        "info": {"bg": PALETTE["light"], "border": PALETTE["primary"], "title": PALETTE["primary"]},
        "warning": {"bg": PALETTE["light_gold"], "border": PALETTE["secondary"], "title": PALETTE["secondary"]},
        "support": {"bg": PALETTE["light_purple"], "border": PALETTE["accent"], "title": PALETTE["accent"]},
        "risk": {"bg": PALETTE["risk_bg"], "border": "CC0000", "title": "CC0000"},
    }
    style = kind_map[kind]

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, style["bg"])
    set_cell_border(
        cell,
        left={"sz": 24, "val": "single", "color": style["border"]},
        top={"sz": 4, "val": "single", "color": style["border"]},
        right={"sz": 4, "val": "single", "color": style["border"]},
        bottom={"sz": 4, "val": "single", "color": style["border"]},
    )
    set_cell_margins(cell, top=0.08, bottom=0.08, left=0.14, right=0.14)
    set_cell_text(cell, title, bold=True, color_hex=style["title"], size=11)
    cell.add_paragraph(content, style="DesignBody")
    doc.add_paragraph("", style="DesignBody")


def add_bullets(doc: Document, lines: List[str]) -> None:
    for line in lines:
        doc.add_paragraph(f"- {line}", style="DesignBody")


def add_numbered_steps(doc: Document, lines: List[str]) -> None:
    for i, line in enumerate(lines, start=1):
        p = doc.add_paragraph(style="DesignBody")
        n = p.add_run(f"{i}. ")
        n.bold = True
        n.font.color.rgb = rgb(PALETTE["accent"])
        p.add_run(line)


def add_table(
    doc: Document,
    headers: List[str],
    rows: List[List[str]],
    *,
    col_widths: List[float] | None = None,
    source: str | None = None,
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    remove_table_borders(table)

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_bg(cell, PALETTE["primary"])
        set_cell_margins(cell, top=0.07, bottom=0.07, left=0.1, right=0.1)
        set_cell_text(
            cell,
            header,
            bold=True,
            color_hex=PALETTE["white"],
            size=10,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    for ridx, row_data in enumerate(rows, start=1):
        for cidx, value in enumerate(row_data):
            cell = table.rows[ridx].cells[cidx]
            if ridx % 2 == 0:
                set_cell_bg(cell, PALETTE["bg_alt"])
            set_cell_margins(cell, top=0.05, bottom=0.05, left=0.1, right=0.1)
            set_cell_text(cell, value, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)

    if source:
        doc.add_paragraph(f"Source: {source}", style="DesignCaption")
    doc.add_paragraph("", style="DesignBody")


def add_flowchart(doc: Document, title: str, steps: List[Dict[str, str]], source: str) -> None:
    doc.add_paragraph(title, style="DesignH2")

    total_rows = len(steps) * 2 - 1
    table = doc.add_table(rows=total_rows, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)

    for row in table.rows:
        row.cells[0].width = Inches(0.75)
        row.cells[1].width = Inches(5.4)
        row.cells[2].width = Inches(0.75)

    idx = 0
    for ridx in range(total_rows):
        if ridx % 2 == 0:
            step = steps[idx]
            style = FLOW_STEP_COLORS.get(step["type"], FLOW_STEP_COLORS["process"])
            cell = table.rows[ridx].cells[1]
            set_cell_bg(cell, style["bg"])
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": style["border"]},
                left={"sz": 12, "val": "single", "color": style["border"]},
                right={"sz": 12, "val": "single", "color": style["border"]},
                bottom={"sz": 12, "val": "single", "color": style["border"]},
            )
            set_cell_margins(cell, top=0.08, bottom=0.08, left=0.12, right=0.12)
            set_cell_text(
                cell,
                f"{step['title']}\n{step['desc']}",
                bold=True,
                color_hex=PALETTE["text"],
                size=10,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            idx += 1
        else:
            arrow = table.rows[ridx].cells[1]
            set_cell_text(
                arrow,
                "v",
                bold=True,
                color_hex=PALETTE["secondary"],
                size=14,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )

    doc.add_paragraph(f"Source: {source}", style="DesignCaption")
    doc.add_paragraph("", style="DesignBody")


def main() -> int:
    parser = argparse.ArgumentParser(description="Create Kath EA training DOCX.")
    parser.add_argument(
        "--output",
        default="docs/reports/KATH_WORLD_CLASS_EA_TRAINING_MANUAL_2026-02-26.docx",
        help="Output DOCX path",
    )
    args = parser.parse_args()

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    setup_page(doc)
    setup_styles(doc)
    add_page_number_footer(doc)

    add_cover_page(
        doc,
        "KATH EA TRAINING MANUAL",
        "90-Day Executive Assistant Training Program",
        datetime.now().strftime("%B %d, %Y"),
    )

    doc.add_paragraph("How to Use This Manual", style="DesignH1")
    doc.add_paragraph(
        "This guide is written for a new EA. It uses plain language, step-by-step actions, and real message examples. "
        "You do not need to be perfect on day one. The goal is steady progress every week.",
        style="DesignBody",
    )
    add_bullets(
        doc,
        [
            "Follow the first 10-day plan exactly. It builds your base.",
            "Use the weekly plan to build speed and confidence.",
            "Use panic rescue steps any time you feel overloaded.",
            "Use the sample message templates until your own style becomes natural.",
        ],
    )
    add_callout(
        doc,
        "Training mindset",
        "If you are unsure, ask early. Asking early prevents small issues from becoming big issues.",
        kind="support",
    )
    add_section_divider(doc)

    doc.add_paragraph("What Success Looks Like After 90 Days", style="DesignH2")
    add_bullets(
        doc,
        [
            "Important emails are surfaced quickly and clearly.",
            "Important chat messages are not missed.",
            "Calendar conflicts are caught early and fixed fast.",
            "Meetings are prepared with clear context.",
            "Daily updates are clear, short, and useful.",
        ],
    )

    doc.add_paragraph("90-Day Roadmap", style="DesignH1")
    roadmap_rows = [
        [
            "Month 1",
            "Build core habits",
            "Email triage, chat monitoring, and calendar basics",
            "Can run daily routine with guidance",
        ],
        [
            "Month 2",
            "Build speed",
            "Faster triage, stronger follow-up, better meeting prep",
            "Can run most days with light support",
        ],
        [
            "Month 3",
            "Build ownership",
            "Proactive reminders, better decision support, cleaner reporting",
            "Can run the role independently",
        ],
    ]
    add_table(
        doc,
        ["Phase", "Focus", "What You Practice", "End-of-Phase Proof"],
        roadmap_rows,
        col_widths=[1.0, 1.3, 2.5, 1.8],
        source="EA training design for Sam and Kath workflow (internal)",
    )

    doc.add_paragraph("First 10 Working Days (Exact Plan)", style="DesignH1")
    first_10_rows = [
        ["Day 1", "Set up inbox labels, chat priorities, and calendar view", "Tools are organized and ready"],
        ["Day 2", "Shadow current flow and capture questions", "List of unclear tasks is complete"],
        ["Day 3", "Do guided email triage in 3 time blocks", "Priority tagging is mostly correct"],
        ["Day 4", "Do guided chat monitoring and alerts", "Urgent chat alerts are on time"],
        ["Day 5", "Manage calendar updates and reminders", "No missed reminders today"],
        ["Day 6", "Run morning briefing draft", "Morning brief is clear and useful"],
        ["Day 7", "Run end-of-day summary draft", "EOD summary is complete and on time"],
        ["Day 8", "Handle one conflict-heavy day with support", "Escalation quality improves"],
        ["Day 9", "Own full-day workflow with review", "Can run workflow with minor corrections"],
        ["Day 10", "Checkpoint review and adjustment plan", "Ready for week-by-week training"],
    ]
    add_table(
        doc,
        ["Day", "Training Focus", "Done When"],
        first_10_rows,
        col_widths=[0.9, 3.8, 1.9],
        source="EA onboarding schedule draft (internal)",
    )

    doc.add_paragraph("12-Week Training Plan", style="DesignH1")
    week_rows = [
        ["Week 1", "Inbox and chat setup", "Use labels, star urgent, clean queue by end of day"],
        ["Week 2", "Priority decisions", "Sort items into now, today, this week"],
        ["Week 3", "Calendar basics", "Create events correctly and avoid overlap"],
        ["Week 4", "Meeting prep basics", "Send prep reminders and key details before meetings"],
        ["Week 5", "Stronger escalations", "Escalate with context and clear ask for help"],
        ["Week 6", "Follow-up discipline", "Track open items until owner confirms done"],
        ["Week 7", "Cross-channel tracking", "Connect email, chat, and calendar updates"],
        ["Week 8", "Busy-day simulation", "Handle high message volume without missing urgent items"],
        ["Week 9", "Proactive planning", "Prepare next-day watchlist before end of day"],
        ["Week 10", "Stakeholder communication", "Adjust message tone for each person"],
        ["Week 11", "Independent operation", "Run day independently with quality checks"],
        ["Week 12", "Final readiness week", "Deliver stable performance across full week"],
    ]
    add_table(
        doc,
        ["Week", "Skill Focus", "Practice Goal"],
        week_rows,
        col_widths=[0.9, 2.2, 3.5],
        source="EA 90-day implementation plan (internal)",
    )

    add_flowchart(
        doc,
        "Daily EA Workflow",
        [
            {"type": "start", "title": "Start of Day", "desc": "Check inbox, priority chats, and today's calendar."},
            {"type": "process", "title": "Sort Work", "desc": "Place each item into: now, today, or this week."},
            {"type": "decision", "title": "Urgent?", "desc": "If urgent, alert immediately. If not, schedule follow-up."},
            {"type": "process", "title": "Calendar Check", "desc": "Fix overlaps and send reminders early."},
            {"type": "support", "title": "Midday Reset", "desc": "Review open tasks and update priorities."},
            {"type": "end", "title": "End of Day", "desc": "Send summary and next-day watchlist."},
        ],
        source="Daily operating model for EA role (internal)",
    )

    add_flowchart(
        doc,
        "Panic Rescue Flow (Use Anytime)",
        [
            {"type": "start", "title": "Pause 30 Seconds", "desc": "Breathe, open notes, and focus on one thing first."},
            {"type": "decision", "title": "Is it urgent?", "desc": "Urgent means meeting/time-critical or business impact."},
            {"type": "process", "title": "Send quick alert", "desc": "Share facts first. Ask for help if needed."},
            {"type": "support", "title": "Take next safe action", "desc": "Book, reply, or escalate based on guidance."},
            {"type": "end", "title": "Log and continue", "desc": "Write what happened and resume normal workflow."},
        ],
        source="Beginner support protocol for high-pressure moments (internal)",
    )

    add_callout(
        doc,
        "Important",
        "In panic mode, speed and clarity matter more than perfect wording. A short clear alert is enough.",
        kind="warning",
    )

    doc.add_paragraph("Ready-to-Use Emergency Message Templates", style="DesignH2")
    panic_msg_rows = [
        [
            "Calendar conflict",
            "Hi Sam, conflict found: 2 meetings at 3:00 PM today. Please choose: move Client A to 4:00 PM or move Team Sync to tomorrow 10:00 AM.",
        ],
        [
            "Urgent email",
            "Hi Sam, urgent email from [name] needs response before 1:00 PM. I can draft a reply now if you want.",
        ],
        [
            "Chat escalation",
            "Hi Sam, [team] is blocked on [topic]. Decision needed today. I can set a 10-minute call if preferred.",
        ],
        [
            "Need help",
            "Hi Sam, I am not fully sure on this one. Here are the facts: [facts]. Can you guide me on best next step?",
        ],
    ]
    add_table(
        doc,
        ["Situation", "Message You Can Send"],
        panic_msg_rows,
        col_widths=[1.5, 4.9],
        source="EA communication templates for real-time use (internal)",
    )

    doc.add_paragraph("Email Playbook (Simple Version)", style="DesignH1")
    add_numbered_steps(
        doc,
        [
            "Check email in focused blocks: morning, midday, and late afternoon.",
            "Mark each email as now, today, or this week.",
            "For urgent emails, send a short alert with deadline.",
            "For non-urgent emails, queue with next action and owner.",
            "Before end of day, close open loops or report what is still pending.",
        ],
    )
    doc.add_paragraph("Email Examples", style="DesignH3")
    email_rows = [
        [
            "Needs quick decision",
            "Subject: Decision needed by 4:00 PM - Supplier quote. Body: Hi Sam, quote expires today 4:00 PM. Option A: approve now. Option B: ask for extension.",
        ],
        [
            "Inform only",
            "Subject: FYI - Meeting moved. Body: Client moved to Friday 10:00 AM. Calendar is already updated.",
        ],
        [
            "Pending follow-up",
            "Subject: Follow-up due tomorrow - Contract draft. Body: Waiting on legal comments. I will follow up at 9:00 AM.",
        ],
    ]
    add_table(
        doc,
        ["Case", "Simple Example"],
        email_rows,
        col_widths=[1.3, 5.1],
        source="Email practice examples for new EA onboarding (internal)",
    )

    doc.add_paragraph("Chat Monitoring Playbook", style="DesignH1")
    add_numbered_steps(
        doc,
        [
            "Pin priority spaces and check them first.",
            "When you see a decision request, log it immediately.",
            "Escalate only if it affects time, money, operations, or reputation.",
            "Keep one running list of open asks and update it every few hours.",
            "End the day by confirming what is done and what carries over.",
        ],
    )
    chat_rows = [
        [
            "Urgent blocker",
            "Hi Sam, Operations is blocked by approval for tonight's delivery. Need decision before 5:30 PM.",
        ],
        [
            "Info update",
            "Hi Sam, HR confirmed all documents received for tomorrow's onboarding.",
        ],
        [
            "Follow-up reminder",
            "Hi Sam, reminder: vendor call in 20 minutes. Agenda is in your calendar notes.",
        ],
    ]
    add_table(
        doc,
        ["Case", "Simple Chat Message"],
        chat_rows,
        col_widths=[1.3, 5.1],
        source="Chat support templates for EA daily operations (internal)",
    )

    doc.add_paragraph("Calendar Playbook", style="DesignH1")
    add_numbered_steps(
        doc,
        [
            "Every morning, review today and the next 2 days.",
            "Protect deep-work windows and transition time between meetings.",
            "Check attendee readiness and links one hour before key meetings.",
            "If conflict appears, propose 2 clear options quickly.",
            "After each meeting block, update action owners and due dates.",
        ],
    )
    cal_rows = [
        [
            "Conflict fix",
            "Hi Sam, 10:30 AM has overlap. Option 1: move internal review to 1:00 PM. Option 2: move supplier call to 11:30 AM.",
        ],
        [
            "Meeting reminder",
            "Reminder: 2:00 PM with Finance. Goal: approve budget draft. Needed file is attached in event notes.",
        ],
        [
            "Prep request",
            "For tomorrow 9:00 AM meeting, please confirm if we should prioritize hiring plan or cost plan.",
        ],
    ]
    add_table(
        doc,
        ["Case", "Simple Calendar Message"],
        cal_rows,
        col_widths=[1.3, 5.1],
        source="Calendar communication templates for EA role (internal)",
    )

    doc.add_paragraph("Gemini for Workspace (Practical Use)", style="DesignH1")
    doc.add_paragraph(
        "Gemini can save time. Use it for drafts and summaries, then always verify names, dates, and commitments before sending.",
        style="DesignBody",
    )
    gem_rows = [
        ["Summarize long email thread", "Summarize in 5 bullets: key request, deadline, owner, risk, next action."],
        ["Draft response options", "Draft 2 response options: short direct option and polite detailed option."],
        ["Meeting notes to actions", "Convert notes to action list with owner and due date."],
        ["Morning briefing", "Create a morning briefing from these items in priority order."],
    ]
    add_table(
        doc,
        ["Task", "Prompt You Can Use"],
        gem_rows,
        col_widths=[2.2, 4.2],
        source="Gemini usage guide for Workspace-support workflows (internal)",
    )
    add_callout(
        doc,
        "Gemini safety check",
        "Before sending any AI-assisted output: verify names, dates, amounts, and decision wording.",
        kind="support",
    )

    doc.add_paragraph("Daily and Weekly Rhythm", style="DesignH1")
    daily_rows = [
        ["8:30-9:00 AM", "Morning sweep", "Inbox, priority chats, calendar risk check"],
        ["9:00-9:15 AM", "Morning brief", "Send top priorities and urgent watch items"],
        ["11:30-11:45 AM", "Midday reset", "Review open items and update queue"],
        ["3:30-3:45 PM", "Calendar prep", "Prepare reminders and next meeting support"],
        ["5:00-5:20 PM", "End-of-day summary", "Send completed items, pending items, and next-day watchlist"],
    ]
    add_table(
        doc,
        ["Time Block", "Routine", "Output"],
        daily_rows,
        col_widths=[1.5, 1.6, 3.3],
        source="EA daily cadence model (internal)",
    )

    weekly_rows = [
        ["Monday", "Set week priorities with Sam", "Clear list of this week's priority outcomes"],
        ["Wednesday", "Midweek quality check", "Fix missed follow-ups and adjust workload"],
        ["Friday", "Weekly review", "Lessons learned and focus for next week"],
    ]
    add_table(
        doc,
        ["Day", "Review Action", "Expected Output"],
        weekly_rows,
        col_widths=[1.0, 2.0, 3.4],
        source="EA weekly coaching rhythm (internal)",
    )

    doc.add_paragraph("Progress Scorecard", style="DesignH1")
    score_rows = [
        ["Urgent item response time", "Within 5 minutes", "Use quick alert template and escalate early"],
        ["Missed urgent items", "Zero per week", "Review queue at midday and end of day"],
        ["Calendar conflicts", "Zero unresolved conflicts", "Check next 48 hours every morning"],
        ["Meeting prep readiness", "At least 95 percent", "Send reminders and notes 1 hour before"],
        ["Daily summary consistency", "100 percent workdays", "Use same format every day"],
    ]
    add_table(
        doc,
        ["Metric", "Good Target", "If Below Target"],
        score_rows,
        col_widths=[2.2, 1.5, 2.7],
        source="EA capability benchmarks for first 90 days (internal)",
    )

    add_callout(
        doc,
        "Final note",
        "This is a training program, not a pressure test. Progress comes from consistent habits, honest communication, and daily follow-through.",
        kind="info",
    )
    doc.add_paragraph("End of manual.", style="DesignCaption")

    doc.save(output)
    print(f"DOCX_WRITTEN {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
