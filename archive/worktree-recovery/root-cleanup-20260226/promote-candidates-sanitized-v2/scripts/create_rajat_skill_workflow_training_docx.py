#!/usr/bin/env python3
"""
Create Rajat-facing trainer DOCX for BEI skill workflow.

This version is intentionally trainer-style and excludes usage-frequency/counts.
"""

from __future__ import annotations

import argparse
import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PALETTE = {
    "primary": "04400A",
    "secondary": "C8900A",
    "accent": "801297",
    "light": "E6ECE7",
    "light_gold": "F8F0D9",
    "light_purple": "F2E5F5",
    "text": "1A1A1A",
    "text_light": "666666",
    "bg_alt": "E6ECE7",
    "white": "FFFFFF",
    "border": "D4D0C8",
}


PLAYBOOKS: Dict[str, dict] = {
    "plan-audit": {
        "module": "Plan Quality Gate",
        "purpose": "Force plan quality checks before any serious execution starts.",
        "when": "Run right after drafting a plan, before assigning implementation work.",
        "expected": "A blocker list that exposes gaps in sequencing, coverage, ownership, testing, and dependencies.",
        "steps": [
            "Draft the full plan in one file with clear sections.",
            "Invoke /plan-audit against that plan file.",
            "Review the blocker list and categorize: must-fix vs accepted risk.",
            "Patch the plan and related docs.",
            "Re-run /plan-audit until the plan is execution-ready.",
        ],
        "pitfalls": [
            "Running it on incomplete/partial plans.",
            "Treating blockers as optional suggestions.",
            "Skipping the re-audit pass after edits.",
        ],
        "before": ["/write-plan", "/architect-reviewer", "/design-review"],
        "after": ["/team-planner", "/teammates", "/tasks"],
    },
    "fact-check": {
        "module": "Evidence Verification",
        "purpose": "Validate claims against sources before final decisions or handoff.",
        "when": "Run before finalizing plans, reports, or architecture statements.",
        "expected": "Clean, source-backed statements with weak claims removed or corrected.",
        "steps": [
            "List critical claims (especially decisions and dependencies).",
            "Run /fact-check using the source-of-truth files.",
            "Patch unsupported statements and ambiguous wording.",
            "Re-check the final version before sharing.",
        ],
        "pitfalls": [
            "Checking only summaries and skipping detailed claims.",
            "Not re-running after edits.",
        ],
        "before": ["/plan-audit", "/write-plan"],
        "after": ["/docx-designer", "/teammates"],
    },
    "adms": {
        "module": "Biometric Device Ops",
        "purpose": "Handle attendance-machine and ADMS operational issues with structured diagnosis.",
        "when": "Run for biometric sync failures, enrollment issues, or device status anomalies.",
        "expected": "Clear root cause and a specific remediation path with verification step.",
        "steps": [
            "Define exact store/device and observed symptom.",
            "Run /adms checks for connectivity and data path.",
            "Apply targeted fix for the identified failure point.",
            "Verify the fix with a fresh operational check.",
        ],
        "pitfalls": [
            "Applying broad changes without device-level confirmation.",
            "Closing issue before re-verification.",
        ],
        "before": ["/find-conversations", "/tasks"],
        "after": ["/fact-check", "/test-full-cycle"],
    },
    "teammates": {
        "module": "Parallel Agent Coordination",
        "purpose": "Run scoped parallel work safely with clean handoffs.",
        "when": "Run when tasks can be split by domain/file ownership.",
        "expected": "Parallel throughput without file ownership collisions.",
        "steps": [
            "Split work into non-overlapping units.",
            "Spawn /teammates in waves, not all-at-once chaos.",
            "Collect outputs to files and integrate in a controlled pass.",
            "Run a verification gate after integration.",
        ],
        "pitfalls": [
            "Parallel agents editing same files.",
            "Merging outputs without validation pass.",
        ],
        "before": ["/team-planner", "/plan-audit"],
        "after": ["/tasks", "/test-full-cycle"],
    },
    "tasks": {
        "module": "Execution Loop",
        "purpose": "Track and drive execution until every blocker is actually closed.",
        "when": "Run for multi-step work with dependencies and blockers.",
        "expected": "No hidden pending work at handoff.",
        "steps": [
            "Create tasks with explicit done criteria.",
            "Execute top priority item and log blockers immediately.",
            "Route blocked items to supporting skills/agents.",
            "Close only after verification evidence exists.",
        ],
        "pitfalls": [
            "Marking done before verification.",
            "Leaving blocked tasks without owner/action.",
        ],
        "before": ["/teammates", "/plan-audit"],
        "after": ["/test-full-cycle", "/save"],
    },
    "test-full-cycle": {
        "module": "End-to-End Validation",
        "purpose": "Validate full workflows and iterate until stable.",
        "when": "Run after implementation changes and before release confidence sign-off.",
        "expected": "Workflow-level pass status with reproducible evidence.",
        "steps": [
            "Run scenario set for impacted modules.",
            "Capture failures with evidence.",
            "Patch root cause and rerun the affected scenarios.",
            "Repeat until cycle is clean.",
        ],
        "pitfalls": [
            "Happy-path-only checks.",
            "No retest after fix.",
        ],
        "before": ["/tasks", "/playwright"],
        "after": ["/fact-check", "/deploy-frappe"],
    },
    "playwright": {
        "module": "Browser Workflow Validation",
        "purpose": "Verify real UI behavior, console/network integrity, and workflow transitions.",
        "when": "Run when user-facing flow quality must be proven.",
        "expected": "Evidence-based pass/fail for front-end workflows.",
        "steps": [
            "Select exact workflow and role context.",
            "Execute browser checks for render + action + backend response.",
            "Capture screenshots/log traces for failures.",
            "Feed findings into fix-and-retest loop.",
        ],
        "pitfalls": [
            "Only checking if page loads.",
            "Ignoring console/network errors.",
        ],
        "before": ["/test-full-cycle", "/tasks"],
        "after": ["/fact-check", "/deploy-frappe"],
    },
    "deploy-frappe": {
        "module": "Production Release Control",
        "purpose": "Deploy backend changes with explicit verification discipline.",
        "when": "Run when API/backend changes must go live.",
        "expected": "Safe rollout with post-deploy verification and rollback readiness.",
        "steps": [
            "Validate release scope and prerequisites.",
            "Run /deploy-frappe with correct migration/build choices.",
            "Verify key endpoints and impacted workflows after release.",
            "Record release status and residual risks.",
        ],
        "pitfalls": [
            "Deploying without verifying prerequisites.",
            "Declaring success without runtime checks.",
        ],
        "before": ["/plan-audit", "/test-full-cycle"],
        "after": ["/playwright", "/fact-check"],
    },
    "rlm": {
        "module": "Large-Context Research",
        "purpose": "Process large sources safely and extract structured findings.",
        "when": "Run when context/files exceed normal prompt size.",
        "expected": "Consolidated findings that can feed planning and validation.",
        "steps": [
            "Define extraction objective and output schema.",
            "Run /rlm chunk and consolidation workflow.",
            "Validate critical findings before downstream use.",
            "Pass outputs into planning or quality gates.",
        ],
        "pitfalls": [
            "No clear extraction objective.",
            "Treating raw chunks as final conclusions.",
        ],
        "before": ["/find-conversations", "/tasks"],
        "after": ["/plan-audit", "/fact-check"],
    },
    "find-conversations": {
        "module": "Historical Context Recovery",
        "purpose": "Retrieve prior solutions before re-solving old problems.",
        "when": "Run at kickoff or when hitting recurring issue patterns.",
        "expected": "Faster setup with proven prior paths and fewer repeated mistakes.",
        "steps": [
            "Search with precise topic keywords.",
            "Identify relevant sessions only.",
            "Extract reusable approach patterns.",
            "Apply to current plan/execution path.",
        ],
        "pitfalls": [
            "Too broad search terms.",
            "Using history without validating relevance.",
        ],
        "before": ["/plan-audit", "/tasks"],
        "after": ["/teammates", "/rlm"],
    },
    "docx-designer": {
        "module": "Executive-Grade Documentation",
        "purpose": "Produce polished, trainer-friendly documents with clear hierarchy and flow.",
        "when": "Run for onboarding docs, handoff docs, and leadership-facing writeups.",
        "expected": "Readable, actionable document that people can execute from.",
        "steps": [
            "Define audience, objective, and flow first.",
            "Build sections with callouts, tables, and flowcharts.",
            "Run extraction QA and patch weak wording.",
            "Ship only after clarity and structure checks pass.",
        ],
        "pitfalls": [
            "Telemetry dump style instead of trainer guidance.",
            "Skipping post-generation QA.",
        ],
        "before": ["/fact-check", "/plan-audit"],
        "after": ["/save", "/restore"],
    },
    "xlsx": {
        "module": "Spreadsheet Evidence Work",
        "purpose": "Handle spreadsheet extraction, transformation, and reconciliation.",
        "when": "Run when structured evidence depends on xlsx/csv sources.",
        "expected": "Clean, validated data outputs ready for decisions and docs.",
        "steps": [
            "Define required fields and expected outputs.",
            "Run /xlsx transformations with schema checks.",
            "Validate totals and spot-check edge rows.",
            "Feed outputs into planning or report generation.",
        ],
        "pitfalls": [
            "Ignoring schema drift.",
            "No spot-check before using outputs downstream.",
        ],
        "before": ["/rlm", "/fact-check"],
        "after": ["/docx-designer", "/plan-audit"],
    },
}


SKILL_ORDER = [
    "plan-audit",
    "fact-check",
    "write-plan",
    "architect-reviewer",
    "design-review",
    "team-planner",
    "teammates",
    "tasks",
    "find-conversations",
    "rlm",
    "adms",
    "playwright",
    "test-full-cycle",
    "deploy-frappe",
    "docx-designer",
    "xlsx",
    "save",
    "restore",
]


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def skill_key(name: str) -> str:
    return name.strip().lower().lstrip("/")


def ensure_style(doc: Document, name: str, style_type: WD_STYLE_TYPE):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def setup_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(PALETTE["text"])

    h1 = ensure_style(doc, "DesignH1", WD_STYLE_TYPE.PARAGRAPH)
    h1.font.name = "Calibri Light"
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = rgb(PALETTE["primary"])
    h1.paragraph_format.space_before = Pt(20)
    h1.paragraph_format.space_after = Pt(8)

    h2 = ensure_style(doc, "DesignH2", WD_STYLE_TYPE.PARAGRAPH)
    h2.font.name = "Calibri"
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = rgb(PALETTE["secondary"])
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after = Pt(6)

    h3 = ensure_style(doc, "DesignH3", WD_STYLE_TYPE.PARAGRAPH)
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = rgb(PALETTE["primary"])
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)

    body = ensure_style(doc, "DesignBody", WD_STYLE_TYPE.PARAGRAPH)
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    body.font.color.rgb = rgb(PALETTE["text"])
    body.paragraph_format.line_spacing = 1.15
    body.paragraph_format.space_after = Pt(8)

    caption = ensure_style(doc, "DesignCaption", WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(PALETTE["text_light"])


def set_cell_bg(cell, color_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = OxmlElement(f"w:{edge}")
            for key, value in edge_data.items():
                tag.set(qn(f"w:{key}"), str(value))
            tc_borders.append(tag)
    tc_pr.append(tc_borders)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    color_hex: str = "1A1A1A",
    size: int = 10,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = rgb(color_hex)


def remove_table_borders(table) -> None:
    tbl = table._element
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tbl_pr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)


def add_callout(doc: Document, title: str, content: str, bg: str, border: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    set_cell_border(
        cell,
        left={"sz": 20, "val": "single", "color": border},
        top={"sz": 4, "val": "single", "color": border},
        right={"sz": 4, "val": "single", "color": border},
        bottom={"sz": 4, "val": "single", "color": border},
    )
    set_cell_text(cell, title, bold=True, color_hex=PALETTE["primary"], size=11)
    cell.add_paragraph(content, style="DesignBody")
    doc.add_paragraph("", style="DesignBody")


def add_flowchart(doc: Document, title: str, steps: List[dict]) -> None:
    doc.add_paragraph(title, style="DesignH2")
    table = doc.add_table(rows=len(steps) * 2 - 1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)

    fills = [PALETTE["light"], PALETTE["light_gold"], PALETTE["light_purple"]]

    sidx = 0
    for ridx in range(len(table.rows)):
        cell = table.rows[ridx].cells[0]
        if ridx % 2 == 0:
            step = steps[sidx]
            set_cell_bg(cell, fills[sidx % len(fills)])
            set_cell_border(
                cell,
                top={"sz": 10, "val": "single", "color": PALETTE["border"]},
                left={"sz": 10, "val": "single", "color": PALETTE["border"]},
                right={"sz": 10, "val": "single", "color": PALETTE["border"]},
                bottom={"sz": 10, "val": "single", "color": PALETTE["border"]},
            )
            set_cell_text(
                cell,
                f"{step['title']}\n{step['desc']}",
                bold=True,
                color_hex=PALETTE["text"],
                size=10,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            sidx += 1
        else:
            set_cell_text(
                cell,
                "↓",
                bold=True,
                color_hex=PALETTE["secondary"],
                size=16,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
    doc.add_paragraph("", style="DesignBody")


def add_bullets(doc: Document, lines: List[str]) -> None:
    for line in lines:
        doc.add_paragraph(f"- {line}", style="DesignBody")


def add_numbered(doc: Document, lines: List[str]) -> None:
    for idx, line in enumerate(lines, start=1):
        p = doc.add_paragraph(style="DesignBody")
        n = p.add_run(f"{idx}. ")
        n.bold = True
        n.font.color.rgb = rgb(PALETTE["accent"])
        p.add_run(line)


def add_cover_page(doc: Document, date_text: str) -> None:
    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(band)
    set_cell_bg(band.rows[0].cells[0], PALETTE["primary"])
    p = band.rows[0].cells[0].paragraphs[0]
    p.paragraph_format.space_before = Pt(68)
    p.paragraph_format.space_after = Pt(68)

    t1 = doc.add_paragraph(style="DesignH1")
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = t1.add_run("Rajat Skill Workflow Training")
    run1.font.size = Pt(30)
    run1.font.color.rgb = rgb(PALETTE["primary"])

    t2 = doc.add_paragraph(style="DesignH2")
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.add_run("How Sam Uses Skills Effectively in BEI ERP")

    t3 = doc.add_paragraph(style="DesignBody")
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t3.add_run(date_text)

    t4 = doc.add_paragraph(style="DesignCaption")
    t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t4.add_run("Trainer format. Sensitive communication/integration command families excluded.")
    doc.add_page_break()


def summarize_pairings(row: Optional[dict], direction: str) -> str:
    if not row:
        return "-"
    key = "recommended_pairings_outgoing" if direction == "next" else "recommended_pairings_incoming"
    items = row.get(key, [])
    if not items:
        return "-"
    return ", ".join(x.get("skill", "") for x in items[:3] if x.get("skill")) or "-"


def add_skill_module(doc: Document, skill: str, row: Optional[dict]) -> None:
    key = skill_key(skill)
    pb = PLAYBOOKS.get(key)
    if pb is None:
        pb = {
            "module": "Execution Module",
            "purpose": f"Use {skill} as part of the delivery flow when appropriate.",
            "when": "Use when task context indicates this skill is the best fit.",
            "expected": "Actionable output with clear next-step handoff.",
            "steps": [
                "Define objective and success criteria.",
                f"Run {skill} with explicit context.",
                "Patch issues and verify output quality.",
                "Handoff to next workflow step.",
            ],
            "pitfalls": [
                "Using the skill without a defined objective.",
                "Skipping verification.",
            ],
            "before": [],
            "after": [],
        }

    doc.add_paragraph(f"{skill} - {pb['module']}", style="DesignH2")
    doc.add_paragraph(pb["purpose"], style="DesignBody")

    add_callout(
        doc,
        "When to use",
        pb["when"],
        PALETTE["light_gold"],
        PALETTE["secondary"],
    )

    add_callout(
        doc,
        "Expected result",
        pb["expected"],
        PALETTE["light"],
        PALETTE["primary"],
    )

    doc.add_paragraph("Recommended run steps", style="DesignH3")
    add_numbered(doc, pb["steps"])

    doc.add_paragraph("Common failure patterns", style="DesignH3")
    add_bullets(doc, pb["pitfalls"])

    prev_text = summarize_pairings(row, "prev")
    next_text = summarize_pairings(row, "next")

    if prev_text == "-" and pb.get("before"):
        prev_text = ", ".join(pb["before"])
    if next_text == "-" and pb.get("after"):
        next_text = ", ".join(pb["after"])

    doc.add_paragraph(f"Typical previous-step pairings: {prev_text}", style="DesignBody")
    doc.add_paragraph(f"Typical next-step pairings: {next_text}", style="DesignBody")
    doc.add_paragraph("", style="DesignBody")


def main() -> int:
    parser = argparse.ArgumentParser(description="Create Rajat trainer DOCX for skill workflow.")
    parser.add_argument(
        "--dataset-json",
        default="scratchpad/reports/skill_training_dataset_full.json",
        help="Input dataset JSON from build_skill_training_dataset.py",
    )
    parser.add_argument(
        "--output-docx",
        default="docs/reports/RAJAT_SKILL_WORKFLOW_TRAINING.docx",
        help="Output DOCX path",
    )
    args = parser.parse_args()

    dataset_path = Path(args.dataset_json)
    if not dataset_path.exists():
        raise SystemExit(f"Dataset file not found: {dataset_path}")

    data = json.loads(dataset_path.read_text(encoding="utf-8"))
    rows = data.get("skills", [])
    if not rows:
        raise SystemExit("Dataset has no skills.")

    by_skill = {skill_key(r.get("skill", "")): r for r in rows}

    selected: List[str] = []
    for key in SKILL_ORDER:
        if key in by_skill:
            selected.append(f"/{key}")

    if not selected:
        selected = [r["skill"] for r in rows[:12]]

    output_path = Path(args.output_docx)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    setup_page(doc)
    setup_styles(doc)

    generated = datetime.now(timezone.utc).strftime("Generated: %Y-%m-%d")
    add_cover_page(doc, generated)

    doc.add_paragraph("How Rajat Should Operate", style="DesignH1")
    doc.add_paragraph(
        "Use this as a practical operator manual: choose the right skill at the right stage, "
        "chain it with the right next step, and always close with a verification gate.",
        style="DesignBody",
    )

    add_callout(
        doc,
        "Trainer Rule",
        "Run /plan-audit immediately after plan drafting. Treat it as a hard gate before execution.",
        PALETTE["light_gold"],
        PALETTE["secondary"],
    )

    add_flowchart(
        doc,
        "Main Delivery Flow",
        [
            {"title": "Plan", "desc": "/write-plan and architecture shaping"},
            {"title": "Audit", "desc": "/plan-audit to force gap discovery"},
            {"title": "Coordinate", "desc": "/team-planner, /teammates, /tasks"},
            {"title": "Build + Domain", "desc": "Implement work and run domain ops such as /adms"},
            {"title": "Validate", "desc": "/playwright, /test-full-cycle, /fact-check"},
            {"title": "Release + Handoff", "desc": "/deploy-frappe then document with /docx-designer"},
        ],
    )

    add_flowchart(
        doc,
        "Mandatory /plan-audit Gate Flow",
        [
            {"title": "Draft plan", "desc": "Create full plan with scope and sequence"},
            {"title": "Run /plan-audit", "desc": "Surface blockers and architecture gaps"},
            {"title": "Patch", "desc": "Fix gaps in docs, ownership, testing, dependencies"},
            {"title": "Re-run audit", "desc": "Repeat until execution-ready"},
            {"title": "Execute", "desc": "Only then proceed to /teammates and /tasks"},
        ],
    )

    doc.add_paragraph("Skill Training Modules", style="DesignH1")
    for skill in selected:
        row = by_skill.get(skill_key(skill))
        add_skill_module(doc, skill, row)

    doc.add_page_break()
    doc.add_paragraph("First Week Training Path", style="DesignH1")
    add_numbered(
        doc,
        [
            "Day 1: Practice plan drafting then hard-gate with /plan-audit.",
            "Day 2: Run /teammates + /tasks on a scoped execution set.",
            "Day 3: Validate workflows with /playwright and /test-full-cycle.",
            "Day 4: Practice release discipline with /deploy-frappe.",
            "Day 5: Create final handoff documentation using /docx-designer.",
        ],
    )

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Internal BEI ERP trainer guide. Non-sensitive skill content only.")
    fr.font.size = Pt(8)
    fr.font.color.rgb = rgb(PALETTE["text_light"])

    doc.save(output_path)
    print(f"DOCX written: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
