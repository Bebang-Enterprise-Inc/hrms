"""S210 designed-guide generator — BEI-branded DOCX with custom layout
per audience. Replaces the lightweight markdown-to-docx generator.

Each guide has its own design (not a generic MD converter):
  - Payment-speed benefit callout at the top of every guide
  - Clickable hyperlinks in BEI green
  - Big button-style CTAs for key actions
  - Numbered step blocks with green circles
  - Tables with zebra-striping, green headers
  - External-facing guides (Supplier FAQ) designed for email attachment
  - Master index one-pager cross-references everything

Run:
    python output/s210/generate_designed_guides.py
"""
from __future__ import annotations

import pathlib
import sys
from typing import Iterable

sys.stdout.reconfigure(encoding='utf-8')
sys.path.insert(0, r'F:\Dropbox\Projects\BEI-ERP\.claude\skills\docx-designer-bei-erp\scripts')

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from bei_docx import (
    PALETTE, add_bei_header, add_callout_box, add_run,
    hex_rgb, remove_cell_borders, set_cell_borders, set_cell_margins,
    set_cell_shading, set_paragraph_spacing, setup_doc,
)

GUIDES_DIR = pathlib.Path(r'F:\Dropbox\Projects\BEI-ERP-s210e\output\s210\guides')

# ------------------------------------------------------------------
# Live resource URLs (canonical; edit here only)
# ------------------------------------------------------------------
FORM_URL = 'https://docs.google.com/forms/d/e/1FAIpQLSc3UC9f_3gefDYNgpOqNx7UCw_5BDrRh9T8-GQeyHHWSxdITw/viewform'
SHEET_A_URL = 'https://docs.google.com/spreadsheets/d/1dambmiLzSMWOQun7MCymK4nHpuqrarFCAOK0G9-6oIU/edit'
SHEET_B_URL = 'https://docs.google.com/spreadsheets/d/10fqnvF_uDl5ky3MkvXUmWvZ1fYat_p6XFGmVFc3vqrw/edit'
SHEET_C_URL = 'https://docs.google.com/spreadsheets/d/1_Ir5O5AW7hOjcvCTXsP06cF3sai9hcefDFrBOTRHOh0/edit'
SHEET_D_URL = 'https://docs.google.com/spreadsheets/d/1mbJiLW9M9e-AmrXSRRTtbRP-xKI16ah5rakOt6qv2As/edit'
SCRIPT_EDITOR_URL = 'https://script.google.com/d/1lsvOlv1rGEvXl_1zms4SURlsLUZk7CxRhg2NyBDrDHh4fDjuioFZhi2S/edit'
SCHEDULER_CONSOLE_URL = 'https://console.cloud.google.com/cloudscheduler?project=quiet-walker-475722-s2'


# ------------------------------------------------------------------
# Hyperlink + custom component helpers
# ------------------------------------------------------------------

def add_hyperlink(paragraph, url, text, color=PALETTE['primary'],
                  bold=False, size=None, underline=True):
    """Add a clickable hyperlink run. BEI green by default, underlined."""
    part = paragraph.part
    r_id = part.relate_to(url,
                          'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                          is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # color
    col = OxmlElement('w:color')
    col.set(qn('w:val'), color)
    rPr.append(col)

    # underline
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    # font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)

    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    if size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(size * 2)))  # half-points
        rPr.append(sz)

    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_button(doc, text, url, width_inches=3.2, bg=None, fg=PALETTE['white'], size=14):
    """Big button-style CTA: rounded filled cell with a hyperlink."""
    bg = bg or PALETTE['primary']
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(width_inches)
    cell = table.cell(0, 0)
    cell.width = Inches(width_inches)
    set_cell_shading(cell, bg)
    set_cell_borders(cell, bg, size='18')
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, 0, 0)

    # Use hyperlink on the button but force white text via colour arg
    add_hyperlink(p, url, text, color=fg.lstrip('#'),
                  bold=True, size=size, underline=False)

    # Spacer below
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, 0, 80)


def add_big_benefit_callout(doc, title, body):
    """Gold-outlined callout box used at the top of each guide to surface
    the payment-speed benefit."""
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = False
    box.columns[0].width = Inches(6.0)
    cell = box.cell(0, 0)
    set_cell_shading(cell, PALETTE['light_gold'])
    set_cell_borders(cell, PALETTE['secondary'], size='18')
    set_cell_margins(cell, top=180, bottom=180, left=240, right=240)

    p = cell.paragraphs[0]
    set_paragraph_spacing(p, 0, 80)
    add_run(p, title, bold=True, size=14, color=PALETTE['secondary'])

    for line in body:
        pp = cell.add_paragraph()
        set_paragraph_spacing(pp, 0, 60)
        add_run(pp, line, size=11, color=PALETTE['text'])

    doc.add_paragraph()


def add_numbered_step(doc, num, title, body_text):
    """A numbered step block: green circle with number, bold title,
    and an indented body paragraph."""
    # Wrap in a 2-col table: circle col, content col
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    t.columns[0].width = Inches(0.55)
    t.columns[1].width = Inches(5.45)

    # Left: filled circle (approximated via colored cell with number)
    circle = t.cell(0, 0)
    set_cell_shading(circle, PALETTE['primary'])
    set_cell_borders(circle, PALETTE['primary'], size='12')
    set_cell_margins(circle, top=60, bottom=60, left=0, right=0)
    cp = circle.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cp, str(num), bold=True, size=15, color=PALETTE['white'])
    set_paragraph_spacing(cp, 0, 0)

    # Right: title + body
    content = t.cell(0, 1)
    remove_cell_borders(content)
    set_cell_margins(content, top=40, bottom=40, left=140, right=40)
    tp = content.paragraphs[0]
    add_run(tp, title, bold=True, size=12, color=PALETTE['primary'])
    set_paragraph_spacing(tp, 0, 40)
    if body_text:
        bp = content.add_paragraph()
        add_run(bp, body_text, size=11, color=PALETTE['text'])
        set_paragraph_spacing(bp, 0, 60)

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, 0, 60)


def add_section_header(doc, text, before=200):
    """Green bold section header with gold underline band."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before, 40)
    add_run(p, text, bold=True, size=14, color=PALETTE['primary'])

    hr = doc.add_table(rows=1, cols=1)
    hr.alignment = WD_TABLE_ALIGNMENT.LEFT
    hr.autofit = False
    hr.columns[0].width = Inches(6.0)
    c = hr.cell(0, 0)
    set_cell_shading(c, PALETTE['secondary'])
    set_cell_margins(c, top=0, bottom=0, left=0, right=0)
    hp = c.paragraphs[0]
    add_run(hp, ' ', size=1)
    set_paragraph_spacing(hp, 0, 0)
    doc.add_paragraph()


def add_zebra_table(doc, headers, rows, col_widths=None, first_col_bold=False):
    """A polished table with green header + zebra body rows."""
    n = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    if col_widths:
        for i, w in enumerate(col_widths):
            t.columns[i].width = Inches(w)

    # Header row
    for i, h in enumerate(headers):
        c = t.cell(0, i)
        set_cell_shading(c, PALETTE['primary'])
        set_cell_borders(c, PALETTE['primary'], size='6')
        set_cell_margins(c, top=70, bottom=70, left=100, right=100)
        p = c.paragraphs[0]
        add_run(p, h, bold=True, size=10.5, color=PALETTE['white'])
        set_paragraph_spacing(p, 0, 0)

    for ri, row in enumerate(rows):
        bg = PALETTE['white'] if ri % 2 == 0 else PALETTE['bg_alt']
        for i, val in enumerate(row):
            c = t.cell(ri + 1, i)
            set_cell_shading(c, bg)
            set_cell_borders(c, PALETTE['border'], size='4')
            set_cell_margins(c, top=60, bottom=60, left=100, right=100)
            p = c.paragraphs[0]
            add_run(p, str(val), bold=(first_col_bold and i == 0),
                    size=10, color=PALETTE['text'])
            set_paragraph_spacing(p, 0, 0)

    doc.add_paragraph()


def add_title(doc, title, subtitle=None):
    """Large centered title + optional gray subtitle."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, 0, 20)
    add_run(p, title, bold=True, size=22, color=PALETTE['primary'])

    if subtitle:
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(s, 0, 160)
        add_run(s, subtitle, italic=True, size=12, color=PALETTE['text_light'])
    else:
        doc.add_paragraph()


def add_body(doc, text, size=11):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 40, 80)
    add_run(p, text, size=size, color=PALETTE['text'])
    return p


def add_footer_note(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, 240, 0)
    add_run(p, text, italic=True, size=9, color=PALETTE['text_light'])


# ------------------------------------------------------------------
# Guide builders
# ------------------------------------------------------------------

def build_supplier_faq():
    doc = Document()
    setup_doc(doc, font_size=11)
    add_bei_header(doc)

    add_title(
        doc,
        'Upload Your Sales Invoice',
        'Fastest path to getting paid — for BEI suppliers',
    )

    add_big_benefit_callout(
        doc,
        'Why this matters for you',
        [
            'BEI used to wait for paper SIs to travel from the 3PL warehouse '
            'to our accounting team before starting your payment — often adding '
            'weeks of delay.',
            '',
            'Upload your SI the moment you deliver. We match it automatically '
            'and queue your payment for release on your contracted Net terms '
            '(15 / 30 / 45 / 60) counted from delivery date. No more paper-chase.',
        ],
    )

    add_section_header(doc, 'How to upload', before=160)

    add_button(doc, 'Open Supplier Upload Form  →', FORM_URL, width_inches=3.8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, 0, 80)
    add_run(p, 'Or type in your browser:  ', size=10, color=PALETTE['text_light'])
    add_hyperlink(p, FORM_URL, FORM_URL, color=PALETTE['primary'], size=10)

    add_numbered_step(doc, 1, 'Pick the warehouse',
                      'Choose 3MD, Pinnacle, or Shaw BLVD — whichever warehouse you delivered to.')
    add_numbered_step(doc, 2, 'PO Number',
                      'Copy from your BEI Purchase Order (e.g. PO-2026-1234).')
    add_numbered_step(doc, 3, 'SI Number, SI Date, Amount',
                      'Type exactly as printed on your Sales Invoice.')
    add_numbered_step(doc, 4, 'Upload SI Copy',
                      'Tap the file button, pick your PDF or phone photo (max 10 MB).')
    add_numbered_step(doc, 5, 'Submit',
                      "That's it — 30 to 60 seconds. Your payment queue starts immediately.")

    add_section_header(doc, 'Frequently asked')

    faqs = [
        ('Can I upload a phone photo instead of a PDF?',
         'Yes. PDF, JPG, and PNG are all accepted up to 10 MB. Just make sure '
         'your SI number, date, line items, and totals are legible.'),
        ('Do I still send the paper SI?',
         'Yes — keep sending paper as you do today. Upload is in addition to '
         'paper, not a replacement. Upload is what speeds up your payment.'),
        ('How fast will I be paid?',
         'Per your contracted payment terms (Net 15, 30, 45, or 60) counted '
         'from the delivery date. Uploading removes the paper-chasing delays '
         'that used to push actual payment past your contract.'),
        ('My upload says "Orphan" — what does that mean?',
         'The system couldn\'t match your upload to a delivery yet. Usually: '
         '(1) you uploaded before the warehouse logged the delivery — we '
         'retry every minute, or (2) a typo in PO or SI — double-check '
         'against your SI. Your upload is never lost.'),
        ('Can I share the link with my team?',
         "Yes. Same link for everyone — share with anyone on your team who "
         'handles BEI invoices. Just don\'t upload the same SI twice.'),
        ('I made a typo — can I fix it?',
         'Submit again with the correct values. Add a note in the Notes field: '
         '"Correction for PO-xxxx SI-yyyy" so our team knows which submission '
         'to keep.'),
        ('Is my upload secure?',
         'Yes. Only BEI staff (accounting, IT, CEO) can see your PDF. It\'s '
         'stored in BEI\'s private Google Drive — not public, not indexed, '
         'not shared externally.'),
        ('What if my delivery covers multiple warehouses on the same day?',
         'Upload the SI once. Our system matches the same SI against every '
         'corresponding delivery record automatically.'),
        ('Where do I find the form link if I lose it?',
         'Bookmark this document, or email Cayla (cayla@bebang.ph) for it '
         'again. Same link always works.'),
    ]
    for q, a in faqs:
        qp = doc.add_paragraph()
        set_paragraph_spacing(qp, 80, 20)
        add_run(qp, q, bold=True, size=11.5, color=PALETTE['secondary'])
        ap = doc.add_paragraph()
        set_paragraph_spacing(ap, 0, 80)
        add_run(ap, a, size=11, color=PALETTE['text'])

    add_section_header(doc, 'Contact', before=160)
    add_body(doc,
             'Questions about the upload tool: sam@bebang.ph (subject "BEI Supplier SI Upload").')
    add_body(doc,
             'Payment status / AP queries: continue using your existing BEI contact.')

    add_footer_note(doc, 'Share this with anyone on your team who handles BEI invoices.  |  v2026-04-21')

    path = GUIDES_DIR / '3_SUPPLIER_FAQ.docx'
    doc.save(str(path))
    return path, 'BEI Supplier SI Upload — Fastest Path to Payment'


def build_dock_card():
    doc = Document()
    setup_doc(doc, font_size=11)
    add_bei_header(doc)

    add_title(
        doc,
        'Receiving Quick Card',
        '3MD / Pinnacle dock staff — print, laminate, post at the receiving area',
    )

    add_big_benefit_callout(
        doc,
        'Why this matters',
        [
            'Every delivery you log here triggers BEI\'s payment process within a '
            'minute. The faster + more accurate you log, the faster suppliers get '
            'paid — and the fewer follow-up calls we make to your warehouse.',
        ],
    )

    add_section_header(doc, 'How to log a delivery', before=120)

    add_numbered_step(doc, 1, 'Open your sheet',
                      'Sheet A (3MD) or Sheet B (Pinnacle) — bookmarked on the receiving terminal.')
    add_numbered_step(doc, 2, 'Click Receipts tab',
                      'Scroll to the first empty row. One row per item per delivery.')
    add_numbered_step(doc, 3, 'Fill 16 columns left-to-right',
                      'Timestamp, RR#, PO#, Supplier, Material, Qty, UoM, SI#, Trucker, Plate, '
                      'Production Date, Expiration Date, Received By, Notes. No photos.')
    add_numbered_step(doc, 4, 'Press Enter — done',
                      "Within ~1 minute BEI's system picks up the row and processes it automatically.")

    add_section_header(doc, 'Multi-item delivery')
    add_body(doc,
             'One truck + one PO + one SI covering multiple materials? Type the header fields '
             '(RR#, PO#, Supplier, SI#, Trucker, Plate) on the first line only. Leave them '
             'blank on lines 2, 3, etc. — the system auto-inherits from the first line.')

    add_zebra_table(
        doc,
        ['RR#', 'PO#', 'Supplier', 'Material', 'Qty', 'UoM'],
        [
            ['RR-0041', 'PO-2026-1234', 'DIMAX', 'FLOUR-25KG', '10', 'BAG'],
            ['(blank)', '(blank)', '(blank)', 'SUGAR-50KG', '5', 'BAG'],
            ['(blank)', '(blank)', '(blank)', 'YEAST-1KG', '2', 'PC'],
        ],
        col_widths=[0.9, 1.3, 1.0, 1.4, 0.6, 0.6],
    )

    add_body(doc,
             'Never leave blank on any row: Material Code, Qty, UoM, Production Date, '
             'Expiration Date. Those must be typed for every line.',
             size=10.5)

    add_section_header(doc, 'Common issues')

    add_zebra_table(
        doc,
        ['Problem', 'What to do'],
        [
            ['PO dropdown doesn\'t show my PO', 'Call Ian — PO may not be routed here'],
            ['Supplier on SI doesn\'t match PO', 'Stop. Call procurement. Do NOT submit.'],
            ['Qty exceeds PO balance', 'Stop. Call procurement first.'],
            ['Multiple production dates for same material', 'Log as separate rows — one per batch'],
            ['Cold chain item without temp log', 'Log normally; add "cold chain broken" in Notes'],
        ],
        col_widths=[2.6, 3.3],
        first_col_bold=True,
    )

    add_section_header(doc, 'Quick access')

    p = doc.add_paragraph()
    set_paragraph_spacing(p, 40, 80)
    add_run(p, '3MD sheet: ', bold=True, size=11, color=PALETTE['text'])
    add_hyperlink(p, SHEET_A_URL, 'BEI 3MD Receiving Log 2026', size=11)

    p2 = doc.add_paragraph()
    set_paragraph_spacing(p2, 0, 80)
    add_run(p2, 'Pinnacle sheet: ', bold=True, size=11, color=PALETTE['text'])
    add_hyperlink(p2, SHEET_B_URL, 'BEI Pinnacle Receiving Log 2026', size=11)

    add_section_header(doc, 'Contacts')
    p3 = doc.add_paragraph()
    set_paragraph_spacing(p3, 40, 40)
    add_run(p3, 'Ian Dionisio (primary): ', bold=True, size=11)
    add_run(p3, 'ian@bebang.ph', size=11, color=PALETTE['primary'])
    p4 = doc.add_paragraph()
    set_paragraph_spacing(p4, 0, 40)
    add_run(p4, 'Jay Sumagui (backup): ', bold=True, size=11)
    add_run(p4, 'jay@bebang.ph', size=11, color=PALETTE['primary'])

    add_footer_note(doc, 'v2026-04-21 · Supersedes all prior dock instructions')

    path = GUIDES_DIR / '2_3PL_DOCK_QUICK_CARD.docx'
    doc.save(str(path))
    return path, '3PL Dock Receiving Quick Card — 3MD & Pinnacle'


def build_supplier_rollout():
    doc = Document()
    setup_doc(doc, font_size=11)
    add_bei_header(doc)

    add_title(
        doc,
        'Supplier SI Upload — Rollout Playbook',
        'Owner: Cayla',
    )

    add_big_benefit_callout(
        doc,
        'Why we\'re rolling this out',
        [
            'Suppliers who upload their SI consistently get paid on their '
            'contracted Net terms counted from delivery date — no paper-SI '
            'delays eating into our DSO.',
            '',
            'Your rollout target: 80% SI match rate within 4 weeks. Below that, '
            'payment releases start lagging contracts and supplier goodwill '
            'suffers.',
        ],
    )

    add_section_header(doc, 'What you have')
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 40, 60)
    add_run(p, 'Form URL (same for all 98 suppliers): ', bold=True, size=11)
    add_hyperlink(p, FORM_URL, FORM_URL, size=11)

    add_body(doc,
             'No per-supplier URL generation — supplier identity is derived from the PO Number '
             'they type. One link, one QR for the dock, one FAQ to attach.')

    add_section_header(doc, 'Rollout waves')

    add_zebra_table(
        doc,
        ['Wave', 'Who', 'Timing', 'Approach'],
        [
            ['1', 'Top 10 Tier A by volume', 'Week 1', 'Individual emails, 48h follow-up, personal call on any non-response'],
            ['2', 'Rest of Tier A (~50-70)', 'Week 2', 'Batch email. Expect 60-70% response rate.'],
            ['3', 'Tier B/C', 'Week 3-4', 'Can batch. One follow-up after a week.'],
            ['4', 'Stragglers', 'Week 5+', 'Direct call to anyone who hasn\'t uploaded once.'],
        ],
        col_widths=[0.6, 2.0, 0.9, 2.7],
        first_col_bold=True,
    )

    add_section_header(doc, 'Email template')
    add_callout_box(
        doc,
        [
            ('Subject: BEI Supplier SI Upload — fastest path to payment',
             True, 11),
            ('', False, 1),
            ('Hi [Supplier Contact Name],', False, 11),
            ('', False, 1),
            ('BEI has moved to a new process that will get your invoices paid faster. '
             'For every delivery you make to BEI or our 3PL warehouses (3MD, Pinnacle), '
             'please upload your Sales Invoice at the link below right after delivery.',
             False, 10.5),
            ('', False, 1),
            (f'Upload link (bookmark this — same link every delivery):', False, 10.5),
            (FORM_URL, False, 10),
            ('', False, 1),
            ('Attached FAQ answers common questions.', False, 10.5),
            ('', False, 1),
            ('— Cayla | cayla@bebang.ph', True, 10.5),
        ],
        bg_color=PALETTE['bg_alt'],
        border_color=PALETTE['primary'],
    )

    add_section_header(doc, 'QR code for the dock')
    add_body(doc,
             'One QR, same for all suppliers. Generate from the form URL, print A5, '
             'laminate, post at each of 3MD, Pinnacle, and Shaw receiving stations.')

    add_section_header(doc, 'Tracking adoption')
    add_body(doc, 'Every Monday 09:00 PHT, open Sheet C Dashboard:')
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 80)
    add_hyperlink(p, SHEET_C_URL, 'BEI Receiving Master 2026 — Dashboard', size=11, bold=True)

    add_body(doc,
             'Target: SI match rate > 60% by end of week 2, > 80% by end of week 4. '
             'For suppliers still not uploading, filter 02_All_Receipts_Consolidated '
             'for SI_Matched=FALSE, group by Supplier, call the top 5 non-adopters.')

    add_section_header(doc, 'Common objections')
    add_zebra_table(
        doc,
        ['Supplier says', 'You say'],
        [
            ['"I don\'t have a scanner"', 'Phone photo is fine — PDF/JPG/PNG accepted.'],
            ['"Our accounts staff don\'t use Gmail"', 'Form is public — any browser works. No Google account needed.'],
            ['"Can I still send paper SI?"', 'Yes, keep doing that. Upload is in addition, not instead.'],
            ['"My PO number is long / has dashes"', 'Copy-paste from the PO exactly as printed. Form accepts any string.'],
            ['"What if I typo after submitting?"', 'Upload again; add "Correction for PO-xxxx" in Notes.'],
            ['"I\'m worried about my PDF being seen"', 'Only BEI staff have access. Not public, not indexed.'],
            ['"How do you know it\'s my company?"', 'We match the PO Number you type to your PO in our system. No supplier list exposed.'],
        ],
        col_widths=[2.3, 3.7],
    )

    add_section_header(doc, 'Escalation')
    add_body(doc, '• Tier A supplier refusing → Cayla → Ian → Sam')
    add_body(doc, '• Form/technical issue → Sam')
    add_body(doc, '• Supplier says we\'re not paying after upload → Luwi (AP)')

    add_footer_note(doc, 'v2026-04-21')

    path = GUIDES_DIR / '4_SUPPLIER_ROLLOUT_GUIDE.docx'
    doc.save(str(path))
    return path, 'Supplier SI Upload — Cayla Rollout Playbook'


def build_ian_playbook():
    doc = Document()
    setup_doc(doc, font_size=11)
    add_bei_header(doc)

    add_title(doc, 'Receiving Ops — Daily Playbook', 'Owner: Ian Dionisio')

    add_big_benefit_callout(
        doc,
        'Your job in one line',
        [
            'Keep the Receipt-Based Payment pipeline clean so suppliers get paid '
            'on their contracted terms. Every unresolved orphan SI or stale DR '
            'directly delays a supplier\'s payment.',
        ],
    )

    add_section_header(doc, 'Quick sheet access')
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 40, 40)
    add_run(p, 'Sheet C (BEI master): ', bold=True, size=11)
    add_hyperlink(p, SHEET_C_URL, 'BEI Receiving Master 2026', size=11)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 40)
    add_run(p, 'Sheet A (3MD): ', bold=True, size=11)
    add_hyperlink(p, SHEET_A_URL, 'BEI 3MD Receiving Log 2026', size=11)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 40)
    add_run(p, 'Sheet B (Pinnacle): ', bold=True, size=11)
    add_hyperlink(p, SHEET_B_URL, 'BEI Pinnacle Receiving Log 2026', size=11)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 40)
    add_run(p, 'Scheduler health: ', bold=True, size=11)
    add_hyperlink(p, SCHEDULER_CONSOLE_URL, 'GCP Cloud Scheduler Console', size=11)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 120)
    add_run(p, 'Apps Script editor: ', bold=True, size=11)
    add_hyperlink(p, SCRIPT_EDITOR_URL, 'S210 Master Handler', size=11)

    add_section_header(doc, 'Morning checklist (07:30–07:45)')

    add_numbered_step(doc, 1, 'Read the CEO daily email',
                      'Subject "[BEI Receiving] Daily KPI digest" — lands at 07:00 PHT in your inbox.')
    add_numbered_step(doc, 2, 'Scan Sheet C Dashboard',
                      'Live formulas. If any cell reads #N/A or #ERROR, ping Sam immediately.')
    add_numbered_step(doc, 3, 'Clear Match Queue (04_Match_Queue)',
                      'Each row = supplier SI that didn\'t auto-match a DR. Use Warehouse column '
                      'to know which 3PL to chase. Resolve or dismiss each one.')
    add_numbered_step(doc, 4, 'Resolve Variance Queue (05_Variance_Queue)',
                      'DRs that failed validation or aged > 72h. Fix upstream, mark Resolved.')

    add_section_header(doc, 'KPI thresholds')

    add_zebra_table(
        doc,
        ['KPI', 'Healthy', 'Yellow flag', 'Red — act now'],
        [
            ['Today\'s receipts (3MD/Pinnacle/Shaw)', 'Matches expected', '1-2 short', 'Zero — scheduler broken or 3PL access lost'],
            ['SI match rate', '> 80%', '60-80%', '< 60% — chase Tier A non-uploaders'],
            ['Stale DR count (>72h)', '< 5', '5-20', '> 20 — structural issue'],
            ['Pending GR depth', '< 10', '10-30', '> 30 — Ashish AppSheet may be stuck'],
            ['Orphan SI count', '< 5', '5-15', '> 15 — Match Queue backlog'],
        ],
        col_widths=[1.9, 1.1, 1.1, 1.9],
        first_col_bold=True,
    )

    add_section_header(doc, 'Afternoon checklist (16:30–16:40)')

    add_numbered_step(doc, 1, 'Confirm expected deliveries landed',
                      'Filter 08_Full_Open_POs by Delivery Needed By = today. Cross-check with '
                      '02_All_Receipts_Consolidated. Call suppliers for missing drops.')
    add_numbered_step(doc, 2, 'Review SCM Chat space',
                      'Scan day\'s automated notifications. Look for unusually high qty, off-hours '
                      'deliveries, or repeated validation failures (pattern = systemic issue).')
    add_numbered_step(doc, 3, 'Verify scheduler health',
                      'GCP Console → all 4 s210-* jobs should be ENABLED with recent success. '
                      'Any red = ping Sam.')

    add_section_header(doc, 'Weekly Monday 09:00')
    add_body(doc,
             '1. Supplier adoption review — compute per-supplier SI match rate from last 7 days, '
             'hand top 5 non-adopters to Cayla.')
    add_body(doc,
             '2. Scheduler job report — GCP Console, History tab for each s210-* job. Zero '
             'AttemptFailed in the last 7 days. Median pollAll response < 10s.')
    add_body(doc,
             '3. Sheet C growth — if 02_All_Receipts_Consolidated > 5000 rows, start planning '
             'year-over-year rotation to 2027 sheet.')

    add_section_header(doc, 'Emergency rollback')
    add_body(doc,
             'If automation misbehaves: GCP Console → Cloud Scheduler → pause all 4 s210-* jobs. '
             'Sheets A/B/D still accept 3PL entries; they just don\'t propagate to Sheet C until '
             'resumed. Notify Sam. Do NOT delete anything. Resume once root cause is fixed — '
             'pollAll will catch up all unprocessed rows.')

    add_section_header(doc, 'Escalation')
    add_zebra_table(
        doc,
        ['Issue', 'Contact'],
        [
            ['Apps Script error (Chat spam, Variance flood, Dashboard #ERROR)', 'Sam'],
            ['Supplier refusing to upload', 'Cayla'],
            ['PO issues (wrong balance, closed POs, routing)', 'Mae → Luwi'],
            ['3PL not logging receipts', 'You (3MD) / Jay (Pinnacle)'],
            ['BIR / compliance question', 'Denise + Ashish'],
        ],
        col_widths=[3.8, 2.2],
        first_col_bold=True,
    )

    add_footer_note(doc, 'v2026-04-21')

    path = GUIDES_DIR / '5_IAN_DAILY_OPS_PLAYBOOK.docx'
    doc.save(str(path))
    return path, 'Receiving Ops — Ian Daily Playbook'


def build_finance_guide():
    doc = Document()
    setup_doc(doc, font_size=11)
    add_bei_header(doc)

    add_title(
        doc,
        'Receiving Pipeline — Finance Reconciliation',
        'Owner: Denise (effective 2026-04-28; Juanna\'s last day 2026-04-27)',
    )

    add_big_benefit_callout(
        doc,
        'Your role in this',
        [
            'The 3PL / supplier / Apps Script pipeline creates the DR and the SI '
            'match. Your job is to confirm that what lands in Frappe (GRs, RFPs, '
            'payment schedules) matches the source of truth in Sheet C, and to '
            'release payment on the contracted Net terms — nothing later.',
        ],
    )

    add_section_header(doc, 'Quick access')
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 40, 40)
    add_run(p, 'Sheet C (master): ', bold=True, size=11)
    add_hyperlink(p, SHEET_C_URL, 'BEI Receiving Master 2026', size=11)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 40)
    add_run(p, 'Frappe ERP: ', bold=True, size=11)
    add_hyperlink(p, 'https://hq.bebang.ph', 'hq.bebang.ph', size=11)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 120)
    add_run(p, 'Upload form (supplier-facing): ', bold=True, size=11)
    add_hyperlink(p, FORM_URL, 'BEI Supplier SI Upload', size=11)

    add_section_header(doc, 'Daily (30 min)')

    add_numbered_step(doc, 1, 'Confirm yesterday\'s pipeline closed',
                      'Sheet C 02_All_Receipts_Consolidated row count for yesterday vs Frappe '
                      'Purchase Receipts posted yesterday. Expected: Frappe count = Sheet count '
                      'minus variance rows. Discrepancy > 5% = ping Ashish. Don\'t hand-fix in '
                      'Frappe; fix upstream.')
    add_numbered_step(doc, 2, 'Verify RFPs match SI uploads',
                      'For each SI_Matched=TRUE row > 1 day old: confirm RFP exists in Frappe '
                      'for (PO, Supplier), amount matches SI upload, scheduled pay date = '
                      'DR date + Supplier.Payment_Terms. Mismatch → correct, note, ping Luwi.')
    add_numbered_step(doc, 3, 'Chase DRs older than 3 days with no SI match',
                      'Filter SI_Matched=FALSE AND Timestamp < today - 3d. Check 04_Match_Queue '
                      'for typos; email the supplier (cc Cayla) if no upload exists. Do NOT '
                      'release payment on a DR without a matching SI.')

    add_section_header(doc, 'Weekly (Monday 10:00)')
    add_body(doc, '1. Sheet C vs Frappe GL reconciliation — sum Amount on last week\'s SI_Matched=TRUE rows = GR entries on Stock in Hand.')
    add_body(doc, '2. EWT + VAT spot-check — pick 3 RFPs at random, confirm Frappe applied supplier\'s EWT Rate + VAT Registered status correctly.')
    add_body(doc, '3. Tier classification review — Sheet C 07_Full_Suppliers_Master. Tier blank or wrong? Rule of thumb: Tier A > ₱2M/year purchases.')

    add_section_header(doc, 'Monthly (first 3 business days)')
    add_body(doc, '1. BIR Form 2307 reconciliation — EWT withheld per Frappe = EWT on BIR Form 2307 issued. Discrepancies trace to manual RFP overrides.')
    add_body(doc, '2. Paper SI spot-check — 10 random rows from 03_Supplier_SI_Uploads. Drive link opens; PDF matches paper on file.')
    add_body(doc, '3. Stale Pending GR audit — 06_Pending_GR rows with PENDING status and Picked_Up_By_AppSheet=FALSE for > 7 days = Ashish AppSheet not polling. Escalate.')

    add_section_header(doc, 'Hard rules (will bite you otherwise)')
    add_zebra_table(
        doc,
        ['Rule', 'What it means'],
        [
            ['Payment releases on DR + SI upload',
             'Don\'t wait for paper SI to arrive before queuing payment. That\'s the old workflow.'],
            ['Disputes happen AFTER payment',
             'Short delivery or defective goods → pay the DR + SI, raise dispute afterward. Never withhold to force the issue.'],
            ['Approval chain while CFO seat vacant',
             'RFPs ≤ ₱1M: Luwi prepares → Mae approves. RFPs > ₱1M: Luwi → Mae → Sam. No CFO step.'],
            ['Never use Internal Customer for a supplier SI',
             'Internal Customers = inter-BEI-entity labor journals only. If a supplier invoice shows Internal Customer as the recipient, stop and ping Sam.'],
        ],
        col_widths=[2.3, 3.7],
        first_col_bold=True,
    )

    add_section_header(doc, 'Tools you need')
    add_body(doc, '• Frappe (hq.bebang.ph) — Finance role via your employee account')
    add_body(doc, '• Sheet C BEI Receiving Master — Cayla or Ian will share')
    add_body(doc, '• Procurement AppSheet — read access from Ashish')
    add_body(doc, '• BIR eFPS credentials — get from Juanna during handover on or before 2026-04-27')

    add_section_header(doc, 'Escalation')
    add_zebra_table(
        doc,
        ['Issue', 'Contact'],
        [
            ['Frappe GR missing for a receipt in Sheet C', 'Ashish'],
            ['RFP amount doesn\'t match supplier\'s upload', 'Luwi'],
            ['Payment terms look wrong', 'Cayla (supplier master owner)'],
            ['Automation broken', 'Sam + Ian'],
            ['BIR interpretation question', 'Sam (until proper CFO hired)'],
        ],
        col_widths=[3.5, 2.5],
        first_col_bold=True,
    )

    add_footer_note(doc, 'v2026-04-21 · Juanna last day 2026-04-27 · Denise effective 2026-04-28')

    path = GUIDES_DIR / '6_FINANCE_RECONCILIATION_GUIDE.docx'
    doc.save(str(path))
    return path, 'Receiving Pipeline — Finance Reconciliation (Denise)'


def build_master_index():
    doc = Document()
    setup_doc(doc, font_size=11)
    add_bei_header(doc)

    add_title(
        doc,
        'S210 Receipt-Based Payment Infrastructure',
        'Master Index — team-by-team guide to which sheets to use',
    )

    add_big_benefit_callout(
        doc,
        'Why this exists',
        [
            'Before S210, BEI waited for paper SIs to travel from 3PL warehouses '
            'to our accounting team before starting supplier payments. Delays '
            'were routine; supplier goodwill suffered.',
            '',
            'S210 replaces paper-SI-forwarding with a parallel DR + supplier-upload '
            'pipeline. Suppliers get paid on their contracted Net terms counted '
            'from delivery date. No paper-chase.',
        ],
    )

    add_section_header(doc, 'Live resources')

    add_zebra_table(
        doc,
        ['Resource', 'Who uses it', 'Link'],
        [
            ['Sheet A — 3MD Receiving Log', '3MD editor + BEI team', 'Sheet A (click below)'],
            ['Sheet B — Pinnacle Receiving Log', 'Pinnacle editor + BEI team', 'Sheet B (click below)'],
            ['Sheet C — BEI Receiving Master', 'BEI internal only', 'Sheet C (click below)'],
            ['Sheet D — Shaw Transitional', 'BEI internal only', 'Sheet D (click below)'],
            ['Supplier SI Upload Form', '98 suppliers', 'Form (click below)'],
            ['Apps Script project', 'Sam, IT', 'Script (click below)'],
            ['Cloud Scheduler console', 'Sam, Ian (read-only)', 'GCP (click below)'],
        ],
        col_widths=[2.4, 2.2, 1.4],
        first_col_bold=True,
    )

    # Per-resource clickable rows
    for label, url in [
        ('Sheet A — 3MD Receiving Log', SHEET_A_URL),
        ('Sheet B — Pinnacle Receiving Log', SHEET_B_URL),
        ('Sheet C — BEI Receiving Master', SHEET_C_URL),
        ('Sheet D — Shaw Transitional', SHEET_D_URL),
        ('Supplier SI Upload Form', FORM_URL),
        ('Apps Script project', SCRIPT_EDITOR_URL),
        ('Cloud Scheduler console', SCHEDULER_CONSOLE_URL),
    ]:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, 20, 40)
        add_run(p, f'• {label}: ', bold=True, size=10.5, color=PALETTE['text'])
        add_hyperlink(p, url, url, size=10)

    add_section_header(doc, 'Guides — by role')

    add_zebra_table(
        doc,
        ['Guide', 'For', 'What they get'],
        [
            ['Supplier SI Upload FAQ',
             '98 BEI suppliers',
             'How to upload + why it speeds payment. Attach to rollout emails.'],
            ['3PL Dock Quick Card',
             '3MD + Pinnacle dock staff',
             'One-page print / laminate. How to log a delivery in 30 seconds.'],
            ['Supplier Rollout Playbook',
             'Cayla',
             '4-wave rollout plan, email template, objection handling.'],
            ['Ian Daily Ops Playbook',
             'Ian',
             'Morning + afternoon + weekly checklists. KPI thresholds. Escalation.'],
            ['Finance Reconciliation',
             'Denise (2026-04-28+)',
             'Daily / weekly / monthly reconciliation. Approval chain. BIR rules.'],
        ],
        col_widths=[1.7, 1.6, 2.7],
        first_col_bold=True,
    )

    add_section_header(doc, 'Who uses which sheet')

    add_zebra_table(
        doc,
        ['Team', 'Writes to', 'Reads from'],
        [
            ['3MD dock staff', 'Sheet A Receipts tab', '—'],
            ['Pinnacle dock staff', 'Sheet B Receipts tab', '—'],
            ['Shaw dock (transitional)', 'Sheet D Receipts tab', '—'],
            ['Suppliers', 'Upload Form', '—'],
            ['Ian (daily ops)', 'Sheet C Match Queue + Variance Queue resolutions', 'All sheets'],
            ['Cayla (supplier master, rollout)', 'Procurement AppSheet Suppliers', 'Sheet C Dashboard'],
            ['Ashish (Frappe automation)', 'Frappe GRs + RFPs', 'Sheet C 06_Pending_GR'],
            ['Luwi (AP prep)', 'Frappe RFPs', 'Sheet C 03_Supplier_SI_Uploads'],
            ['Mae (approvals ≤ ₱1M)', 'Frappe RFP approval', 'Frappe pending RFPs'],
            ['Denise (reconciliation)', 'Frappe adjustments, BIR filings', 'Sheet C + Frappe'],
            ['Sam (CEO oversight)', '—', 'Daily 07:00 KPI email'],
        ],
        col_widths=[2.1, 2.1, 1.8],
        first_col_bold=True,
    )

    add_section_header(doc, 'Automation — what runs on its own')

    add_zebra_table(
        doc,
        ['Job', 'Schedule', 'What it does'],
        [
            ['s210-poll-all', 'every minute', 'Picks up new 3PL receipts + supplier uploads → Sheet C'],
            ['s210-age-variance-hourly', 'every hour', 'Moves DRs > 72h w/o SI match into Variance Queue'],
            ['s210-refresh-masters-06', 'daily 06:00 PHT', 'Pulls Suppliers + POs from Procurement AppSheet'],
            ['s210-ceo-email-07', 'daily 07:00 PHT', 'Sends KPI digest to Sam + Ian'],
        ],
        col_widths=[2.0, 1.5, 2.5],
        first_col_bold=True,
    )

    add_footer_note(doc, 'v2026-04-21 · Sprint S210 · Receipt-Based Payment Infrastructure')

    path = GUIDES_DIR / '1_MASTER_INDEX.docx'
    doc.save(str(path))
    return path, 'S210 Master Index — Team-by-Team Guide'


# ------------------------------------------------------------------
# Orchestrator
# ------------------------------------------------------------------

def main():
    print('\n=== S210 designed guide generator ===\n')
    outputs = []
    for name, builder in [
        ('MASTER_INDEX', build_master_index),
        ('3PL_DOCK_QUICK_CARD', build_dock_card),
        ('SUPPLIER_FAQ', build_supplier_faq),
        ('SUPPLIER_ROLLOUT_GUIDE', build_supplier_rollout),
        ('IAN_DAILY_OPS_PLAYBOOK', build_ian_playbook),
        ('FINANCE_RECONCILIATION_GUIDE', build_finance_guide),
    ]:
        path, title = builder()
        size_kb = path.stat().st_size / 1024
        print(f'  [{name}] {path.name}  ({size_kb:.0f} KB)  — "{title}"')
        outputs.append((name, path, title))

    print('\nAll 6 DOCX generated.')
    return outputs


if __name__ == '__main__':
    main()
