"""Generate BEI-branded DOCX versions of the 5 user-facing Phase 11 guides.

Source: output/s210/guides/*.md
Output: output/s210/guides/*.docx (alongside source)

SHEET_INSTRUCTIONS_3PL.md is intentionally SKIPPED — its content is pushed
directly into the live Sheet A + Sheet B `_Instructions` tabs and a DOCX
version would be unused.

Lightweight markdown parser handles: H1/H2/H3, pipe tables, bullets
(- or *), numbered lists, bold (**text**), italic (*text*), fenced code
blocks (```), blockquotes (>), horizontal rules (---).

Run:
    python output/s210/generate_guide_docx.py
"""
from __future__ import annotations

import re
import sys
import subprocess
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')
sys.path.insert(0,
    r'F:\Dropbox\Projects\BEI-ERP\.claude\skills\docx-designer-bei-erp\scripts')

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from bei_docx import (
    PALETTE, setup_doc, add_bei_header, add_section_header,
    add_run, add_body, set_cell_shading, set_cell_borders,
    set_cell_margins, set_paragraph_spacing, remove_cell_borders,
)

GUIDES_DIR = Path(r'F:\Dropbox\Projects\BEI-ERP-s210e\output\s210\guides')

# Doc title + subtitle derived from filename
DOC_META = {
    '3PL_DOCK_QUICK_CARD': (
        'BEI Receiving Quick Card',
        '3MD / Pinnacle dock staff — print + laminate',
    ),
    'SUPPLIER_FAQ': (
        'BEI Supplier SI Upload',
        'Frequently Asked Questions',
    ),
    'SUPPLIER_ROLLOUT_GUIDE': (
        'Supplier SI Upload — Rollout Playbook',
        'Owner: Cayla',
    ),
    'IAN_DAILY_OPS_PLAYBOOK': (
        'BEI Receiving — Daily Ops Playbook',
        'Owner: Ian Dionisio',
    ),
    'FINANCE_RECONCILIATION_GUIDE': (
        'BEI Receiving Pipeline — Finance Reconciliation Playbook',
        'Owner: Denise (effective 2026-04-28)',
    ),
}

# Files to convert (SHEET_INSTRUCTIONS_3PL deliberately excluded)
TO_CONVERT = list(DOC_META.keys())

# ------------------------------------------------------------------
# Lightweight markdown parser
# ------------------------------------------------------------------

_RE_BOLD = re.compile(r'\*\*(.+?)\*\*')
_RE_ITALIC = re.compile(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)')
_RE_CODE = re.compile(r'`([^`]+?)`')
_RE_LINK = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')


def _add_inline(p, text, base_size=11):
    """Render inline markdown (bold, italic, inline code, links) into runs."""
    # Process in passes. Simple tokenizer: find earliest match, emit run.
    remaining = text
    while remaining:
        # Find earliest of: bold, italic, code, link
        candidates = []
        for pattern, kind in [
            (_RE_BOLD, 'bold'),
            (_RE_CODE, 'code'),
            (_RE_LINK, 'link'),
            (_RE_ITALIC, 'italic'),
        ]:
            m = pattern.search(remaining)
            if m:
                candidates.append((m.start(), m, kind))
        if not candidates:
            add_run(p, remaining, size=base_size)
            return
        candidates.sort(key=lambda t: t[0])
        start, m, kind = candidates[0]
        # Preamble before the match
        if start > 0:
            add_run(p, remaining[:start], size=base_size)
        if kind == 'bold':
            add_run(p, m.group(1), bold=True, size=base_size)
        elif kind == 'italic':
            add_run(p, m.group(1), italic=True, size=base_size)
        elif kind == 'code':
            add_run(p, m.group(1), size=base_size - 1, font='Consolas',
                    color=PALETTE['primary'])
        elif kind == 'link':
            label, url = m.group(1), m.group(2)
            add_run(p, label, color=PALETTE['primary'], underline=True,
                    size=base_size)
        remaining = remaining[m.end():]


def _render_table(doc, header, rows):
    """Render a table with BEI-styled header."""
    if not header:
        return
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row: green background, white bold text
    for i, text in enumerate(header):
        c = table.cell(0, i)
        set_cell_shading(c, PALETTE['primary'])
        set_cell_borders(c, PALETTE['primary'], '6')
        set_cell_margins(c, 60, 60, 100, 100)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, text.strip(), bold=True, size=10.5,
                color=PALETTE['white'])
        set_paragraph_spacing(p, 0, 0)

    # Body rows with zebra striping
    for r_idx, row in enumerate(rows):
        bg = PALETTE['white'] if r_idx % 2 == 0 else PALETTE['bg_alt']
        for i in range(n_cols):
            c = table.cell(r_idx + 1, i)
            set_cell_shading(c, bg)
            set_cell_borders(c, PALETTE['border'], '4')
            set_cell_margins(c, 50, 50, 100, 100)
            text = row[i] if i < len(row) else ''
            p = c.paragraphs[0]
            _add_inline(p, text.strip(), base_size=10)
            set_paragraph_spacing(p, 0, 0)

    doc.add_paragraph()


def _render_quote(doc, lines):
    """Render a blockquote as a gold-bordered callout."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.columns[0].width = Inches(6.0)
    cell = box.cell(0, 0)
    set_cell_shading(cell, PALETTE['light_gold'])
    set_cell_borders(cell, PALETTE['secondary'], '8')
    set_cell_margins(cell, 100, 100, 180, 180)

    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        _add_inline(p, line, base_size=11)
        set_paragraph_spacing(p, 0, 40)
    doc.add_paragraph()


def _render_code_block(doc, lines):
    """Render a fenced code block as a gray-bg monospace box."""
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.columns[0].width = Inches(6.0)
    cell = box.cell(0, 0)
    set_cell_shading(cell, 'F4F2ED')
    set_cell_borders(cell, PALETTE['border'], '4')
    set_cell_margins(cell, 80, 80, 140, 140)

    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        add_run(p, line, size=9, font='Consolas', color=PALETTE['text'])
        set_paragraph_spacing(p, 0, 0)
    doc.add_paragraph()


def _render_bullet(doc, line, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3 + 0.3 * level)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    add_run(p, '• ', bold=True, size=11, color=PALETTE['primary'])
    _add_inline(p, line, base_size=11)
    set_paragraph_spacing(p, 0, 40)


def _render_numbered(doc, num, line):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    add_run(p, f'{num}. ', bold=True, size=11, color=PALETTE['primary'])
    _add_inline(p, line, base_size=11)
    set_paragraph_spacing(p, 0, 40)


# ------------------------------------------------------------------
# Main parser
# ------------------------------------------------------------------

def _strip_emojis(text: str) -> str:
    """Remove emojis the user doesn't want in generated docs."""
    import re
    # Common emojis we used (cross, check, warnings, traffic lights, etc.)
    # Range covers most emoji blocks; plus a few specific ones we used.
    emoji_pattern = re.compile(
        '['
        '\U0001F300-\U0001F9FF'  # misc symbols + pictographs
        '\U0001F600-\U0001F64F'  # emoticons
        '\U0001F680-\U0001F6FF'  # transport
        '\u2600-\u27BF'          # misc symbols incl. ✗✅❌⚠
        '\u2700-\u27BF'          # dingbats
        ']+',
        flags=re.UNICODE,
    )
    cleaned = emoji_pattern.sub('', text)
    # Clean up trailing whitespace from lines that had only emoji+space
    cleaned = re.sub(r'^[\s]*-[\s]+-[\s]+', '- ', cleaned, flags=re.MULTILINE)
    return cleaned


def _join_multiline_italic(text: str) -> str:
    """Join '_...[newline]..._' footer italics into a single line."""
    import re
    # Pattern: _ on its own, content possibly spanning lines, _ on its own
    # Match _<content>_ where content may include newlines; collapse to one line
    def _collapse(m):
        inner = m.group(1).replace('\n', ' ').replace('  ', ' ')
        return '_' + inner + '_'
    # Only match if the pair is on its own (preceded by ^ or newline, followed by $ or newline)
    return re.sub(r'_([^_]+?)_', _collapse, text, flags=re.DOTALL)


def md_to_docx(md_path: Path, docx_path: Path, title: str, subtitle: str):
    """Parse markdown and render as a BEI-branded DOCX."""
    text = md_path.read_text(encoding='utf-8')
    text = _strip_emojis(text)
    text = _join_multiline_italic(text)
    lines = text.splitlines()

    doc = Document()
    setup_doc(doc)
    add_bei_header(doc)

    # Document title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, title, bold=True, size=18, color=PALETTE['primary'])
    set_paragraph_spacing(p, 0, 40)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, subtitle, italic=True, size=11, color=PALETTE['text_light'])
    set_paragraph_spacing(p, 0, 160)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Skip the markdown H1 (title already rendered)
        if line.startswith('# ') and i < 3:
            i += 1
            continue

        # Horizontal rule
        if stripped in ('---', '***', '___'):
            # Render as a subtle gold line
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            hr = doc.add_table(rows=1, cols=1)
            hr.alignment = WD_TABLE_ALIGNMENT.CENTER
            hr.columns[0].width = Inches(6.0)
            c = hr.cell(0, 0)
            set_cell_shading(c, PALETTE['secondary'])
            set_cell_margins(c, 0, 0, 0, 0)
            hp = c.paragraphs[0]
            add_run(hp, ' ', size=1)
            set_paragraph_spacing(hp, 0, 0)
            doc.add_paragraph()
            i += 1
            continue

        # Headings
        if line.startswith('## '):
            add_section_header(doc, line[3:].strip())
            i += 1
            continue
        if line.startswith('### '):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, 120, 40)
            add_run(p, line[4:].strip(), bold=True, size=12,
                    color=PALETTE['secondary'])
            i += 1
            continue
        if line.startswith('#### '):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, 80, 20)
            add_run(p, line[5:].strip(), bold=True, size=11,
                    color=PALETTE['text'])
            i += 1
            continue
        if line.startswith('# '):
            add_section_header(doc, line[2:].strip())
            i += 1
            continue

        # Fenced code block
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            _render_code_block(doc, code_lines)
            continue

        # Blockquote
        if stripped.startswith('>'):
            quote_lines = []
            while i < n and lines[i].strip().startswith('>'):
                quote_lines.append(lines[i].strip().lstrip('>').strip())
                i += 1
            _render_quote(doc, quote_lines)
            continue

        # Pipe table
        if stripped.startswith('|') and i + 1 < n and re.match(r'^\s*\|[\s:|\-]+\|\s*$', lines[i + 1]):
            header_cells = [c.strip() for c in stripped.strip('|').split('|')]
            i += 2  # skip header + separator
            body = []
            while i < n and lines[i].strip().startswith('|'):
                row_cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                body.append(row_cells)
                i += 1
            _render_table(doc, header_cells, body)
            continue

        # Numbered list
        m_num = re.match(r'^(\d+)\.\s+(.+)$', line)
        if m_num:
            _render_numbered(doc, m_num.group(1), m_num.group(2))
            i += 1
            # Handle continuation lines (indented) — fold in
            while i < n and (lines[i].startswith('   ') or lines[i].startswith('\t')) and lines[i].strip():
                # Append continuation to the previous para (keep simple: new para indented)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.6)
                _add_inline(p, lines[i].strip(), base_size=11)
                set_paragraph_spacing(p, 0, 20)
                i += 1
            continue

        # Bullet list
        m_bul = re.match(r'^[-*]\s+(.+)$', line)
        if m_bul:
            _render_bullet(doc, m_bul.group(1), level=0)
            i += 1
            # Continuation lines
            while i < n:
                m_nested = re.match(r'^\s{2,}[-*]\s+(.+)$', lines[i])
                if m_nested:
                    _render_bullet(doc, m_nested.group(1), level=1)
                    i += 1
                    continue
                if (lines[i].startswith('   ') or lines[i].startswith('\t')) and lines[i].strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.6)
                    _add_inline(p, lines[i].strip(), base_size=11)
                    set_paragraph_spacing(p, 0, 20)
                    i += 1
                    continue
                break
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Italic trailing metadata (e.g. "_Version 2026-04-21._")
        if stripped.startswith('_') and stripped.endswith('_') and len(stripped) > 2:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, 80, 0)
            add_run(p, stripped.strip('_').strip(), italic=True, size=9,
                    color=PALETTE['text_light'])
            i += 1
            continue

        # Regular paragraph (may span multiple lines until blank)
        para_lines = [line]
        i += 1
        while i < n and lines[i].strip() and not (
            lines[i].startswith('#') or
            lines[i].strip().startswith('|') or
            lines[i].strip().startswith('```') or
            lines[i].strip().startswith('>') or
            lines[i].strip() in ('---', '***', '___') or
            re.match(r'^[-*]\s+', lines[i]) or
            re.match(r'^\d+\.\s+', lines[i])
        ):
            para_lines.append(lines[i])
            i += 1
        p = doc.add_paragraph()
        _add_inline(p, ' '.join(line.strip() for line in para_lines), base_size=11)
        set_paragraph_spacing(p, 40, 60)

    doc.save(str(docx_path))
    print(f'  Wrote {docx_path.name} ({docx_path.stat().st_size:,} bytes)')


def main():
    out_dir = GUIDES_DIR
    print(f'\nGenerating DOCX guides in {out_dir}\n')
    for stem in TO_CONVERT:
        md = out_dir / f'{stem}.md'
        docx = out_dir / f'{stem}.docx'
        title, subtitle = DOC_META[stem]
        print(f'[{stem}]')
        md_to_docx(md, docx, title, subtitle)
    print('\nAll 5 DOCX files generated.')


if __name__ == '__main__':
    main()
