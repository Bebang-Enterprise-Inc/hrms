"""S258 Phase 6 redo — Build SIGNOFF.docx with BEI brand theme.

Bridge Consulting QBO Handoff Package — formal sign-off document.
"""
from __future__ import annotations
import sys
from pathlib import Path

sys.path.insert(0, r"F:\Dropbox\Projects\BEI-ERP\.claude\skills\docx-designer-bei-erp\scripts")
from bei_docx import (
    PALETTE, LOGO_PATH, setup_doc, add_bei_header, add_section_header,
    add_run, add_body, add_page_break, add_signature_block,
    add_callout_box, set_paragraph_spacing, set_cell_shading,
    set_cell_borders, set_cell_margins, hex_rgb,
)

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor


def add_cover(doc):
    """Cover page: BEI letterhead band + title + metadata block."""
    # Top green band
    band = doc.add_paragraph()
    band.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(band, before=0, after=0)
    pPr = band._element.get_or_add_pPr()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), PALETTE['primary'])
    pPr.append(shd)
    add_run(band, "BEBANG ENTERPRISE INC.", bold=True, size=18, color=PALETTE['white'])

    # Sub-band gold
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(sub, before=0, after=0)
    pPr = sub._element.get_or_add_pPr()
    shd2 = OxmlElement('w:shd')
    shd2.set(qn('w:val'), 'clear')
    shd2.set(qn('w:color'), 'auto')
    shd2.set(qn('w:fill'), PALETTE['light_gold'])
    pPr.append(shd2)
    add_run(sub, "Chart of Accounts + General Ledger Finalization", bold=True, size=12, color=PALETTE['text'])

    # Spacer
    sp = doc.add_paragraph()
    set_paragraph_spacing(sp, before=240, after=0)

    # Logo (centered)
    if LOGO_PATH.exists():
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = logo_p.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(2.5))
        set_paragraph_spacing(logo_p, before=200, after=200)

    # Title
    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(t1, before=400, after=80)
    add_run(t1, "S258 SIGN-OFF", bold=True, size=28, color=PALETTE['primary'])

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(t2, before=0, after=40)
    add_run(t2, "Bridge Consulting QBO Handoff Package", bold=True, size=14, color=PALETTE['secondary'])

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(t3, before=0, after=400)
    add_run(t3, "COA + GL Finalization for ALL 58 Companies", italic=True, size=12, color=PALETTE['text_light'])

    # Metadata table
    meta = doc.add_table(rows=6, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_rows = [
        ("Sprint ID", "S258"),
        ("Sprint Title", "COA + GL Finalization for ALL 58 Companies (Bridge QBO Handoff)"),
        ("Date", "2026-06-04"),
        ("Approver", "Sam Karazi — Chief Executive Officer, Bebang Enterprise Inc."),
        ("Migration Partner", "Bridge Consulting (APEX → QuickBooks Online)"),
        ("PR / Merge SHA", "Bebang-Enterprise-Inc/hrms#770 — merged at 8b0636f22"),
    ]
    for i, (k, v) in enumerate(meta_rows):
        meta.cell(i, 0).text = ""
        meta.cell(i, 1).text = ""
        kp = meta.cell(i, 0).paragraphs[0]
        add_run(kp, k, bold=True, color=PALETTE['white'], size=10)
        set_cell_shading(meta.cell(i, 0), PALETTE['primary'])
        vp = meta.cell(i, 1).paragraphs[0]
        add_run(vp, v, size=10, color=PALETTE['text'])
        set_cell_shading(meta.cell(i, 1), PALETTE['white'])
        set_cell_borders(meta.cell(i, 0), color=PALETTE['border'])
        set_cell_borders(meta.cell(i, 1), color=PALETTE['border'])
        set_cell_margins(meta.cell(i, 0), top=80, bottom=80, left=120, right=120)
        set_cell_margins(meta.cell(i, 1), top=80, bottom=80, left=120, right=120)
    meta.columns[0].width = Inches(2.0)
    meta.columns[1].width = Inches(4.5)

    add_page_break(doc)


def add_kpi_dashboard(doc):
    add_section_header(doc, "EXECUTIVE SUMMARY — KEY METRICS")
    p = add_body(doc,
        "S258 closed out on 2026-06-04 with all 100 work units executed across 11 phases. "
        "All 58 Frappe Companies now carry the canonical 5-root tree + Butch Formoso's "
        "27-account Sales tree (per-role population) + Fork 1 scaffolding for BFC + per-store "
        "P&L for the 4 BEI-TIN stub stores. The Bridge Consulting QBO handoff package is "
        "ready for sandbox import.",
        first_indent=True,
    )

    # KPI Grid (4 cards in 1 row)
    kpi = doc.add_table(rows=2, cols=4)
    kpi.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpis = [
        ("58", "COMPANIES"),
        ("6,928", "ACTIVE ACCOUNTS"),
        ("290", "ROOT GROUPS"),
        ("100%", "GL PRESERVED"),
    ]
    for i, (val, label) in enumerate(kpis):
        # Top row: big number
        c = kpi.cell(0, i)
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, val, bold=True, size=28, color=PALETTE['accent'])
        set_cell_shading(c, PALETTE['light_purple'])
        set_cell_borders(c, color=PALETTE['border'])
        set_cell_margins(c, top=200, bottom=80, left=80, right=80)
        # Bottom row: label
        c2 = kpi.cell(1, i)
        c2.text = ""
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p2, label, bold=True, size=9, color=PALETTE['text'])
        set_cell_shading(c2, PALETTE['light_purple'])
        set_cell_borders(c2, color=PALETTE['border'])
        set_cell_margins(c2, top=40, bottom=200, left=80, right=80)

    doc.add_paragraph()


def add_directive_compliance(doc):
    add_section_header(doc, "SAM DIRECTIVE COMPLIANCE — CEO Directive 2026-06-04")
    add_callout_box(doc, [
        ("Verbatim CEO Directive — 2026-06-04", True, 11),
        ('"GLs and COA for all companies ready for Bridge. Per-store P&L. '
         'BKI = Commissary. BEI = Head Office. Nothing cancelled or deferred."', False, 10),
    ])

    rows = [
        ("Directive", "Live State", "Verdict"),
        ("GLs and COA for all companies ready for Bridge",
         "58 CSVs in per_company_coa.zip; 6,928 active accounts", "YES"),
        ("Per-store P&L for ALL companies",
         "49 stores + 4 BEI-TIN stubs each with own Sales tree", "YES"),
        ("BKI P&L = Commissary",
         "BKI populates 4000200 sub-tree only per COA-175-011", "YES"),
        ("BEI = Head Office P&L (NOT store)",
         "BEI Head Office; 4 stubs are separate Frappe Companies", "YES"),
        ("BFC GO-LIVE structure ready",
         "Fork 1 scaffolding (1104200 + 2104200 + 2102205)", "YES"),
        ("BFI2 → BFT abbr rename (SEC name unchanged)",
         "abbr=BFT; SEC TIN 663-440-106-00000 preserved", "YES"),
        ("Nothing cancelled or deferred",
         "10 in-session deviations (D0-* + D1-*) all RESOLVED", "YES"),
    ]
    t = doc.add_table(rows=len(rows), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j)
            c.text = ""
            p = c.paragraphs[0]
            if i == 0:
                add_run(p, val, bold=True, color=PALETTE['white'], size=10)
                set_cell_shading(c, PALETTE['primary'])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                add_run(p, val, size=9, color=PALETTE['text'])
                set_cell_shading(c, PALETTE['white'] if i % 2 else PALETTE['light'])
                if j == 2:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_cell_shading(c, '8FBC8F')  # success green tint
                    p.runs[0].font.bold = True
            set_cell_borders(c, color=PALETTE['border'])
            set_cell_margins(c, top=60, bottom=60, left=100, right=100)
    t.columns[0].width = Inches(2.6)
    t.columns[1].width = Inches(3.6)
    t.columns[2].width = Inches(0.8)


def add_phase_results(doc):
    add_section_header(doc, "PHASE RESULTS — All 11 phases PASS")
    rows = [
        ("Phase", "Description", "Result"),
        ("0", "Boot + Preflight + Audit + DECISIONS.md ratification", "PASS"),
        ("1", "Safe Sync — A1+A2+A3+A4+A5", "PASS"),
        ("2", "Templates + B1 BFC + B2 BFT + B3 4 BEI-TIN stubs", "PASS"),
        ("2.0", "Per-Company Apex→canonical migration map (BEI/BKI/III)", "PASS"),
        ("3a", "5-root tree seed on all 58 Companies (SSM bench)", "PASS"),
        ("3b", "BKI Commissary rewrite (combined into C2 mass-normalize)", "PASS"),
        ("3c", "BEI Head Office rewrite (combined into C2)", "PASS"),
        ("3.5", "BEI AP/AR suffix → abbr (286 long-suffix renames)", "PASS"),
        ("4", "4000900 DISCOUNTS AND PROMO renumber", "PASS"),
        ("5", "UPPER CASE + drop number-prefix (157 renames)", "PASS"),
        ("6", "Bridge QBO handoff package (58 CSVs + manifests)", "PASS"),
        ("7", "Closeout (DECISIONS.md + plan COMPLETED + PR)", "PASS"),
    ]
    t = doc.add_table(rows=len(rows), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j)
            c.text = ""
            p = c.paragraphs[0]
            if i == 0:
                add_run(p, val, bold=True, color=PALETTE['white'], size=10)
                set_cell_shading(c, PALETTE['primary'])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                add_run(p, val, size=9, color=PALETTE['text'])
                set_cell_shading(c, PALETTE['white'] if i % 2 else PALETTE['light'])
                if j in (0, 2):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if j == 2:
                    set_cell_shading(c, '8FBC8F')
                    p.runs[0].font.bold = True
            set_cell_borders(c, color=PALETTE['border'])
            set_cell_margins(c, top=50, bottom=50, left=80, right=80)
    t.columns[0].width = Inches(0.7)
    t.columns[1].width = Inches(5.5)
    t.columns[2].width = Inches(0.8)


def add_deltas(doc):
    add_section_header(doc, "KNOWN COSMETIC DELTAS")
    add_body(doc,
        "Five non-blocking cosmetic items remain in Frappe state. Bridge can import "
        "the 58 CSVs as-is and address these during QBO sandbox iteration or in a "
        "follow-up S259 cleanup sprint if stricter classification is desired.",
        first_indent=True)

    rows = [
        ("#", "Delta", "Affected", "Severity", "Recommendation"),
        ("1", "III still carries both old (4000201-206) AND new (4000901-903) discount account_numbers.",
         "III only", "LOW", "Keep both or delete 4000201-206 on III in sandbox."),
        ("2", "BEI `STOCK ADJUSTMENT - BEI` sits under `LOCAL AD & PROMO - BEI` Expense group instead of `Direct Expenses - BEI`.",
         "BEI only", "LOW", "Reparent in QBO post-import; no operational impact."),
        ("3", "Some BEI / III accounts retained number prefix in docname (e.g. `4000100 - STORE SALES - BEI`).",
         "BEI ~283, III ~36", "COSMETIC", "QBO uses AccountNumber + AccountName columns separately. Zero impact."),
        ("4", "BFC Fork 1 scaffolding docnames retain number prefix (`1104200 - DUE FROM BEI - BFC`).",
         "BFC, BEI", "COSMETIC", "Same as #3."),
        ("5", "QBO DetailType is best-effort for unmapped Frappe account_types.",
         "All Companies", "KNOWN", "QBO sandbox import surfaces rejected DetailType; iterate. Map handles 18 explicit + 5 fallbacks."),
    ]
    sev_colors = {"LOW": "CDD8CE", "COSMETIC": PALETTE['light_gold'], "KNOWN": PALETTE['light_purple']}
    t = doc.add_table(rows=len(rows), cols=5)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j)
            c.text = ""
            p = c.paragraphs[0]
            if i == 0:
                add_run(p, val, bold=True, color=PALETTE['white'], size=10)
                set_cell_shading(c, PALETTE['primary'])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                add_run(p, val, size=9, color=PALETTE['text'])
                set_cell_shading(c, PALETTE['white'] if i % 2 else PALETTE['light'])
                if j in (0, 3):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if j == 3:
                    set_cell_shading(c, sev_colors.get(val, PALETTE['white']))
                    p.runs[0].font.bold = True
            set_cell_borders(c, color=PALETTE['border'])
            set_cell_margins(c, top=50, bottom=50, left=80, right=80)
    t.columns[0].width = Inches(0.4)
    t.columns[1].width = Inches(2.9)
    t.columns[2].width = Inches(1.2)
    t.columns[3].width = Inches(0.8)
    t.columns[4].width = Inches(1.7)


def add_signatures(doc):
    add_section_header(doc, "AUTHORIZATION & ACKNOWLEDGMENT")
    add_body(doc,
        "By signing below, the BEI CEO releases the S258 deliverable (Bridge Consulting "
        "QBO Handoff Package) for use in the QuickBooks Online migration. Bridge "
        "Consulting acknowledges receipt and assumes responsibility for QBO sandbox import "
        "and any DetailType iteration required during testing.")

    # Two signature blocks side-by-side
    sigt = doc.add_table(rows=1, cols=2)
    sigt.alignment = WD_TABLE_ALIGNMENT.CENTER

    # BEI CEO sign-off (left)
    left = sigt.cell(0, 0)
    left.text = ""
    p1 = left.paragraphs[0]
    add_run(p1, "RELEASED BY (BEI):", bold=True, size=10, color=PALETTE['primary'])
    left.add_paragraph()
    left.add_paragraph()
    line1 = left.add_paragraph()
    add_run(line1, "_____________________________________", size=10)
    n1 = left.add_paragraph()
    add_run(n1, "Sam Karazi", bold=True, size=11, color=PALETTE['text'])
    t1 = left.add_paragraph()
    add_run(t1, "Chief Executive Officer", size=9, color=PALETTE['text_light'])
    cn1 = left.add_paragraph()
    add_run(cn1, "Bebang Enterprise Inc.", italic=True, size=9, color=PALETTE['text_light'])
    d1 = left.add_paragraph()
    add_run(d1, "Date: ______________________", size=10, color=PALETTE['text'])
    set_cell_shading(left, PALETTE['light'])
    set_cell_borders(left, color=PALETTE['secondary'], size="12")
    set_cell_margins(left, top=200, bottom=200, left=200, right=200)

    # Bridge acknowledgment (right)
    right = sigt.cell(0, 1)
    right.text = ""
    p2 = right.paragraphs[0]
    add_run(p2, "ACKNOWLEDGED BY (BRIDGE):", bold=True, size=10, color=PALETTE['primary'])
    right.add_paragraph()
    right.add_paragraph()
    line2 = right.add_paragraph()
    add_run(line2, "_____________________________________", size=10)
    n2 = right.add_paragraph()
    add_run(n2, "_______________________________", bold=True, size=11, color=PALETTE['text'])
    t2 = right.add_paragraph()
    add_run(t2, "(Print Name & Title)", size=9, color=PALETTE['text_light'])
    cn2 = right.add_paragraph()
    add_run(cn2, "Bridge Consulting", italic=True, size=9, color=PALETTE['text_light'])
    d2 = right.add_paragraph()
    add_run(d2, "Date: ______________________", size=10, color=PALETTE['text'])
    set_cell_shading(right, PALETTE['light'])
    set_cell_borders(right, color=PALETTE['secondary'], size="12")
    set_cell_margins(right, top=200, bottom=200, left=200, right=200)

    sigt.columns[0].width = Inches(3.5)
    sigt.columns[1].width = Inches(3.5)


def add_footer_block(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=400, after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "— END OF SIGN-OFF —", italic=True, size=9, color=PALETTE['text_light'])

    p2 = doc.add_paragraph()
    set_paragraph_spacing(p2, before=40, after=0)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "Bebang Enterprise Inc. — Confidential — For Bridge Consulting use only",
            size=8, color=PALETTE['text_light'])

    p3 = doc.add_paragraph()
    set_paragraph_spacing(p3, before=20, after=0)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p3, "Generated 2026-06-04 from S258 sprint closeout — see hrms PR #770",
            size=8, color=PALETTE['text_light'])


def main():
    sys.stdout.reconfigure(encoding="utf-8")
    doc = Document()
    setup_doc(doc)
    add_cover(doc)
    add_kpi_dashboard(doc)
    add_directive_compliance(doc)
    add_phase_results(doc)
    add_deltas(doc)
    add_signatures(doc)
    add_footer_block(doc)

    out = Path("output/s258/bridge_handoff/SIGNOFF.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"[OK] Wrote {out}")


if __name__ == "__main__":
    main()
