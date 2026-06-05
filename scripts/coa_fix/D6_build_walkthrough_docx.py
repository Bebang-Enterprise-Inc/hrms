"""S258 Bridge handoff — One-document walkthrough DOCX.

Comprehensive step-by-step guide for Bridge Consulting to set up QBO mirroring
Frappe. 13 sections, all embedded links, BEI-branded.
"""
from __future__ import annotations
import sys
from pathlib import Path

sys.path.insert(0, r"F:\Dropbox\Projects\BEI-ERP\.claude\skills\docx-designer-bei-erp\scripts")
from bei_docx import (
    PALETTE, LOGO_PATH, setup_doc, add_section_header,
    add_run, add_body, add_page_break, add_callout_box,
    set_paragraph_spacing, set_cell_shading, set_cell_borders, set_cell_margins,
)

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Hyperlink helper (BEI Green)
# ---------------------------------------------------------------------------
def add_hyperlink(paragraph, url, text, color=None, size=10, bold=False):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color or PALETTE['primary'])
    rPr.append(c)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size * 2))
    rPr.append(sz)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def shaded_paragraph(doc, color, content_runs, align='center'):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0)
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)
    for text, kwargs in content_runs:
        add_run(p, text, **kwargs)
    return p


def add_step(doc, num, title, body_text=None):
    """Add a numbered step block — green number badge + title + body."""
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    badge = t.cell(0, 0)
    badge.text = ""
    p = badge.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, str(num), bold=True, size=16, color=PALETTE['white'])
    set_cell_shading(badge, PALETTE['primary'])
    set_cell_borders(badge, color=PALETTE['primary'])
    set_cell_margins(badge, top=120, bottom=120, left=60, right=60)
    badge.width = Inches(0.6)

    title_cell = t.cell(0, 1)
    title_cell.text = ""
    p2 = title_cell.paragraphs[0]
    add_run(p2, title, bold=True, size=12, color=PALETTE['primary'])
    set_cell_shading(title_cell, PALETTE['light'])
    set_cell_borders(title_cell, color=PALETTE['secondary'])
    set_cell_margins(title_cell, top=120, bottom=120, left=200, right=120)
    title_cell.width = Inches(6.0)
    t.columns[0].width = Inches(0.6)
    t.columns[1].width = Inches(6.0)

    if body_text:
        for para in body_text.split('\n\n'):
            add_body(doc, para.strip())


def add_link_para(doc, label, url):
    """Add a label + clickable Drive URL on one line."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=20, after=20)
    add_run(p, f"  {label}: ", bold=True, size=10, color=PALETTE['text'])
    add_hyperlink(p, url, url, size=9)


def build_cover(doc):
    # Top green band
    shaded_paragraph(doc, PALETTE['primary'], [
        ("BEBANG ENTERPRISE INC.", {'bold': True, 'size': 18, 'color': PALETTE['white']}),
    ])
    shaded_paragraph(doc, PALETTE['light_gold'], [
        ("BEI ERP → QuickBooks Online Setup", {'bold': True, 'size': 13, 'color': PALETTE['text']}),
    ])
    shaded_paragraph(doc, PALETTE['light'], [
        ("Bridge Consulting Walkthrough Pack", {'italic': True, 'size': 11, 'color': PALETTE['text']}),
    ])

    if LOGO_PATH.exists():
        sp = doc.add_paragraph()
        set_paragraph_spacing(sp, before=400, after=200)
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sp.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(2.0))

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(t, before=200, after=40)
    add_run(t, "ONE-DOCUMENT WALKTHROUGH", bold=True, size=24, color=PALETTE['primary'])

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(sub, before=0, after=40)
    add_run(sub,
            "Everything Bridge needs to set up QuickBooks Online to mirror BEI's Frappe ERP",
            italic=True, size=11, color=PALETTE['text_light'])

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(sub2, before=0, after=400)
    add_run(sub2,
            "13 sections | All Drive links embedded | Step-by-step QBO setup checklist",
            size=10, color=PALETTE['text_light'])

    # Metadata table
    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_rows = [
        ("Audience", "Bridge Consulting (Anna R., Kim C., Flor A., accountant.outsource@bridge-ph.com)"),
        ("Author", "Sam Karazi — Chief Executive Officer, Bebang Enterprise Inc."),
        ("Date", "2026-06-05"),
        ("Source sprint", "S258 — COA + GL Finalization (PR #770, merged 8b0636f22)"),
        ("Important", "Frappe ERP is structurally ready but NOT yet operationally live. Bridge migrates from APEX (current live system) → QBO using the structure defined here as the target."),
    ]
    for i, (k, v) in enumerate(meta_rows):
        meta.cell(i, 0).text = ""
        meta.cell(i, 1).text = ""
        kp = meta.cell(i, 0).paragraphs[0]
        add_run(kp, k, bold=True, color=PALETTE['white'], size=10)
        set_cell_shading(meta.cell(i, 0), PALETTE['primary'])
        vp = meta.cell(i, 1).paragraphs[0]
        add_run(vp, v, size=10, color=PALETTE['text'])
        set_cell_shading(meta.cell(i, 1), PALETTE['white'])
        for cell in (meta.cell(i, 0), meta.cell(i, 1)):
            set_cell_borders(cell, color=PALETTE['border'])
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
    meta.columns[0].width = Inches(1.6)
    meta.columns[1].width = Inches(5.0)
    add_page_break(doc)


def build_toc(doc):
    add_section_header(doc, "TABLE OF CONTENTS")
    toc = [
        ("1.  Read This First", "Why this document exists + how to use it"),
        ("2.  The Company Universe", "Why BEI has 58 Frappe Companies + role tagging"),
        ("3.  The Canonical Chart of Accounts", "5-root tree + Butch's 27-account Sales tree + per-Company population"),
        ("4.  Tax & Regulatory Setup", "VAT 12% + EWT + BIR forms + currency + fiscal year"),
        ("5.  Intercompany Scheme", "Fork 1 (BFC) + BKI→store paired SI/PI + S206 labor cost-sharing"),
        ("6.  APEX → Canonical Translation", "How to translate old APEX data using migration maps"),
        ("7.  5 Known Cosmetic Deltas", "Non-blocker items to address during sandbox import"),
        ("8.  Master Decision Ledger", "Butch's locks + Sam's directives — full DECISIONS.md"),
        ("9.  What Bridge Needs from APEX (not BEI)", "Opening balances, AR/AP, bank balances"),
        ("10. QBO Setup Checklist", "Step-by-step ordered checklist for sandbox + production"),
        ("11. Reference File Index", "All Drive links in one place"),
        ("12. Contact Directory", "Bridge + BEI stakeholders + escalation"),
        ("13. Sign-off", "CEO release + Bridge acknowledgment"),
    ]
    for title, desc in toc:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=20, after=20)
        add_run(p, title, bold=True, size=11, color=PALETTE['primary'])
        add_run(p, f"  —  {desc}", size=10, color=PALETTE['text'])
    add_page_break(doc)


def build_section_1(doc):
    add_section_header(doc, "1.  READ THIS FIRST")
    add_callout_box(doc, [
        ("WHAT THIS DOCUMENT IS", True, 11),
        ("A single navigable walkthrough for Bridge Consulting to set up QuickBooks "
         "Online so that QBO mirrors BEI's Frappe ERP structure. Every supporting file "
         "is linked in the Reference File Index (Section 11) and uploaded to the same "
         "Drive folder where Bridge already has Editor access.", False, 10),
    ])

    add_body(doc, "Three facts to anchor everything below:", first_indent=False)
    facts = [
        ("Frappe is the TARGET ERP, not the current live system.",
         "Bebang Enterprise's accounting was on Apex Solutions. Apex is being retired. "
         "Bridge is migrating Apex transactions into QuickBooks Online (QBO) AND BEI is "
         "transitioning its operating ERP from Apex to Frappe simultaneously. The COA we "
         "ship from Frappe defines the TARGET shape; Apex provides the historical data."),
        ("BEI runs 58 Frappe Companies, not 1.",
         "Per Sam's directive 2026-06-04, every store gets its own per-store P&L. The 58 "
         "Companies represent 9 group/holdco entities + 49 store-level Companies (including "
         "4 BEI-TIN stub stores filed under BEI's TIN for management reporting). Bridge "
         "should create the corresponding 58 QBO Companies."),
        ("Butch Formoso (former CFO) locked the COA design in April 2026.",
         "All canonical decisions are in DECISIONS.md as COA-175-001..030. Butch resigned "
         "2026-04-15 but his signed locks remain BINDING per the 2026-05-07 supersession "
         "banner. Sam (CEO) ratifies all post-resignation supersessions."),
    ]
    for title, body in facts:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=40, after=20)
        add_run(p, "  •  ", bold=True, size=11, color=PALETTE['primary'])
        add_run(p, title, bold=True, size=11, color=PALETTE['text'])
        add_body(doc, body, first_indent=False)

    add_callout_box(doc, [
        ("HOW TO USE THIS DOCUMENT", True, 11),
        ("Sections 2-9 explain the WHAT and WHY. Section 10 is the operational "
         "step-by-step CHECKLIST for QBO setup — work through it in order. Every "
         "reference file linked is in your BEI COA Handoff Drive folder.", False, 10),
    ])

    add_section_header(doc, "Quick Numbers")
    qn_table = doc.add_table(rows=2, cols=5)
    qn_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (val, label) in enumerate([("58", "Companies"), ("6,928", "Active Accounts"),
                                       ("290", "Root Groups"), ("100%", "GL Preserved"),
                                       ("27", "COA-175 Locks")]):
        top = qn_table.cell(0, i)
        bot = qn_table.cell(1, i)
        top.text = ""
        bot.text = ""
        tp = top.paragraphs[0]
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(tp, val, bold=True, size=20, color=PALETTE['accent'])
        bp = bot.paragraphs[0]
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(bp, label, bold=True, size=9, color=PALETTE['text'])
        for c in (top, bot):
            set_cell_shading(c, PALETTE['light_purple'])
            set_cell_borders(c, color=PALETTE['border'])
            set_cell_margins(c, top=120, bottom=60, left=40, right=40)
    add_page_break(doc)


def build_section_2(doc):
    add_section_header(doc, "2.  THE COMPANY UNIVERSE — 58 Frappe Companies")
    add_body(doc,
        "BEI's accounting structure is a per-store P&L model. Each of the 49 physical stores "
        "is its own Frappe Company (with its own Chart of Accounts, Sales tree, and rolled-up "
        "P&L). Group entities (legal-entity parents, holdcos, the franchisor, the commissary) "
        "round out the 58. Bridge should create the same 58 in QBO — one QuickBooks Online "
        "Company per Frappe Company.")

    add_callout_box(doc, [
        ("REFERENCE FILE — Live Company Register", True, 11),
        ("COMPANY_REGISTER.xlsx in the Drive folder: 4 sheets including all 58 Companies "
         "with role / ABBR / parent / TIN / RDO / VAT status / account counts, plus the "
         "Per-Role P&L Population Rules and Butch's verbatim 27-account Sales tree.", False, 10),
    ])

    add_section_header(doc, "Role distribution")
    rows = [
        ("Role", "Count", "Frappe Company examples"),
        ("Holdco (top of tree)", "1", "IRRESISTIBLE INFUSIONS INC. (III) — is_group=1, parent of all"),
        ("Head Office (BEI HQ)", "1", "BEBANG ENTERPRISE INC. (BEI) — JV fees + Brand Growth + corporate overhead"),
        ("Commissary", "1", "BEBANG KITCHEN INC. (BKI) — Manufacturing, BKI→store deliveries"),
        ("Franchisor (Fork 1)", "1", "BEBANG FRANCHISE CORP. (BFC) — Royalty / Mgmt / Franchise / Marketing / E-Comm fees"),
        ("Legal Entity Parents (group=1)", "~6", "BEBANG MEGA INC. (BMI2), TAJ FOOD CORP., TUNGSTEN CAPITAL HOLDINGS OPC, etc."),
        ("Per-store P&L Companies (49)", "49", "AYALA VERMOSA - BEBANG MEGA INC., SM TANZA - BEBANG MEGA INC., etc."),
        ("BEI-TIN stub stores (4 of 49)", "(4 incl)", "ROBINSONS ANTIPOLO / SM MANILA / SM MEGAMALL / SM SOUTHMALL — file under BEI TIN"),
    ]
    t = doc.add_table(rows=len(rows), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j)
            c.text = ""
            p = c.paragraphs[0]
            if i == 0:
                add_run(p, val, bold=True, color=PALETTE['white'], size=10)
                set_cell_shading(c, PALETTE['primary'])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                add_run(p, val, size=10, color=PALETTE['text'])
                set_cell_shading(c, PALETTE['white'] if i % 2 else PALETTE['light'])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 1 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_borders(c, color=PALETTE['border'])
            set_cell_margins(c, top=50, bottom=50, left=100, right=100)
    t.columns[0].width = Inches(2.1)
    t.columns[1].width = Inches(0.8)
    t.columns[2].width = Inches(4.0)

    add_section_header(doc, "ABBR naming convention")
    add_body(doc,
        "Each Frappe Company has a short ABBR used as the suffix on every account name "
        "(e.g. `IN-STORE SALES - BMI2`, `Round Off - BKI`). When Bridge creates the QBO "
        "Companies, use the ABBR consistently as the QBO Company short-name and as the "
        "AccountName suffix to keep the link between Frappe and QBO unambiguous.")

    add_page_break(doc)


def build_section_3(doc):
    add_section_header(doc, "3.  THE CANONICAL CHART OF ACCOUNTS")
    add_body(doc,
        "Every Company shares the SAME structural Chart of Accounts. The differences are "
        "in WHICH sub-tree of the 27-account Sales tree gets postings (i.e. which Income "
        "leaves carry transactional revenue) and in the role-specific Expense classifications.")

    add_section_header(doc, "5-root tree (Asset / Liability / Equity / Income / Expense)")
    add_body(doc,
        "Every Company now carries 5 root group accounts (`Asset - <ABBR>`, `Liability - <ABBR>`, "
        "etc.). These were seeded in Phase 3a of S258 via bench execute with "
        "`frappe.local.flags.ignore_root_company_validation = True`. QBO will infer hierarchy "
        "from the ParentAccount column in the per-Company CSVs — Bridge does not need to "
        "manually re-create the 5-root scaffolding.")

    add_section_header(doc, "Butch's 27-account Sales tree (COA-175-001, locked 2026-04-08)")
    add_body(doc,
        "Butch's locked Sales hierarchy is structurally identical across all 58 Companies. "
        "The number space is `4000000` SALES → 4 sub-trees: `4000100` STORE SALES, `4000200` "
        "BKI SALES, `4000230` FEES, `4000900` DISCOUNTS AND PROMO. See Sheet 4 of "
        "COMPANY_REGISTER.xlsx for the verbatim 27-row table.")

    add_callout_box(doc, [
        ("PER-ROLE POPULATION (which sub-tree gets postings)", True, 11),
        ("Head Office (BEI): 4000234 + 4000235 + 4000005    |    "
         "Commissary (BKI): 4000200 sub-tree only    |    "
         "Franchisor (BFC): 4000230 FEES only    |    "
         "Stores: 4000110 + 4000121 + 4000122 + 4000123    |    "
         "Stubs (ROA/SMM/SMMM/SMS): same as Stores", False, 9),
    ])

    add_section_header(doc, "Two key COA-175 decisions Bridge should know cold")
    add_body(doc,
        "(a) COA-175-002 — DISCOUNTS AND PROMO. All contra-revenue (PWD, Senior Citizen, "
        "promo, employee, refund) lives at `4000900` as a group with 8 leaves `4000901..4000908`. "
        "Bridge should set up QBO discount classes mapped to these exact numbers. The legacy "
        "Apex range `4000201..4000208` (used in early 2026) has been renumbered.")
    add_body(doc,
        "(b) COA-175-011 — BKI is Commissary only. BKI populates ONLY `4000200 BKI SALES` "
        "and its sub-tree. BKI has no in-store sales, no online sales, no franchise income. "
        "If you see Apex postings on BKI to `4000110` or `4000300`, those are migration "
        "artifacts to be cleared (or reclassified to the actual selling Company).")

    add_page_break(doc)


def build_section_4(doc):
    add_section_header(doc, "4.  TAX & REGULATORY SETUP")
    add_body(doc,
        "All 58 Companies are PHP-denominated and operate under Philippine BIR regulation. "
        "Standard VAT is 12% (mid-2025 onward). Below is the canonical tax taxonomy Bridge "
        "should mirror in QBO.")

    add_section_header(doc, "VAT structure (4-stem Input + 1 Output)")
    vat = [
        ("Code", "Account number", "Account name", "QBO suggested DetailType"),
        ("Output VAT", "2102205", "OUTPUT VAT PAYABLE", "SalesTaxPayable"),
        ("Input VAT — Goods", "1105103", "INPUT VAT - GOODS", "OtherCurrentAssets"),
        ("Input VAT — Services", "1105104", "INPUT VAT - SERVICES", "OtherCurrentAssets"),
        ("Input VAT — Importation", "1105105", "INPUT VAT - IMPORTATION", "OtherCurrentAssets"),
        ("Input VAT — Inter-Co", "1106210", "INPUT VAT - BKI INTER-CO", "OtherCurrentAssets"),
    ]
    t = doc.add_table(rows=len(vat), cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(vat):
        for j, val in enumerate(row):
            c = t.cell(i, j)
            c.text = ""
            p = c.paragraphs[0]
            if i == 0:
                add_run(p, val, bold=True, color=PALETTE['white'], size=10)
                set_cell_shading(c, PALETTE['primary'])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                add_run(p, val, size=10, color=PALETTE['text'])
                set_cell_shading(c, PALETTE['white'] if i % 2 else PALETTE['light'])
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j in (2, 3) else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_borders(c, color=PALETTE['border'])
            set_cell_margins(c, top=40, bottom=40, left=80, right=80)
    t.columns[0].width = Inches(1.8)
    t.columns[1].width = Inches(1.2)
    t.columns[2].width = Inches(2.5)
    t.columns[3].width = Inches(1.5)

    add_body(doc,
        "The 4-stem Input VAT split is required by BIR Form 2550Q Schedule 3 — "
        "Purchases / Importations (Goods / Services / Importation as separate columns). "
        "The Inter-Co stem is management-control per Butch's ICT-001. Do NOT consolidate "
        "the four stems in QBO — Bridge will need to file 2550Q quarterly per Company.")

    add_section_header(doc, "EWT codes (TAX-002 — Butch locked 2026-02-19)")
    ewt = [
        ("Category", "Rate", "ATC", "Trigger"),
        ("Professional fees — Individual ≤ 3M", "5%", "WI050 (mgmt/tech) / WI010 (lawyers/CPAs)", "All professional payments"),
        ("Professional fees — Individual > 3M or VAT-reg", "10%", "WI051 / WI011", "All professional payments"),
        ("Professional fees — Corp ≤ 720K", "10%", "WC050 / WC010", "Corporate professional payments"),
        ("Professional fees — Corp > 720K", "15%", "WC051 / WC011", "Corporate professional payments"),
        ("Rentals — Individual lessor", "5%", "WI100", "Lease payments to individual"),
        ("Rentals — Corporate lessor", "5%", "WC100", "Lease payments to corporate"),
    ]
    t2 = doc.add_table(rows=len(ewt), cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(ewt):
        for j, val in enumerate(row):
            c = t2.cell(i, j)
            c.text = ""
            p = c.paragraphs[0]
            if i == 0:
                add_run(p, val, bold=True, color=PALETTE['white'], size=10)
                set_cell_shading(c, PALETTE['primary'])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                add_run(p, val, size=9, color=PALETTE['text'])
                set_cell_shading(c, PALETTE['white'] if i % 2 else PALETTE['light'])
            set_cell_borders(c, color=PALETTE['border'])
            set_cell_margins(c, top=40, bottom=40, left=80, right=80)
    t2.columns[0].width = Inches(2.6)
    t2.columns[1].width = Inches(0.6)
    t2.columns[2].width = Inches(2.3)
    t2.columns[3].width = Inches(1.5)

    add_callout_box(doc, [
        ("BIR Filing Calendar — what each Company files", True, 11),
        ("Monthly: 1601-EQ (EWT) within 10 days after month-end. "
         "Quarterly: 2550Q (VAT) within 25 days after Q-end + 1701Q/1702Q (Income Tax) within 60 days. "
         "Annually: 1604CF/1604E (alphabetical lists) and AFS. "
         "TAX-007 confirms BIR Form 2550M was abolished per RMC 5-2023 — do NOT set up monthly VAT.", False, 9),
    ])

    add_body(doc,
        "Currency: PHP across all 58 Companies. Fiscal year: calendar year (Jan 1 – Dec 31).")
    add_page_break(doc)


def build_section_5(doc):
    add_section_header(doc, "5.  INTERCOMPANY SCHEME")
    add_body(doc,
        "BEI's group has three intercompany flows Bridge needs to mirror exactly. Each is "
        "documented with a verbatim JE pattern below.")

    add_section_header(doc, "(a) Fork 1 — BFC collection-agent model (COA-175-013)")
    add_body(doc,
        "Until BFC's BIR Authority-to-Print (OR booklet) is approved by RDO 044, BEI collects "
        "franchise fees AS AGENT for BFC. The economic substance is BFC's revenue but the "
        "cash lands on BEI's bank. Both sides book intercompany so consolidation eliminates "
        "to zero.")
    add_callout_box(doc, [
        ("JE pattern — Franchisee pays BEI for BFC services", True, 10),
        ("On BEI (agent):    Dr Cash on Hand - BEI    /    Cr 2104200 DUE TO BFC - BEI", False, 9),
        ("On BFC (principal): Dr 1104200 DUE FROM BEI - BFC    /    Cr 4000231 ROYALTY FEES - BFC + 2102205 OUTPUT VAT - BFC", False, 9),
        ("At cutover (BFC bank opens):  BEI sweeps DUE TO BFC balance to BFC's BDO; both intercompany accounts zero out.", False, 9),
    ])

    add_section_header(doc, "(b) BKI→Store paired SI/PI — Butch ICT-001..007")
    add_body(doc,
        "When BKI ships goods to a store, it's an EXTERNAL purchase with 12% VAT (NOT an "
        "internal transfer). Butch's ICT-001 locks this as a paired Sales Invoice (on BKI's "
        "side) + Purchase Invoice (on the buyer Company's side) with 2.75% markup on goods "
        "and 8% markup on services.")
    add_callout_box(doc, [
        ("JE pattern — BKI ships PHP 10,000 of goods (markup 2.75%) to AYALA VERMOSA store", True, 10),
        ("On BKI (seller):      Dr 1103210 AR - AYVER - BKI 11,500    /    Cr 4000210 DELIVERIES - BKI 10,275 + 2102205 OUTPUT VAT - BKI 1,225", False, 9),
        ("On AYVER (buyer):     Dr 1104210 Inventory-from-Commissary - AYVER 10,275 + 1106210 INPUT VAT - BKI Inter-Co 1,225    /    Cr 2103210 AP - BKI Inter-Co - AYVER 11,500", False, 9),
    ])

    add_section_header(doc, "(c) Reliever labor cost-sharing (S206)")
    add_body(doc,
        "When a reliever from Store A covers a shift at Store B, the labor cost is JE'd "
        "between the two per-store Companies using Internal Customers (S206). Internal "
        "Customers exist ONLY for labor JEs — never used for regular Sales Invoices. They "
        "have `is_internal_customer=1` and no TIN (since paired JEs need no BIR receipt).")
    add_page_break(doc)


def build_section_6(doc):
    add_section_header(doc, "6.  APEX → CANONICAL TRANSLATION")
    add_body(doc,
        "APEX (the system Bridge is migrating FROM) uses a flat-tree dialect with different "
        "account names and number conventions than Frappe canonical. S258 Phase 2.0 generated "
        "topologically-sorted migration maps for the 3 Apex-dialect parent Companies "
        "(BEI / BKI / III). Bridge uses these maps when importing historical APEX transactions "
        "into QBO so the transactions land on the right canonical account.")

    add_callout_box(doc, [
        ("REFERENCE FILES in Drive (migration maps as CSV)", True, 11),
        ("migration_map_BEI.csv (330 rows — Head Office Apex→canonical)", False, 10),
        ("migration_map_BKI.csv (325 rows — Commissary Apex→canonical)", False, 10),
        ("migration_map_III.csv (327 rows — Holdco Apex→canonical)", False, 10),
    ])

    add_body(doc,
        "Each map has columns: old_name, old_parent, old_account_number, canonical_name, "
        "canonical_parent, canonical_account_number, canonical_root_type, canonical_account_type, "
        "migration_action, is_group. Rows are topologically sorted (parents before children). "
        "migration_action is one of NOOP (already canonical), RENAME (apply translation), "
        "UNRESOLVED (no canonical match — Bridge decides during sandbox import).")

    add_section_header(doc, "Frappe → QBO type mapping (Appendix F)")
    add_body(doc,
        "Sheet 3 of master_reconciliation.xlsx is the full Frappe-account-type → QBO-AccountType "
        "+ DetailType map (18 explicit + 5 root_type fallback). Apply it after the APEX→canonical "
        "translation. The mapping is best-effort; the QBO sandbox import will surface any "
        "DetailType rejections, then Bridge and BEI iterate.")
    add_page_break(doc)


def build_section_7(doc):
    add_section_header(doc, "7.  5 KNOWN COSMETIC DELTAS")
    add_body(doc,
        "Five non-blocking items remain in Frappe state. Bridge can import the 58 CSVs "
        "as-is and address these during sandbox iteration. Full detail in Sheet 4 of "
        "master_reconciliation.xlsx and in BRIDGE_READINESS_ASSESSMENT.")

    deltas = [
        ("1", "III still carries both old (4000201-206) AND new (4000901-903) discount account_numbers.",
         "LOW — III has 0 GL postings on either."),
        ("2", "BEI STOCK ADJUSTMENT - BEI sits under LOCAL AD & PROMO - BEI Expense group instead of Direct Expenses - BEI.",
         "LOW — pre-existing orphan; no operational impact."),
        ("3", "Some BEI/III accounts retained number prefix in docname (e.g. `4000100 - STORE SALES - BEI`).",
         "COSMETIC — QBO uses AccountNumber + AccountName columns separately; zero impact."),
        ("4", "BFC Fork 1 scaffolding docnames have number prefix.",
         "COSMETIC — same as #3."),
        ("5", "QBO DetailType is best-effort for unmapped Frappe account_types.",
         "KNOWN — sandbox import surfaces rejections; iterate."),
    ]
    rows = [("#", "Delta", "Severity / Action")] + deltas
    t = doc.add_table(rows=len(rows), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    sev_color = {"LOW": "CDD8CE", "COSMETIC": PALETTE['light_gold'],
                 "KNOWN": PALETTE['light_purple']}
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j)
            c.text = ""
            p = c.paragraphs[0]
            if i == 0:
                add_run(p, val, bold=True, color=PALETTE['white'], size=10)
                set_cell_shading(c, PALETTE['primary'])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                add_run(p, val, size=9, color=PALETTE['text'])
                set_cell_shading(c, PALETTE['white'] if i % 2 else PALETTE['light'])
                if j == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif j == 2:
                    # Color-code by severity word at start
                    for sev, color in sev_color.items():
                        if val.startswith(sev):
                            set_cell_shading(c, color)
                            break
            set_cell_borders(c, color=PALETTE['border'])
            set_cell_margins(c, top=50, bottom=50, left=80, right=80)
    t.columns[0].width = Inches(0.4)
    t.columns[1].width = Inches(4.0)
    t.columns[2].width = Inches(2.6)
    add_page_break(doc)


def build_section_8(doc):
    add_section_header(doc, "8.  MASTER DECISION LEDGER")
    add_body(doc,
        "DECISIONS.md is BEI's locked policy ledger. Every COA, tax, accounting policy, "
        "P&L, AR/AP, procurement, and intercompany decision is recorded with the date "
        "Butch (or Sam in supersession) confirmed it. Bridge should treat DECISIONS.md "
        "as the canonical source of truth for any \"why is it set up this way?\" question.")

    add_section_header(doc, "Structure (sections in DECISIONS.md)")
    sections = [
        "Company Structure",
        "Chart of Accounts (COA-001..010 + COA-175-001..030)",
        "Tax Policy (TAX-001..014)",
        "Cost Center",
        "Accounting Policy",
        "AR / Collections",
        "P&L Structure",
        "RBAC & Permissions",
        "Gap Engineering",
        "Billing & Franchise Fees",
        "Inter-Company Transactions (BKI ↔ BEI)",
        "Procurement & AP Controls",
        "Revenue Recognition & SOA",
    ]
    for s in sections:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=10, after=10)
        add_run(p, "  •  ", bold=True, color=PALETTE['primary'], size=11)
        add_run(p, s, size=10, color=PALETTE['text'])

    add_callout_box(doc, [
        ("REFERENCE FILE in Drive", True, 11),
        ("DECISIONS.md (uploaded as Google Doc — comment / search / cross-reference live). "
         "27 COA-175 rows + 10 COA + 14 TAX + ~30 cross-cutting decisions.", False, 10),
    ])
    add_page_break(doc)


def build_section_9(doc):
    add_section_header(doc, "9.  WHAT BRIDGE NEEDS FROM APEX (NOT BEI)")
    add_body(doc,
        "BEI ships Bridge the TARGET COA structure (this pack). Bridge gets the historical "
        "transactional data from APEX directly — the BEI accounting team / Anna R. at Bridge "
        "already has that access from the 2026-06-01 turnover. Below is what Bridge needs to "
        "pull from APEX to populate the new QBO companies.")

    needs = [
        ("Opening Balances (CRITICAL for cutover)",
         "Per-Company trial balance at the cutover date. Source: APEX P&L + Balance Sheet "
         "as of cutover. Map each Apex account → canonical name via migration_map_*.csv, "
         "then load opening balances into QBO."),
        ("AR Aging Snapshot",
         "Per-Company AR aging from APEX as of cutover date. Map customers → QBO Customers "
         "(use the BEI Customer master per Company)."),
        ("AP Aging Snapshot",
         "Per-Company AP aging from APEX. Map suppliers → QBO Vendors. Note that BEI runs "
         "the Procurement AppSheet with payment terms tracked separately — confirm with Anna."),
        ("Bank Balances",
         "Per-Company bank balances from APEX. Match each Apex bank account to canonical "
         "`1101100 - Cash on Hand` / `1101105 - UNIONBANK` / etc."),
        ("Inventory Valuation",
         "BKI commissary + per-store inventory at cutover. Frappe will run Stock Reconciliation "
         "post-cutover; QBO can mirror via Other Current Asset / Inventory category."),
        ("Fixed Asset Schedule",
         "PPE schedule per Company from APEX. Map to `Fixed Asset - <ABBR>` group in QBO."),
    ]
    for title, body in needs:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=30, after=10)
        add_run(p, "  ●  ", bold=True, color=PALETTE['primary'], size=12)
        add_run(p, title, bold=True, size=11, color=PALETTE['text'])
        add_body(doc, body, first_indent=False)

    add_callout_box(doc, [
        ("WHY THIS COMES FROM APEX, NOT BEI", True, 11),
        ("Frappe has structural definitions but no historical operational data (the BKI→store "
         "paired-doc work in early 2026 generated test transactions only, not real ops). "
         "APEX is the system of record for historical BEI transactions through cutover. "
         "Anna R. (anna.r@bridge-ph.com) has APEX access via the 2025 Apex Turnover folder.", False, 9),
    ])
    add_page_break(doc)


def build_section_10(doc):
    add_section_header(doc, "10.  QBO SETUP CHECKLIST — Step-by-Step")
    add_body(doc,
        "Work through these steps in order. Each step links to the relevant reference file "
        "in your Drive folder. Tick the checkboxes as you go — this becomes the audit trail.")

    steps = [
        ("1", "Create 58 QBO Companies",
         "One QBO Company per Frappe Company. Use the ABBR as QBO short-name. Source: "
         "All 58 Companies sheet of COMPANY_REGISTER.xlsx. Set each Company's currency = PHP, "
         "fiscal year = calendar. Use BIR TIN + RDO from the register."),
        ("2", "Import the per-Company COA CSVs",
         "Unzip per_company_coa.zip → 58 CSVs. Import each CSV into the matching QBO Company. "
         "QBO will create accounts respecting the ParentAccount hierarchy. Expected: 6,928 "
         "active accounts total (range per Company: 19 for fresh stubs, ~336 for BEI)."),
        ("3", "Verify Frappe → QBO type map on 3 anchor Companies first",
         "Validate import on BEI (Head Office), BKI (Commissary), BFC (Franchisor) before "
         "doing the other 55. Use Sheet 3 of master_reconciliation.xlsx to confirm "
         "AccountType + DetailType. Surface any DetailType rejections to Sam."),
        ("4", "Set up Tax Codes (VAT 12% + EWT)",
         "Create QBO Tax Codes: VAT-12 (standard), VAT-0 (zero-rated), VAT-Exempt. EWT codes "
         "per Section 4 table (WI050, WI051, WC050, WC051, WI100, WC100). Each Tax Code maps "
         "to the canonical Input/Output VAT account."),
        ("5", "Configure intercompany clearing accounts",
         "Verify the Fork 1 scaffolding is live in QBO: 2104200 DUE TO BFC - BEI on BEI; "
         "1104200 DUE FROM BEI - BFC + 2102205 OUTPUT VAT - BFC on BFC. Mark these as "
         "Intercompany category for QBO consolidation."),
        ("6", "Apply Opening Balances from APEX",
         "Per Section 9 — pull APEX trial balance per Company, translate via migration_map "
         "CSVs, load OBs into QBO using Journal Entry with `Opening Balance Equity` as the "
         "offset. Date OB JEs at the cutover date."),
        ("7", "Import historical transactions using migration maps",
         "From APEX, export transactions per Company. Apply the appropriate "
         "migration_map_BEI.csv / BKI.csv / III.csv to translate account names. For the "
         "49 store Companies + BFC + BFT + 4 stubs, accounts are already canonical (no "
         "translation needed)."),
        ("8", "Reconcile QBO P&L vs Frappe P&L (when Frappe goes live)",
         "Once Frappe is operationally live, run side-by-side P&L for both systems for the "
         "first month. Investigate any line variance > PHP 100 per Company. Sam to chair "
         "the reconciliation review weekly until variance < 1%."),
        ("9", "Sign-off (SIGNOFF.docx)",
         "When all 58 QBO Companies are confirmed mirroring Frappe + opening balances "
         "loaded + first month reconciled, Bridge countersigns SIGNOFF.docx (or its Google "
         "Doc copy). Sam Karazi (CEO) provides BEI sign-off."),
    ]
    for num, title, body in steps:
        add_step(doc, num, title, body_text=body)
        doc.add_paragraph()
    add_page_break(doc)


def build_section_11(doc, links):
    add_section_header(doc, "11.  REFERENCE FILE INDEX — All Drive Links")
    add_body(doc,
        "All files are in the BEI COA Handoff Drive folder. The 4 @bridge-ph.com Editor "
        "accounts (Anna R., Kim C., Flor A., accountant.outsource) have full access.")

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=20, after=20)
    add_run(p, "BEI COA Handoff Drive folder", bold=True, size=11, color=PALETTE['primary'])
    p2 = doc.add_paragraph()
    set_paragraph_spacing(p2, before=0, after=20)
    add_hyperlink(p2, links['folder'], links['folder'], size=10)

    # Group links by category
    groups = [
        ("Native files (authoritative — for QBO import + signing)", [
            ("per_company_coa.zip — 58 CSVs for QBO import", links['per_company_coa_zip']),
            ("COMPANY_REGISTER.xlsx — 58 Companies with role + TIN + RDO", links.get('company_register_xlsx', '(uploaded with this pack)')),
            ("master_reconciliation.xlsx — 5-sheet BEI workbook", links['master_recon_xlsx']),
            ("SIGNOFF.docx — BEI-branded sign-off document", links['signoff_docx']),
            ("coa_export_zip_manifest.csv — SHA-256 + counts", links['manifest_csv']),
            ("upload_manifest.json — file integrity", links['upload_manifest']),
        ]),
        ("Google-format mirrors (for Bridge to comment / edit)", [
            ("master_reconciliation (Google Sheets)", links['master_recon_gs']),
            ("SIGNOFF (Google Doc)", links['signoff_gdoc']),
            ("validation (Google Doc) — import readiness", links['validation_gdoc']),
            ("BRIDGE_READINESS_ASSESSMENT (Google Doc) — 35-point fact-check", links['bra_gdoc']),
            ("WALKTHROUGH (this document, as Google Doc)", links.get('walkthrough_gdoc', '(uploaded next)')),
        ]),
        ("APEX → Canonical Migration Maps (for translating Apex transactions)", [
            ("migration_map_BEI.csv (330 rows, Head Office)", links.get('migration_BEI', '(uploaded with this pack)')),
            ("migration_map_BKI.csv (325 rows, Commissary)", links.get('migration_BKI', '(uploaded with this pack)')),
            ("migration_map_III.csv (327 rows, Holdco)", links.get('migration_III', '(uploaded with this pack)')),
        ]),
    ]
    for group_title, items in groups:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=200, after=80)
        add_run(p, group_title, bold=True, size=11, color=PALETTE['secondary'])
        for label, url in items:
            p2 = doc.add_paragraph()
            set_paragraph_spacing(p2, before=20, after=20)
            add_run(p2, "  •  ", color=PALETTE['primary'])
            add_run(p2, f"{label}: ", size=10, color=PALETTE['text'])
            if url.startswith("http"):
                add_hyperlink(p2, url, url, size=9)
            else:
                add_run(p2, url, size=9, italic=True, color=PALETTE['text_light'])
    add_page_break(doc)


def build_section_12(doc):
    add_section_header(doc, "12.  CONTACT DIRECTORY")

    add_section_header(doc, "Bridge Consulting (Editor access on Drive folder)")
    bridge = [
        ("Anna R.", "anna.r@bridge-ph.com", "Primary contact — Apex turnover lead"),
        ("Kim C.", "kim.c@bridge-ph.com", "QBO implementation"),
        ("Flor A.", "flor.a@bridge-ph.com", "Accounting / tax compliance"),
        ("Outsource", "accountant.outsource@bridge-ph.com", "Bookkeeping / data entry"),
    ]
    for name, email, role in bridge:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=20, after=20)
        add_run(p, f"  •  {name}: ", bold=True, color=PALETTE['primary'], size=10)
        add_run(p, email, color=PALETTE['text'], size=10)
        add_run(p, f"  ({role})", italic=True, color=PALETTE['text_light'], size=10)

    add_section_header(doc, "BEI (Editor access on Drive folder)")
    bei = [
        ("Sam Karazi", "sam@bebang.ph", "CEO — escalation + final sign-off"),
        ("Denise", "denise@bebang.ph", "Finance lead (post-Butch)"),
        ("Anthony", "anthony@bebang.ph", "Accounting operations"),
        ("Sheena", "sheena@bebang.ph", "AP / supplier admin"),
        ("Drew", "drew@bebang.ph", "BIR compliance"),
    ]
    for name, email, role in bei:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=20, after=20)
        add_run(p, f"  •  {name}: ", bold=True, color=PALETTE['primary'], size=10)
        add_run(p, email, color=PALETTE['text'], size=10)
        add_run(p, f"  ({role})", italic=True, color=PALETTE['text_light'], size=10)

    add_callout_box(doc, [
        ("ESCALATION", True, 11),
        ("Sandbox import question → Anthony or Denise. Tax / regulatory → Drew. Policy / "
         "supersession decision → Sam. CFO seat is vacant indefinitely (Butch resigned "
         "2026-04-15); Sam Karazi is single-owner authority for all post-resignation locks.", False, 10),
    ])
    add_page_break(doc)


def build_section_13(doc):
    add_section_header(doc, "13.  SIGN-OFF")
    add_body(doc,
        "By signing below, Bebang Enterprise Inc. (BEI) releases this BEI ERP → "
        "QuickBooks Online Setup Walkthrough Pack to Bridge Consulting for use in the QBO "
        "migration. Bridge Consulting acknowledges receipt of this document and the "
        "supporting reference files in the BEI COA Handoff Drive folder.")

    sigt = doc.add_table(rows=1, cols=2)
    sigt.alignment = WD_TABLE_ALIGNMENT.CENTER

    left = sigt.cell(0, 0)
    left.text = ""
    p1 = left.paragraphs[0]
    add_run(p1, "RELEASED BY (BEI):", bold=True, size=10, color=PALETTE['primary'])
    left.add_paragraph()
    left.add_paragraph()
    line1 = left.add_paragraph()
    add_run(line1, "_____________________________________", size=10)
    n1 = left.add_paragraph()
    add_run(n1, "Sam Karazi", bold=True, size=11, color=PALETTE['text'])
    t1 = left.add_paragraph()
    add_run(t1, "Chief Executive Officer", size=9, color=PALETTE['text_light'])
    cn1 = left.add_paragraph()
    add_run(cn1, "Bebang Enterprise Inc.", italic=True, size=9, color=PALETTE['text_light'])
    d1 = left.add_paragraph()
    add_run(d1, "Date: ______________________", size=10, color=PALETTE['text'])
    set_cell_shading(left, PALETTE['light'])
    set_cell_borders(left, color=PALETTE['secondary'], size="12")
    set_cell_margins(left, top=200, bottom=200, left=200, right=200)

    right = sigt.cell(0, 1)
    right.text = ""
    p2 = right.paragraphs[0]
    add_run(p2, "ACKNOWLEDGED BY (BRIDGE):", bold=True, size=10, color=PALETTE['primary'])
    right.add_paragraph()
    right.add_paragraph()
    line2 = right.add_paragraph()
    add_run(line2, "_____________________________________", size=10)
    n2 = right.add_paragraph()
    add_run(n2, "_______________________________", bold=True, size=11, color=PALETTE['text'])
    t2 = right.add_paragraph()
    add_run(t2, "(Print Name & Title)", size=9, color=PALETTE['text_light'])
    cn2 = right.add_paragraph()
    add_run(cn2, "Bridge Consulting", italic=True, size=9, color=PALETTE['text_light'])
    d2 = right.add_paragraph()
    add_run(d2, "Date: ______________________", size=10, color=PALETTE['text'])
    set_cell_shading(right, PALETTE['light'])
    set_cell_borders(right, color=PALETTE['secondary'], size="12")
    set_cell_margins(right, top=200, bottom=200, left=200, right=200)

    sigt.columns[0].width = Inches(3.5)
    sigt.columns[1].width = Inches(3.5)

    # Footer
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=400, after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "— END OF WALKTHROUGH —", italic=True, size=9, color=PALETTE['text_light'])

    p2 = doc.add_paragraph()
    set_paragraph_spacing(p2, before=40, after=0)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "Bebang Enterprise Inc. — Confidential — For Bridge Consulting use only",
            size=8, color=PALETTE['text_light'])

    p3 = doc.add_paragraph()
    set_paragraph_spacing(p3, before=20, after=0)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p3, "Generated 2026-06-05 from S258 sprint closeout — see hrms PR #770",
            size=8, color=PALETTE['text_light'])


def main():
    sys.stdout.reconfigure(encoding="utf-8")

    # Drive links from previous upload + new ones we'll fill at upload time
    import json
    try:
        idx = json.load(open("output/s258/bridge_handoff/drive_upload_index.json"))
        file_urls = {f['name']: f['url'] for f in idx.get('files', [])}
    except Exception:
        file_urls = {}
    folder_url = "https://drive.google.com/drive/folders/1GnrFKICFYN6xz9IKeAFCM0xYtFN02OcE"

    links = {
        'folder': folder_url,
        'per_company_coa_zip': file_urls.get('per_company_coa.zip', folder_url),
        'master_recon_xlsx': file_urls.get('master_reconciliation.xlsx', folder_url),
        'signoff_docx': file_urls.get('SIGNOFF.docx', folder_url),
        'manifest_csv': file_urls.get('coa_export_zip_manifest.csv', folder_url),
        'upload_manifest': file_urls.get('upload_manifest.json', folder_url),
        'master_recon_gs': file_urls.get('master_reconciliation (Google Sheets)', folder_url),
        'signoff_gdoc': file_urls.get('SIGNOFF (Google Doc)', folder_url),
        'validation_gdoc': file_urls.get('validation (Google Doc)', folder_url),
        'bra_gdoc': file_urls.get('BRIDGE_READINESS_ASSESSMENT (Google Doc)', folder_url),
    }

    doc = Document()
    setup_doc(doc)
    build_cover(doc)
    build_toc(doc)
    build_section_1(doc)
    build_section_2(doc)
    build_section_3(doc)
    build_section_4(doc)
    build_section_5(doc)
    build_section_6(doc)
    build_section_7(doc)
    build_section_8(doc)
    build_section_9(doc)
    build_section_10(doc)
    build_section_11(doc, links)
    build_section_12(doc)
    build_section_13(doc)

    out = Path("output/s258/bridge_handoff/WALKTHROUGH.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"[OK] Wrote {out}")


if __name__ == "__main__":
    main()
