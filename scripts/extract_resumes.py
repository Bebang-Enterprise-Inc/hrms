"""Extract text from all resume files (PDF, DOCX, DOC, TXT) using benchmark-proven routing."""
import os
import json
import asyncio
import subprocess
import sys
from pathlib import Path
from datetime import datetime

# === Config ===
BASE = Path("F:/Dropbox/Projects/BEI-ERP/recruitment")
JOBS = [
    ("head-of-finance-and-accounting-controller", "Head of Finance"),
    ("accounting-manager", "Accounting Manager"),
]
GEMINI_KEY = subprocess.check_output(
    [r"C:\Users\Sam\bin\doppler.exe", "secrets", "get", "GEMINI_API_KEY", "--plain",
     "--project", "bei-erp", "--config", "dev"],
    text=True
).strip()

# === Extractors ===

def extract_pdf_pymupdf(path: Path) -> str:
    """Fast local PDF extraction via PyMuPDF."""
    import fitz
    doc = fitz.open(str(path))
    text = []
    for page in doc:
        text.append(page.get_text())
    doc.close()
    return "\n".join(text).strip()


async def extract_pdf_gemini(path: Path, semaphore: asyncio.Semaphore) -> str:
    """PDF via Gemini 2.5 Flash — best quality for PDFs."""
    async with semaphore:
        from google import genai
        from google.genai import types

        client = genai.Client(api_key=GEMINI_KEY)

        file_bytes = path.read_bytes()
        mime = "application/pdf"

        try:
            response = client.models.generate_content(
                model="gemini-2.5-flash-preview-05-20",
                contents=[
                    "Extract ALL text from this resume/CV document. Preserve structure: headings, bullet points, dates, job titles, companies, education, skills. Output as clean text, no commentary.",
                    types.Part.from_bytes(data=file_bytes, mime_type=mime),
                ],
            )
            return response.text.strip() if response.text else ""
        except Exception as e:
            # Fallback to PyMuPDF
            try:
                return extract_pdf_pymupdf(path)
            except:
                return f"[EXTRACTION ERROR: {e}]"


def extract_docx(path: Path) -> str:
    """DOCX via python-docx (Docling fallback)."""
    try:
        from docx import Document
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also get tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs).strip()
    except Exception as e:
        return f"[DOCX EXTRACTION ERROR: {e}]"


def extract_doc(path: Path) -> str:
    """Legacy DOC via PyMuPDF."""
    try:
        import fitz
        doc = fitz.open(str(path))
        text = []
        for page in doc:
            text.append(page.get_text())
        doc.close()
        return "\n".join(text).strip()
    except Exception as e:
        return f"[DOC EXTRACTION ERROR: {e}]"


def extract_txt(path: Path) -> str:
    """Plain text."""
    try:
        return path.read_text(encoding="utf-8", errors="replace").strip()
    except:
        return path.read_text(encoding="latin-1", errors="replace").strip()


# === Main Pipeline ===

async def process_file(path: Path, semaphore: asyncio.Semaphore) -> dict:
    """Route file to the right extractor and return result."""
    ext = path.suffix.lower()
    start = datetime.now()
    text = ""
    method = ""

    try:
        if ext == ".pdf":
            # Use PyMuPDF first (fast, local) — if result is too short, use Gemini
            text = extract_pdf_pymupdf(path)
            method = "pymupdf"
            if len(text) < 100:
                text = await extract_pdf_gemini(path, semaphore)
                method = "gemini-2.5-flash"
        elif ext == ".docx":
            text = extract_docx(path)
            method = "python-docx"
        elif ext == ".doc":
            text = extract_doc(path)
            method = "pymupdf"
        elif ext == ".txt":
            text = extract_txt(path)
            method = "plaintext"
        else:
            text = f"[UNSUPPORTED FORMAT: {ext}]"
            method = "skip"
    except Exception as e:
        text = f"[ERROR: {e}]"
        method = "error"

    elapsed = (datetime.now() - start).total_seconds()
    return {
        "file": path.name,
        "method": method,
        "chars": len(text),
        "seconds": round(elapsed, 2),
        "text": text,
    }


async def process_wave(files: list[Path], semaphore: asyncio.Semaphore) -> list[dict]:
    """Process a wave of files concurrently."""
    tasks = [process_file(f, semaphore) for f in files]
    return await asyncio.gather(*tasks)


async def main():
    semaphore = asyncio.Semaphore(8)  # Max 8 concurrent Gemini calls
    wave_size = 16

    for folder, job_name in JOBS:
        resume_dir = BASE / folder / "resumes"
        output_dir = BASE / folder / "extracted"
        output_dir.mkdir(exist_ok=True)

        # Get all files
        files = sorted([f for f in resume_dir.iterdir() if f.is_file()])
        print(f"\n{'='*60}")
        print(f"  {job_name}: {len(files)} files")
        print(f"{'='*60}")

        # Check manifest for already-extracted
        manifest_path = output_dir / "_manifest.json"
        manifest = {}
        if manifest_path.exists():
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))

        to_process = [f for f in files if f.name not in manifest]
        print(f"  Already extracted: {len(manifest)}, remaining: {len(to_process)}")

        all_results = []
        for wave_start in range(0, len(to_process), wave_size):
            wave = to_process[wave_start:wave_start + wave_size]
            print(f"\n  Wave {wave_start // wave_size + 1}: {len(wave)} files...")

            results = await process_wave(wave, semaphore)

            for r in results:
                # Save individual extracted text
                safe_name = Path(r["file"]).stem + ".txt"
                (output_dir / safe_name).write_text(r["text"], encoding="utf-8")

                # Update manifest
                manifest[r["file"]] = {
                    "method": r["method"],
                    "chars": r["chars"],
                    "seconds": r["seconds"],
                    "output": safe_name,
                }
                all_results.append(r)

                status = "OK" if r["chars"] > 50 else "WEAK"
                print(f"    {r['file'][:50]:50s} {r['method']:15s} {r['chars']:6d} chars  {r['seconds']:5.1f}s  {status}")

            # Save manifest after each wave
            manifest_path.write_text(json.dumps(manifest, indent=2, ensure_ascii=False), encoding="utf-8")

            # Brief pause between waves
            await asyncio.sleep(1)

        # Summary
        total_chars = sum(r["chars"] for r in all_results)
        total_time = sum(r["seconds"] for r in all_results)
        weak = sum(1 for r in all_results if r["chars"] < 50)
        print(f"\n  {job_name} DONE: {len(all_results)} extracted, {total_chars:,} total chars, {total_time:.1f}s total, {weak} weak")
        print(f"  Output: {output_dir}")

    print("\n=== ALL DONE ===")


if __name__ == "__main__":
    asyncio.run(main())
