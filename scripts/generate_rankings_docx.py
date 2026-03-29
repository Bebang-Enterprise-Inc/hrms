"""Generate BEI Candidate Rankings DOCX with per-candidate summaries."""
import json, sys
from pathlib import Path
from datetime import datetime

sys.path.insert(0, r"F:\Dropbox\Projects\BEI-ERP\.claude\skills\docx-designer-bei-erp\scripts")

from bei_docx import (
    PALETTE, LOGO_PATH, setup_doc, add_bei_header, add_section_header,
    add_run, add_body, add_page_break, set_cell_shading, set_cell_borders,
    set_cell_margins, hex_rgb, format_php,
)
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

BASE = Path("F:/Dropbox/Projects/BEI-ERP/recruitment")

def add_score_bar(cell, score, max_score=10):
    """Add a visual score indicator."""
    filled = int(score)
    bar = "█" * filled + "░" * (max_score - filled)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{score}")
    run.font.size = Pt(9)
    run.font.bold = True
    if score >= 8:
        run.font.color.rgb = RGBColor(0x04, 0x40, 0x0A)  # Green
    elif score >= 5:
        run.font.color.rgb = RGBColor(0xC8, 0x90, 0x0A)  # Gold
    else:
        run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)  # Red

def style_header_row(row, bg_color="04400A"):
    """Style a table header row."""
    for cell in row.cells:
        set_cell_shading(cell, bg_color)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(8)

def add_candidate_detail(doc, rank, c, job_type):
    """Add a detailed candidate card."""
    ai = c.get('ai_scores', {})
    name = c.get('name', 'Unknown')
    cpa = "CPA" if c.get('is_cpa') else "Not CPA"
    fit = c.get('fit', 'N/A') or 'N/A'
    interviewed = c.get('already_interviewed', False)

    # Candidate header
    p = doc.add_paragraph()
    run = p.add_run(f"#{rank}  ")
    run.font.size = Pt(14)
    run.font.color.rgb = hex_rgb(PALETTE["accent"])
    run.font.bold = True
    run = p.add_run(name)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = hex_rgb(PALETTE["primary"])
    if interviewed:
        run = p.add_run("  [ALREADY INTERVIEWED]")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
        run.font.italic = True

    # Info table (2 cols)
    info = doc.add_table(rows=4, cols=2)
    info.columns[0].width = Inches(3.0)
    info.columns[1].width = Inches(3.5)

    fields = [
        ("Current Role", f"{c.get('role', 'N/A')}"),
        ("Company", f"{c.get('company', 'N/A')}"),
        ("Tenure", f"{c.get('months', 0) // 12} years {c.get('months', 0) % 12} months"),
        ("CPA / Fit Level", f"{cpa}  |  {fit}"),
    ]
    for i, (label, value) in enumerate(fields):
        cell_l = info.cell(i, 0)
        cell_l.paragraphs[0].clear()
        run = cell_l.paragraphs[0].add_run(label + ": ")
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = hex_rgb("555555")
        run = cell_l.paragraphs[0].add_run(value)
        run.font.size = Pt(9)

    # Contact in second column
    info.cell(0, 1).paragraphs[0].clear()
    run = info.cell(0, 1).paragraphs[0].add_run(f"Email: {c.get('email', 'N/A')}")
    run.font.size = Pt(9)
    info.cell(1, 1).paragraphs[0].clear()
    run = info.cell(1, 1).paragraphs[0].add_run(f"Phone: {c.get('phone', 'N/A')}")
    run.font.size = Pt(9)
    info.cell(2, 1).paragraphs[0].clear()
    run = info.cell(2, 1).paragraphs[0].add_run(f"Location: {c.get('location', 'N/A')}")
    run.font.size = Pt(9)
    info.cell(3, 1).paragraphs[0].clear()
    run = info.cell(3, 1).paragraphs[0].add_run(f"Final Score: {c.get('final_score', 0)}/110")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = hex_rgb(PALETTE["primary"])

    # Remove borders from info table
    for row in info.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = tcPr.find(qn('w:tcBorders'))
            if borders is not None:
                tcPr.remove(borders)
            set_cell_margins(cell, top=20, bottom=20)

    # Score breakdown table
    score_table = doc.add_table(rows=2, cols=7)
    headers = ["Leadership", "QSR/Retail", "Multi-Entity", "ERP/Systems", "PH Compliance", "Career", "TOTAL"]
    keys = ["leadership", "qsr_retail", "multi_entity", "erp_systems", "ph_compliance", "career_trajectory"]

    for j, h in enumerate(headers):
        cell = score_table.cell(0, j)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.font.size = Pt(7)
        run.font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    style_header_row(score_table.rows[0], "04400A")

    for j, k in enumerate(keys):
        add_score_bar(score_table.cell(1, j), ai.get(k, 0))
    # Total
    total_cell = score_table.cell(1, 6)
    total_cell.paragraphs[0].clear()
    total_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = total_cell.paragraphs[0].add_run(f"{ai.get('total', 0)}/60")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = hex_rgb(PALETTE["primary"])
    set_cell_shading(total_cell, "E8F5E9")

    for row in score_table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=40, bottom=40, left=40, right=40)

    # AI Summary
    summary = ai.get('summary', 'No AI summary available.')
    if summary and isinstance(summary, str) and len(summary) > 10:
        p = doc.add_paragraph()
        run = p.add_run("Assessment: ")
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = hex_rgb(PALETTE["accent"])
        run = p.add_run(summary)
        run.font.size = Pt(9)
        run.font.italic = True

    # Separator
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("─" * 85)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def main():
    doc = Document()
    setup_doc(doc)
    add_bei_header(doc)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("CANDIDATE RANKINGS REPORT")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = hex_rgb(PALETTE["primary"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CFO Replacement Program — March 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = hex_rgb(PALETTE["accent"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y %I:%M %p')} PHT")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Methodology section
    add_section_header(doc, "METHODOLOGY")
    add_body(doc, "274 candidates across two job postings were evaluated using a two-phase scoring system:")
    add_body(doc, "Phase 1 — Deterministic Filter (max 50 points): CPA qualification (20), relevant experience (15), SEEK AI fit level (10), industry match (5), and current role tenure (5). This phase uses structured data from screening questions with no subjective judgment.")
    add_body(doc, "Phase 2 — AI Assessment (max 60 points): Each candidate's full resume was evaluated by Gemini AI against six BEI-specific criteria: Leadership presence (10), QSR/Retail experience (10), Multi-entity accounting (10), ERP/Systems knowledge (10), Philippine compliance expertise (10), and Career trajectory (10).")
    add_body(doc, "Total possible score: 110 points. Candidates already interviewed by Ronald Caringal are marked but included for comparison.")

    # Context box
    p = doc.add_paragraph()
    run = p.add_run("BEI Context: ")
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = hex_rgb(PALETTE["accent"])
    run = p.add_run("47+ QSR stores (Bebang Halo-Halo), two legal entities (BEI + BKI), mid-migration to Frappe ERPNext, replacing CFO who resigned March 2026. Priority: strong executive leadership personality, CPA, retail/F&B experience, Philippine tax compliance.")
    run.font.size = Pt(9)
    run.font.italic = True

    # Already interviewed note
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.font.bold = True
    run.font.size = Pt(9)
    run = p.add_run("Ronald's 4 interviewed candidates (Jovelynne Tamayo, Maria Echevarria, Dolores Mejia, Eleaser Calayag) were excluded from this ranking. Sam's assessment: \"I'm not sold with the candidates. If I had to choose one it would be Mejia from Ajinomoto\" — but Mejia is not a CPA.")
    run.font.size = Pt(9)
    run.font.italic = True

    # Process both jobs
    ALREADY_INTERVIEWED = {'Jovelynne Tamayo', 'Maria Echevarria', 'Dolores Mejia', 'Eleaser Calayag'}

    for folder, job_type, top_n in [
        ('head-of-finance-and-accounting-controller', 'Head of Finance and Accounting (Controller)', 10),
        ('accounting-manager', 'Accounting Manager', 10),
    ]:
        add_page_break(doc)

        data = json.loads((BASE / folder / 'rankings.json').read_text(encoding='utf-8'))
        ranked = data['rankings']
        # Remove already interviewed
        filtered = [r for r in ranked if r.get('name', '') not in ALREADY_INTERVIEWED]

        add_section_header(doc, f"TOP {top_n} — {job_type.upper()}")

        total = data.get('total_candidates', len(ranked))
        qualified = data.get('qualified', len(ranked))
        p = doc.add_paragraph()
        run = p.add_run(f"Total applicants: {total}  |  Qualified after Phase 1: {qualified}  |  AI-scored: {len(ranked)}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Summary table
        summary_table = doc.add_table(rows=top_n + 1, cols=6)
        summary_table.columns[0].width = Inches(0.3)
        summary_table.columns[1].width = Inches(1.8)
        summary_table.columns[2].width = Inches(1.8)
        summary_table.columns[3].width = Inches(0.5)
        summary_table.columns[4].width = Inches(0.5)
        summary_table.columns[5].width = Inches(1.3)

        headers = ["#", "Name", "Current Role", "CPA", "Score", "Top Strength"]
        for j, h in enumerate(headers):
            cell = summary_table.cell(0, j)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(h)
            run.font.size = Pt(8)
        style_header_row(summary_table.rows[0])

        for i, c in enumerate(filtered[:top_n]):
            ai = c.get('ai_scores', {})
            row = summary_table.rows[i + 1]

            vals = [
                str(i + 1),
                c.get('name', ''),
                f"{c.get('role', '')} at {c.get('company', '')}",
                "Yes" if c.get('is_cpa') else "No",
                str(c.get('final_score', 0)),
                _top_strength(ai),
            ]
            for j, v in enumerate(vals):
                cell = row.cells[j]
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run(v[:60])
                run.font.size = Pt(8)
                if j == 1:
                    run.font.bold = True
                if j == 4:
                    run.font.bold = True
                    score = c.get('final_score', 0)
                    if score >= 100:
                        run.font.color.rgb = RGBColor(0x04, 0x40, 0x0A)
                    elif score >= 90:
                        run.font.color.rgb = RGBColor(0xC8, 0x90, 0x0A)

            # Alternate row shading
            if i % 2 == 0:
                for cell in row.cells:
                    set_cell_shading(cell, "F5F5F5")

        for row in summary_table.rows:
            for cell in row.cells:
                set_cell_margins(cell, top=30, bottom=30, left=40, right=40)

        # Detailed cards
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        run = p.add_run("DETAILED CANDIDATE PROFILES")
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = hex_rgb(PALETTE["accent"])

        for i, c in enumerate(filtered[:top_n]):
            add_candidate_detail(doc, i + 1, c, job_type)

    # Save
    output_path = BASE / "BEI_Candidate_Rankings_2026-03-29.docx"
    doc.save(str(output_path))
    print(f"Saved: {output_path}")
    return output_path


def _top_strength(ai):
    """Get the highest scoring criteria."""
    if not ai or ai.get('total', 0) == 0:
        return "N/A"
    keys = {
        'leadership': 'Leadership',
        'qsr_retail': 'QSR/Retail',
        'multi_entity': 'Multi-Entity',
        'erp_systems': 'ERP/Systems',
        'ph_compliance': 'PH Compliance',
        'career_trajectory': 'Career',
    }
    best_key = max(keys.keys(), key=lambda k: ai.get(k, 0))
    return f"{keys[best_key]} ({ai.get(best_key, 0)}/10)"


if __name__ == '__main__':
    main()
